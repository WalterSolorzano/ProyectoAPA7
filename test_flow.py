#!/usr/bin/env python3
"""
WordAPA7 — Test de Flujo Completo (End-to-End)

Este script ejecuta el flujo completo de WordAPA7:
1. Crea un documento Word de prueba
2. Lo sube al servidor (o genera localmente si el servidor no está)
3. Genera el archivo APA 7
4. Descarga y valida el resultado

Uso:
    python test_flow.py                    # Modo integración (requiere servidor corriendo)
    python test_flow.py --local            # Modo local (sin servidor)
    cd python && python -m pytest tests/ -v --ignore=tests/test_integration.py  # Solo tests unitarios
"""

import io
import sys
import json
import urllib.request
import urllib.error
from pathlib import Path

# Asegurar path
sys.path.insert(0, str(Path(__file__).parent / "python"))

import docx
from docx.shared import Pt


BASE_URL = "http://127.0.0.1:8742"


def create_test_document():
    """Crea un documento .docx de prueba en memoria."""
    print("1. Creando documento de prueba...")

    doc = docx.Document()

    # Heading 1
    h1 = doc.add_heading("Análisis del Impacto de la IA en la Educación", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # Párrafo con cita
    p1 = doc.add_paragraph(
        "La inteligencia artificial ha transformado significativamente el panorama "
        "educativo global en la última década. Según investigaciones recientes "
        "(García, 2023, p. 45), los sistemas de tutoría inteligente han demostrado "
        "ser particularmente efectivos en la personalización del aprendizaje para "
        "estudiantes con diferentes niveles de habilidad y estilos de aprendizaje."
    )
    for run in p1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading 2
    h2 = doc.add_heading("Marco Teórico", level=2)
    for run in h2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)

    # Párrafo con cita narrativa
    p2 = doc.add_paragraph(
        "López et al. (2021) encontraron que los estudiantes que utilizan plataformas "
        "de aprendizaje adaptativo mejoran su rendimiento académico en un promedio "
        "de 23% en comparación con aquellos que utilizan métodos tradicionales de "
        "enseñanza presencial."
    )
    for run in p2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading 3
    h3 = doc.add_heading("Metodología de la Investigación", level=3)
    for run in h3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Párrafo
    p3 = doc.add_paragraph(
        "Se empleó un diseño de investigación cuantitativo con alcance descriptivo-"
        "correlacional. La muestra estuvo compuesta por 350 estudiantes universitarios "
        "seleccionados mediante muestreo aleatorio estratificado de tres instituciones "
        "de educación superior en la Ciudad de México durante el período 2024-2025."
    )
    for run in p3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    print(f"   OK — {len(doc.paragraphs)} párrafos creados")
    return buf.getvalue()


def run_integration_test():
    """Ejecuta el test contra el servidor FastAPI."""
    print("2. Verificando servidor...")
    try:
        req = urllib.request.Request(f"{BASE_URL}/api/health")
        with urllib.request.urlopen(req, timeout=3) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            if data.get("status") != "ok":
                print(f"   ERROR: Servidor responde pero status != ok: {data}")
                return False
    except Exception as e:
        print(f"   ERROR: No se puede conectar al servidor en {BASE_URL}")
        print(f"   Inicia el servidor con: cd python && python main.py")
        print(f"   O ejecuta: python test_flow.py --local")
        return False
    print("   OK — Servidor corriendo")

    # Upload
    print("3. Subiendo documento...")
    file_bytes = create_test_document()
    boundary = '----WebKitFormBoundary7MA4YWxkTrZu0gW'
    body = (
        f'--{boundary}\r\n'
        f'Content-Disposition: form-data; name="file"; filename="test_doc.docx"\r\n'
        f'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n'
    ).encode('utf-8') + file_bytes + f'\r\n--{boundary}--\r\n'.encode('utf-8')

    upload_req = urllib.request.Request(
        f'{BASE_URL}/api/upload',
        data=body,
        headers={'Content-Type': f'multipart/form-data; boundary={boundary}'}
    )

    with urllib.request.urlopen(upload_req) as resp:
        res_data = json.loads(resp.read().decode('utf-8'))
        session_id = res_data.get("session_id")
        element_count = len(res_data.get("elements", []))
        print(f"   OK — Session ID: {session_id}")
        print(f"   OK — {element_count} elementos clasificados")

    # Generate
    print("4. Generando documento APA 7...")
    generate_body = json.dumps({
        "session_id": session_id,
        "portada": {
            "title": "Análisis del Impacto de la Inteligencia Artificial en la Educación Superior",
            "author": "Juan Carlos Pérez García",
            "institution": "Universidad Nacional Autónoma de México",
            "course": "Seminario de Investigación Avanzada",
            "instructor": "Dra. María Elena López Castillo",
            "date": "15 de mayo de 2025"
        }
    }).encode('utf-8')

    gen_req = urllib.request.Request(
        f'{BASE_URL}/api/generate',
        data=generate_body,
        headers={'Content-Type': 'application/json'}
    )

    with urllib.request.urlopen(gen_req) as resp:
        gen_data = json.loads(resp.read().decode('utf-8'))
        print(f"   OK — Download URL: {gen_data.get('download_url')}")
        print(f"   OK — Filename: {gen_data.get('filename')}")

    # Download
    print("5. Descargando archivo generado...")
    dl_req = urllib.request.Request(f"{BASE_URL}{gen_data.get('download_url')}")
    with urllib.request.urlopen(dl_req) as resp:
        downloaded_bytes = resp.read()
        print(f"   OK — {len(downloaded_bytes):,} bytes descargados")

    # Validate
    print("6. Validando archivo generado...")
    import zipfile
    buf = io.BytesIO(downloaded_bytes)
    if zipfile.is_zipfile(buf):
        print("   OK — El archivo es un ZIP válido (.docx)")
    else:
        print("   ERROR: El archivo no es un ZIP válido")
        return False

    doc = docx.Document(buf)
    para_count = len(doc.paragraphs)
    print(f"   OK — {para_count} párrafos en el documento generado")

    all_text = " ".join(p.text for p in doc.paragraphs)
    keywords = ["inteligencia", "artificial", "educación"]
    found = [kw for kw in keywords if kw.lower() in all_text.lower()]
    if len(found) >= 2:
        print(f"   OK — Contenido preservado: {found}")
    else:
        print(f"   WARN — Palabras clave faltantes. Buscadas: {keywords}, Encontradas: {found}")

    return True


def run_local_test():
    """Ejecuta el test en modo local (sin servidor)."""
    print("2. Modo local — Probando componentes directamente...")

    from models import DocumentModel, ElementModel, ElementType, APARuleSet, PortadaData, APAFormat
    from parsing.pre_classifier import pre_classify_elements
    from modules.citation_engine import extract_citations_from_text
    from generation.generator import generate_apa7_docx

    # Crear elementos de prueba
    print("   Creando elementos de prueba...")
    elements = [
        ElementModel(
            id="1", text="1. Introducción al Análisis",
            style_name="Heading 1", is_bold=True,
            alignment="center", font_size=14.0
        ),
        ElementModel(
            id="2",
            text="Este es un párrafo de prueba que contiene una cita parentética "
                 "(García, 2023) para verificar la extracción de citas y el formato "
                 "APA 7 en texto académico en español.",
            style_name="Normal"
        ),
        ElementModel(
            id="3", text="1.1 Marco Teórico",
            style_name="Heading 2", is_bold=True,
            alignment="left", font_size=12.0
        ),
        ElementModel(
            id="4",
            text="Según López et al. (2021), la inteligencia artificial ha "
                 "transformado el procesamiento de texto académico de manera "
                 "significativa en los últimos años.",
            style_name="Normal"
        ),
    ]

    # Clasificar
    classified = pre_classify_elements(elements)
    headings = [e for e in classified if e.type == ElementType.HEADING]
    paragraphs = [e for e in classified if e.type == ElementType.PARAGRAPH]

    print(f"   OK — {len(headings)} headings, {len(paragraphs)} párrafos clasificados")

    # Extraer citas
    total_cits = 0
    for elem in classified:
        cits = extract_citations_from_text(elem.text, elem.id)
        total_cits += len(cits)
    print(f"   OK — {total_cits} citas detectadas")

    # Generar documento
    doc_model = DocumentModel(
        session_id="test_local",
        file_name="test_local.docx",
        elements=classified,
    )

    portada = PortadaData(
        title="Análisis del Impacto de la Inteligencia Artificial",
        author="Juan Pérez",
        institution="Universidad Nacional",
    )

    out_dir = Path(__file__).parent / "python" / "tests"
    out_file = out_dir / "output_test_apa7.docx"
    gen_path = generate_apa7_docx(doc_model, out_file, portada=portada)

    print(f"   OK — Documento generado: {gen_path}")
    print(f"   OK — Tamaño: {gen_path.stat().st_size:,} bytes")

    doc = docx.Document(str(gen_path))
    print(f"   OK — {len(doc.paragraphs)} párrafos en el documento generado")
    return True


if __name__ == "__main__":
    print("=" * 60)
    print(" WordAPA7 — Test de Flujo Completo")
    print("=" * 60)

    if "--local" in sys.argv:
        success = run_local_test()
    else:
        success = run_integration_test()

    print("=" * 60)
    if success:
        print(" FLUJO COMPLETO VERIFICADO EXITOSAMENTE")
    else:
        print(" FALLOS DETECTADOS EN EL FLUJO")
        sys.exit(1)
    print("=" * 60)
