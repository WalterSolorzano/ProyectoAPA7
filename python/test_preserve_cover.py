import copy
from pathlib import Path

import docx

src_path = Path(r"C:\Users\--X\.gemini\antigravity\scratch\wordapa7\10mo Trabajo Contabilidad.docx")
out_path = Path(r"C:\Users\--X\.gemini\antigravity\scratch\wordapa7\python\tests\output_test_preserve_cover.docx")

src_doc = docx.Document(str(src_path))
dst_doc = docx.Document()

# Copiar parrafos de portada (P0 a P17) directamente intactos
for p in src_doc.paragraphs[:18]:
    new_p = dst_doc.add_paragraph()
    new_p._element.getparent().replace(new_p._element, copy.deepcopy(p._element))

# Aplicar APA 7 a partir del cuerpo (P18 "Introduccion" en adelante)
p_intro = dst_doc.add_paragraph()
r = p_intro.add_run("Introducción")
r.bold = True

dst_doc.save(str(out_path))
print("Documento guardado con portada intacta preservada en:", out_path)
