import sys
from pathlib import Path
import docx

sys.path.insert(0, str(Path(__file__).resolve().parent))

from parsing.docx_parser import parse_docx_bytes
from generation.generator import generate_apa7_docx
from models import APARuleSet, ElementType

TARGET_FILE = Path(r"C:\Users\--X\Desktop\Estudio_Trabajo (1).docx")

def main():
    with open(TARGET_FILE, "rb") as f:
        file_bytes = f.read()

    doc_model = parse_docx_bytes(file_bytes, TARGET_FILE.name, "session_trace", Path("storage"))
    rules = APARuleSet()
    
    print("\n--- Diagnóstico de Imágenes en doc_model ---")
    for idx, e in enumerate(doc_model.elements[:10]):
        print(f"  elem #{idx:02d} | type={e.type} | is_cover={e.is_cover_section}")
    
    img_elems = [(idx, e.is_cover_section) for idx, e in enumerate(doc_model.elements) if e.type == "image" or e.type == ElementType.IMAGE]
    print("Índices de imágenes en doc_model:", img_elems)

    out_path = generate_apa7_docx(doc_model, "output_trace.docx", rules)
    doc = docx.Document(out_path)
    print("P#00 text:", repr(doc.paragraphs[0].text))
    print("P#01 text:", repr(doc.paragraphs[1].text))
    print("P#02 text:", repr(doc.paragraphs[2].text))

if __name__ == "__main__":
    main()
