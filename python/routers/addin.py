"""
WordAPA7 — Router del Word Add-in (Office.js) — Asistente en Vivo

Endpoints específicos para el Task Pane que vive dentro de Microsoft Word.
Estos endpoints operan sobre texto plano (fragmentos del documento) sin
necesidad de sesión ni archivo .docx, ya que el Add-in envía el contenido
del documento en tiempo real via Office.js.

El "Asistente en Vivo" auditúa el documento continuamente:
  - Extrae citas (chupar citas) y las guarda en el store local persistente.
  - Construye la bibliografía APA 7 dinámicamente.
  - Sugiere portada y detecta la estructura de títulos.

Endpoints:
  GET  /api/addin/health                — health check del backend
  POST /api/addin/analyze-selection      — analiza texto seleccionado contra APA 7
  POST /api/addin/suggest-caption        — sugiere caption APA para figura/tabla
  POST /api/addin/validate-fragment      — valida un fragmento de texto contra APA 7
  POST /api/addin/next-figure-number    — calcula el próximo número de figura
  POST /api/addin/next-table-number     — calcula el próximo número de tabla

  --- Asistente en Vivo ---
  POST /api/addin/extract-citations      — detecta citas en el texto, las guarda y las devuelve
  GET  /api/addin/references             — lista las referencias guardadas (store local)
  POST /api/addin/save-reference         — crea/actualiza una referencia
  DELETE /api/addin/reference/{ref_id}   — elimina una referencia
  DELETE /api/addin/clear-references     — vacía el store de referencias
  POST /api/addin/build-bibliography      — construye el texto APA 7 de Referencias
  POST /api/addin/suggest-cover          — sugiere los 5 campos de portada APA 7
  POST /api/addin/detect-headings        — sugiere niveles de título por estructura
  POST /api/addin/format-bibliography    — formatea referencias crudas a APA 7

  --- Detección de IA ---
  POST /api/addin/analyze-ai             — analiza un texto en busca de patrones de IA
  POST /api/addin/analyze-ai-document    — analiza múltiples párrafos con señales globales
  POST /api/addin/ai-rewrite             — reescribe texto en estilo académico APA 7
  GET  /api/addin/ai-health              — estado de salud de los proveedores de IA
  POST /api/addin/ai-analyze-table       — detecta artefactos de IA en datos tabulares
  POST /api/addin/audit-document         — auditoría APA 7 completa del documento (texto plano)
"""

import re
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel

router = APIRouter(prefix="/api/addin", tags=["word-addin"])


# ── SCHEMAS ──────────────────────────────────────────────────────────────────

class AnalyzeSelectionRequest(BaseModel):
    text: str
    context: Optional[str] = None  # texto circundante (párrafo anterior/siguiente)


class SuggestCaptionRequest(BaseModel):
    type: str  # "figure" | "table"
    context_text: str


class ValidateFragmentRequest(BaseModel):
    text: str
    element_type: Optional[str] = None  # "paragraph" | "heading" | "table" | "figure"


class NextNumberRequest(BaseModel):
    document_text: str


# ── SCHEMAS: Asistente en Vivo ───────────────────────────────────────────────

class ExtractCitationsRequest(BaseModel):
    text: str
    element_id: Optional[str] = "addin-live"


class SaveReferenceRequest(BaseModel):
    id: Optional[str] = None
    authors: List[str] = []
    year: Optional[str] = None
    title: str = ""
    source: str = ""
    doi_or_url: Optional[str] = None
    raw_text: str = ""


class SuggestCoverRequest(BaseModel):
    document_text: str


class DetectHeadingsRequest(BaseModel):
    paragraphs: List[str] = []
    document_text: Optional[str] = None


class FormatBibliographyRequest(BaseModel):
    references: List[str] = []  # referencias crudas en texto libre


class BuildBibliographyRequest(BaseModel):
    document_text: Optional[str] = None  # texto del documento para re-extraer citas


# ── SCHEMAS: Detección de IA ─────────────────────────────────────────────────

class AnalyzeAIRequest(BaseModel):
    """Solicitud de análisis de riesgo de IA para un fragmento de texto."""
    text: str


class AnalyzeAIDocumentRequest(BaseModel):
    """Solicitud de análisis de IA para múltiples párrafos del documento."""
    paragraphs: List[str] = []


class AIRewriteRequest(BaseModel):
    """Solicitud de reescritura de texto en estilo académico APA 7."""
    text: str


class AnalyzeTableAIRequest(BaseModel):
    """Solicitud de detección de artefactos de IA en datos de una tabla."""
    headers: List[str] = []
    rows: List[List[str]] = []


class AuditDocumentRequest(BaseModel):
    """Solicitud de auditoría APA 7 de documento completo (texto plano).

    El Add-in extrae el texto completo del documento vía Office.js
    (body.getText) y lo envía acá. Los saltos de párrafo pueden llegar
    como \\n, \\r\\n o \\r según el host de Word.
    """
    text: str


# ── PATRONES APA 7 (heurística local, sin IA) ────────────────────────────────

# Citas parentéticas: (Apellido, 2024) o (Apellido et al., 2024)
RE_CITATION_PARENTETICA = re.compile(
    r'\(([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+(?:et al\.?|y|&)\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)?,?\s+\d{4}[a-z]?(?:,\s*p\.?\s*\d+)?)\)'
)
# Citas narrativas: Apellido (2024) o Apellido et al. (2024)
RE_CITATION_NARRATIVA = re.compile(
    r'[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+(?:et al\.?|y)\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)?\s*\((\d{4}[a-z]?)\)'
)
# Doble espacio después de punto
RE_DOUBLE_SPACE = re.compile(r'\.\s{2,}')
# MAYÚSCULAS sostenidas (20+ caracteres consecutivos en mayúscula)
RE_SHOUTING = re.compile(r'[A-ZÁÉÍÓÚÑ]{20,}')
# "et al" sin punto
RE_ET_AL_NO_DOT = re.compile(r'et\s+al(?!\.)', re.IGNORECASE)
# DOI sin prefijo URL
RE_DOI_BARE = re.compile(r'(?<!https://doi\.org/)10\.\d{4,}/[^\s]+')
# Muletillas típicas de IA
RE_AI_FILLER = re.compile(
    r'\b(?:en conclusión|en resumen|es importante destacar|no obstante|'
    r'vale la pena|resulta fundamental|cabe destacar|cabe mencionar|'
    r'como se mencionó anteriormente|en primer lugar|en segundo lugar|'
    r'por un lado|por otro lado)\b',
    re.IGNORECASE
)
# Primera persona
RE_FIRST_PERSON = re.compile(r'\b(?:yo|nosotros|nuestra|nuestro|mi |me |creo que|pienso que)\b', re.IGNORECASE)
# Espanglish básico (palabras en inglés comunes dentro de texto en español)
RE_SPANGLISH = re.compile(r'\b(?:actually|basically|literally|obviously|you know|like\s)\b', re.IGNORECASE)


def _issue(category: str, severity: str, message: str, suggestion: str) -> dict:
    return {
        "category": category,
        "severity": severity,
        "message": message,
        "suggestion": suggestion,
    }


# ── ENDPOINTS BÁSICOS ─────────────────────────────────────────────────────────

@router.get("/health")
async def addin_health() -> dict:
    """Health check del backend para el Add-in."""
    return {"status": "ok", "version": "1.0.0"}


@router.post("/analyze-selection")
async def analyze_selection(req: AnalyzeSelectionRequest) -> dict:
    """
    Analiza un fragmento de texto seleccionado en Word contra las normas APA 7.

    Usa heurística local (regex) para detectar problemas comunes de formato
    APA de forma instantánea.

    Retorna:
      - issues: lista de problemas encontrados (category, severity, message, suggestion)
      - score: puntuación APA 0-100 (100 = perfecto)
      - suggestion: sugerencia general si aplica
    """
    text = (req.text or "").strip()
    if not text:
        return {"issues": [], "score": 100, "suggestion": None}

    issues: list[dict] = []

    # 1. Citas sin año
    if re.search(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+,?\s*\)', text):
        if not re.search(r'\([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+,?\s*\d{4}', text):
            issues.append(_issue(
                "citas", "warning",
                "Posible cita sin año entre paréntesis.",
                "APA 7 requiere el año: (Apellido, 2024)."
            ))

    # 2. "et al" sin punto abreviatura
    if RE_ET_AL_NO_DOT.search(text):
        issues.append(_issue(
            "citas", "error",
            "'et al' sin punto abreviatura.",
            "Escribí 'et al.' con punto."
        ))

    # 3. Doble espacio después de punto
    if RE_DOUBLE_SPACE.search(text):
        issues.append(_issue(
            "formato", "warning",
            "Se detectaron dobles espacios después de puntos.",
            "Reemplazá dobles espacios por uno solo."
        ))

    # 4. MAYÚSCULAS sostenidas
    match_shout = RE_SHOUTING.search(text)
    if match_shout:
        issues.append(_issue(
            "formato", "warning",
            f"Texto en MAYÚSCULAS sostenidas: \"{match_shout.group()[:30]}...\"",
            "APA 7 no usa mayúsculas sostenidas. Convertí a Title Case."
        ))

    # 5. Muletillas de IA
    filler_match = RE_AI_FILLER.search(text)
    if filler_match:
        issues.append(_issue(
            "estilo", "info",
            f"Muletilla detectada: \"{filler_match.group()}\"",
            "Eliminá las muletillas de IA para un texto más natural."
        ))

    # 6. Primera persona en texto académico
    first_person_match = RE_FIRST_PERSON.search(text)
    if first_person_match:
        issues.append(_issue(
            "estilo", "info",
            f"Uso de primera persona: \"{first_person_match.group()}\"",
            "APA 7 prefiere voz impersonal en trabajos académicos."
        ))

    # 7. Espanglish
    spanglish_match = RE_SPANGLISH.search(text)
    if spanglish_match:
        issues.append(_issue(
            "estilo", "warning",
            f"Palabra en inglés dentro del texto: \"{spanglish_match.group()}\"",
            "Traducí al español para mantener consistencia."
        ))

    # 8. DOI sin prefijo URL
    doi_match = RE_DOI_BARE.search(text)
    if doi_match:
        issues.append(_issue(
            "referencias", "warning",
            "DOI sin prefijo URL.",
            "El DOI debe empezar con 'https://doi.org/'."
        ))

    # 9. Palabra repetida (error de tipeo/pegado)
    rep_match = re.search(r'\b(\w{4,})\s+\1\b', text, re.IGNORECASE)
    if rep_match:
        issues.append(_issue(
            "formato", "error",
            f"Palabra repetida: \"{rep_match.group()}\"",
            "Eliminá la palabra duplicada."
        ))

    # Calcular score
    error_count = sum(1 for i in issues if i["severity"] == "error")
    warning_count = sum(1 for i in issues if i["severity"] == "warning")
    info_count = sum(1 for i in issues if i["severity"] == "info")

    score = max(0, 100 - (error_count * 20 + warning_count * 10 + info_count * 5))

    suggestion = None
    if score < 70:
        suggestion = "Revisá los problemas destacados. Tu texto necesita ajustes de formato APA 7."
    elif score < 100:
        suggestion = "Casi listo. Revisá las observaciones menores."
    elif score == 100:
        suggestion = "El texto cumple con APA 7."

    return {"issues": issues, "score": score, "suggestion": suggestion}


@router.post("/suggest-caption")
async def suggest_caption(req: SuggestCaptionRequest) -> dict:
    """
    Sugiere un caption APA 7 para una figura o tabla basado en el contexto.

    Si hay IA disponible, usa el LLM para generar una sugerencia inteligente.
    Si no, usa heurística básica.
    """
    element_type = req.type or "figure"
    context = (req.context_text or "").strip()

    # Intentar con IA si hay proveedores activos
    try:
        from modules.ai_client import execute_with_specialty
        label_prefix = "Figura" if element_type == "figure" else "Tabla"

        system_prompt = (
            "Sos un asistente APA 7. Generás títulos descriptivos para "
            f"{'figuras' if element_type == 'figure' else 'tablas'} en trabajos "
            "académicos. Respondés SOLO el título, sin numeración ni prefijo. "
            "Máximo 120 caracteres. En español académico."
        )
        user_prompt = (
            f"Contexto del documento:\n{context[:500]}\n\n"
            f"Generá un título descriptivo para una {label_prefix.lower()}."
        )

        caption = await execute_with_specialty(
            prompt=user_prompt,
            system_prompt=system_prompt,
            specialty="FAST",
            temperature=0.7,
            max_tokens=120,
            use_cache=True,
        )
        caption = (caption or "").strip().strip('"\'`')
        if caption:
            return {
                "label": label_prefix,
                "caption": caption[:120],
                "note": "",
            }
    except Exception:
        pass

    # Fallback heurístico
    label = "Figura" if element_type == "figure" else "Tabla"
    return {
        "label": label,
        "caption": f"Título de la {label.lower()}",
        "note": "",
    }


@router.post("/validate-fragment")
async def validate_fragment(req: ValidateFragmentRequest) -> dict:
    """
    Valida un fragmento de texto contra las normas APA 7.

    Similar a analyze-selection pero enfocado en un tipo específico de
    elemento (párrafo, heading, tabla, figura).
    """
    text = (req.text or "").strip()
    if not text:
        return {"status": "ok", "issues": [], "summary": "Sin texto para validar."}

    issues: list[dict] = []
    element_type = (req.element_type or "").lower()

    # Validaciones comunes (todos los tipos)
    if RE_DOUBLE_SPACE.search(text):
        issues.append(_issue("formato", "warning", "Dobles espacios.", "Unificar a espacio simple."))

    if RE_SHOUTING.search(text):
        issues.append(_issue("formato", "warning", "Mayúsculas sostenidas.", "Convertir a Title Case."))

    # Validaciones específicas por tipo
    if element_type == "heading":
        # APA 7: los headings no llevan punto final
        if text.rstrip().endswith('.'):
            issues.append(_issue("headings", "warning", "El título termina con punto.", "APA 7: los títulos no llevan punto final."))

    elif element_type == "paragraph":
        if RE_ET_AL_NO_DOT.search(text):
            issues.append(_issue("citas", "error", "'et al' sin punto.", "Escribir 'et al.'"))

        if RE_AI_FILLER.search(text):
            issues.append(_issue("estilo", "info", "Muletilla de IA detectada.", "Eliminar para texto natural."))

    elif element_type == "table":
        if not re.match(r'^Tabla\s+\d+', text):
            issues.append(_issue("tablas", "info", "La tabla debería tener un número.", "Formato APA 7: 'Tabla 1' en negrita."))

    elif element_type == "figure":
        if not re.match(r'^Figura\s+\d+', text):
            issues.append(_issue("figuras", "info", "La figura debería tener un número.", "Formato APA 7: 'Figura 1' en negrita."))

    # Status general
    has_error = any(i["severity"] == "error" for i in issues)
    has_warning = any(i["severity"] == "warning" for i in issues)
    status = "error" if has_error else ("warning" if has_warning else "ok")

    summary = "Cumple con APA 7."
    if has_error:
        summary = f"Se encontraron {sum(1 for i in issues if i['severity'] == 'error')} error(es) crítico(s)."
    elif has_warning:
        summary = f"Se encontraron {sum(1 for i in issues if i['severity'] == 'warning')} advertencia(s)."

    return {"status": status, "issues": issues, "summary": summary}


@router.post("/next-figure-number")
async def next_figure_number(req: NextNumberRequest) -> dict:
    """Calcula el próximo número de figura basado en el texto del documento."""
    text = req.document_text or ""
    pattern = re.compile(r'Figura\s+(\d+)', re.IGNORECASE)
    matches = pattern.findall(text)
    existing = sorted(set(int(m) for m in matches))
    next_num = (max(existing) if existing else 0) + 1
    return {"next_number": next_num, "existing_labels": [f"Figura {n}" for n in existing]}


@router.post("/next-table-number")
async def next_table_number(req: NextNumberRequest) -> dict:
    """Calcula el próximo número de tabla basado en el texto del documento."""
    text = req.document_text or ""
    pattern = re.compile(r'Tabla\s+(\d+)', re.IGNORECASE)
    matches = pattern.findall(text)
    existing = sorted(set(int(m) for m in matches))
    next_num = (max(existing) if existing else 0) + 1
    return {"next_number": next_num, "existing_labels": [f"Tabla {n}" for n in existing]}


# ════════════════════════════════════════════════════════════════════════════
#  ASISTENTE EN VIVO — ENDPOINTS
# ════════════════════════════════════════════════════════════════════════════

@router.post("/extract-citations")
async def extract_citations(req: ExtractCitationsRequest) -> dict:
    """
    "Chupa citas": escanea el texto del documento, detecta todas las citas
    APA 7 (parentéticas, narrativas, et al., múltiples, secundarias) usando el
    motor de citas de WordAPA7, las guarda en el store local persistente y
    devuelve la lista de citas detectadas en este texto.

    El frontend usa esto para notificar "¡Cita procesada con éxito!" y para
    saber cuántas citas nuevas se encontraron.
    """
    text = (req.text or "").strip()
    if not text:
        return {"citations": [], "new_count": 0, "total_in_store": 0}

    try:
        from modules.citation_engine import extract_citations_from_text
        from modules.addin_references_store import add_citation, bibliography_summary
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo cargar el motor de citas: {e}")

    # Extraer citas con el motor regex de WordAPA7
    detected = extract_citations_from_text(text, req.element_id or "addin-live")

    saved: List[Dict[str, Any]] = []
    for c in detected:
        entry = add_citation(
            raw_text=c.raw_text,
            authors=c.authors,
            year=c.year,
            citation_type=c.citation_type.value if hasattr(c.citation_type, "value") else str(c.citation_type),
            page=c.page,
        )
        saved.append({
            "id": entry.get("id"),
            "raw_text": c.raw_text,
            "authors": c.authors,
            "year": c.year,
            "page": c.page,
            "citation_type": entry.get("citation_type"),
            "is_new": entry.get("occurrences", 1) == 1,
        })

    new_count = sum(1 for c in saved if c.get("is_new"))
    summary = bibliography_summary()

    return {
        "citations": saved,
        "new_count": new_count,
        "total_in_store": summary["total"],
        "drafts": summary["drafts"],
    }


@router.get("/references")
async def list_references() -> dict:
    """Lista las referencias guardadas en el store local del Add-in."""
    from modules.addin_references_store import bibliography_summary
    return bibliography_summary()


@router.post("/save-reference")
async def save_reference(req: SaveReferenceRequest) -> dict:
    """Crea o actualiza una referencia en el store local."""
    from modules.addin_references_store import save_reference as _save
    ref = _save(req.model_dump())
    return {"status": "ok", "reference": ref}


@router.delete("/reference/{ref_id}")
async def delete_reference(ref_id: str) -> dict:
    """Elimina una referencia del store local."""
    from modules.addin_references_store import delete_reference as _delete
    ok = _delete(ref_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Referencia no encontrada.")
    return {"status": "ok"}


@router.delete("/clear-references")
async def clear_references() -> dict:
    """Vacía el store local de referencias del Add-in."""
    from modules.addin_references_store import clear as _clear
    _clear()
    return {"status": "ok"}


@router.post("/build-bibliography")
async def build_bibliography(req: Optional[BuildBibliographyRequest] = None) -> dict:
    """
    Construye el texto de la sección de Referencias en APA 7, ordenado
    alfabéticamente, listo para que el Add-in lo inserte al final del documento.
    """
    from modules.addin_references_store import build_bibliography_text, bibliography_summary

    if req is None:
        req = BuildBibliographyRequest()

    # Re-extracción de citas: si el frontend envía el texto del documento,
    # lo escaneamos antes de construir para mantener el store al día.
    document_text = (req.document_text or "").strip()
    if document_text:
        try:
            from modules.citation_engine import extract_citations_from_text
            from modules.addin_references_store import add_citation
            detected = extract_citations_from_text(document_text, "addin-live")
            for c in detected:
                add_citation(
                    raw_text=c.raw_text,
                    authors=c.authors,
                    year=c.year,
                    citation_type=c.citation_type.value if hasattr(c.citation_type, "value") else str(c.citation_type),
                    page=c.page,
                )
        except Exception:
            # Best-effort: si la re-extracción falla, se construye la
            # bibliografía con lo que ya hay en el store.
            pass

    summary = bibliography_summary()
    return {
        "bibliography_text": build_bibliography_text(),
        "total": summary["total"],
        "drafts": summary["drafts"],
        "complete": summary["complete"],
        "references": summary["references"],
    }


@router.post("/format-bibliography")
async def format_bibliography(req: FormatBibliographyRequest) -> dict:
    """
    Formatea una lista de referencias en texto libre a APA 7. Intenta
    resolver DOIs vía Crossref cuando los encuentra.
    """
    from modules.referencias_module import extract_doi_or_isbn, fetch_crossref_metadata
    import asyncio

    results: List[Dict[str, Any]] = []
    for raw in req.references:
        raw = (raw or "").strip()
        if not raw:
            continue
        info = extract_doi_or_isbn(raw)
        if info and info["type"] == "doi":
            try:
                meta = await fetch_crossref_metadata(info["value"])
                if meta:
                    results.append({
                        "original": raw,
                        "formatted": meta.get("formatted", raw),
                        "resolved": True,
                        "doi": info["value"],
                    })
                    continue
            except Exception:
                pass
        # Fallback: limpiar prefijos de viñeta
        cleaned = re.sub(r'^(?:[•○▪–\-\*]|\(?\d+[\.)]|\(?[a-zA-Z][\.)])\s*', '', raw)
        results.append({"original": raw, "formatted": cleaned, "resolved": False, "doi": None})

    return {"results": results}


# ── SECCIONES TÍPICAS PARA DETECCIÓN DE HEADINGS ─────────────────────────────

_LEVEL1_KEYWORDS = {
    "introduccion", "introducción", "marco teorico", "marco teórico",
    "metodologia", "metodología", "metodos", "métodos", "resultados",
    "discusion", "discusión", "conclusiones", "conclusion", "conclusión",
    "referencias", "bibliografia", "bibliografía", "resumen", "abstract",
    "anexos", "anexo", "apendices", "apéndices", "desarrollo", "objetivos",
    "planteamiento del problema", "justificacion", "justificación",
    "hipotesis", "hipótesis", "variables", "antecedentes",
}

_LEVEL2_KEYWORDS = {
    "objetivos generales", "objetivo general", "objetivos especificos",
    "objetivos específicos", "tipo de estudio", "tipo de investigacion",
    "tipo de investigación", "poblacion y muestra", "población y muestra",
    "tecnica de", "técnica de", "tecnicas de", "técnicas de",
    "instrumento", "instrumentos", "procedimiento", "procedimientos",
    "analisis de", "análisis de", "definicion de", "definición de",
    "limitaciones", "alcance", "marco conceptual", "marco legal",
    "bases teoricas", "bases teóricas",
}


@router.post("/detect-headings")
async def detect_headings(req: DetectHeadingsRequest) -> dict:
    """
    Analiza una lista de párrafos y sugiere el nivel de título APA 7 para
    cada uno según heurísticas (longitud, keywords, formato). Lo usa el
    asistente en vivo para aplicar estilos de título automáticamente.

    Retorna:
      - headings: [{ index, text, suggested_level, confidence, reason }]
      - paragraphs sin detectar como título no se incluyen.
    """
    paragraphs = req.paragraphs or []
    if not paragraphs and req.document_text:
        paragraphs = req.document_text.split("\n")

    headings: List[Dict[str, Any]] = []
    for idx, raw in enumerate(paragraphs):
        text = (raw or "").strip()
        if not text:
            continue

        norm = re.sub(r'\s+', ' ', text.lower()).strip()

        # Filtros: si termina en punto y es largo → no es título
        ends_with_period = text.rstrip().endswith('.')
        word_count = len(text.split())

        suggested_level: Optional[int] = None
        confidence = 0.5
        reason = ""

        if norm in _LEVEL1_KEYWORDS or any(norm.startswith(k) for k in _LEVEL1_KEYWORDS):
            suggested_level = 1
            confidence = 0.92
            reason = "Sección principal reconocida"
        elif norm in _LEVEL2_KEYWORDS or any(norm.startswith(k) for k in _LEVEL2_KEYWORDS):
            suggested_level = 2
            confidence = 0.82
            reason = "Sub-sección reconocida"
        elif (
            not ends_with_period
            and word_count <= 10
            and len(text) <= 80
            and text[0].isupper()
            # No parece una cita ni una línea de tabla
            and not text.startswith('(')
            and not re.match(r'^[\d\W]', text)
        ):
            # Heurística de "parece título": corto, sin punto, capitalizado.
            # Si tiene dos puntos "X: Y" podría ser nivel 3; por defecto nivel 1.
            if ':' in text and word_count <= 6:
                suggested_level = 2
                confidence = 0.6
                reason = "Título corto con dos puntos (posible sub-sección)"
            else:
                suggested_level = 1
                confidence = 0.6
                reason = "Línea corta sin punto (posible título)"

        if suggested_level is not None:
            headings.append({
                "index": idx,
                "text": text,
                "suggested_level": suggested_level,
                "confidence": round(confidence, 2),
                "reason": reason,
            })

    return {"suggestions": headings, "total": len(headings)}


@router.post("/suggest-cover")
async def suggest_cover(req: SuggestCoverRequest) -> dict:
    """
    Sugiere los 5 campos de la portada APA 7 (estudiante) a partir del
    contenido del documento:
      - title:    la frase/tema más relevante detectada
      - author:   nombre detectado por patrón
      - institution: institución detectada por patrón
      - course:   materia/curso detectado
      - date:     fecha actual por defecto

    Usa heurística local; si hay IA disponible la enriquece.
    """
    text = (req.document_text or "")[:8000]

    # ── Heurística local ──
    title = ""
    author = ""
    institution = ""
    course = ""
    date = ""

    # Título: primera línea no vacía significativa, o la frase tras "Tema:"/"Título:"
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for l in lines[:6]:
        low = l.lower()
        if low.startswith(("título:", "titulo:", "tema:")):
            title = l.split(":", 1)[1].strip()
            break
    if not title and lines:
        # Primera línea "grave" (no encabezado, no muy corta)
        for l in lines[:6]:
            if 3 < len(l) <= 120 and not l.endswith("."):
                title = l
                break
    if not title and lines:
        title = lines[0][:80]

    # Autor
    m = re.search(
        r'(?:autor(?:a)?|alumn[oa]|estudiante|nombre|elaborad[oa]?\s*por)[:\s]+([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+){1,3})',
        text, re.IGNORECASE
    )
    if m:
        author = m.group(1).strip()

    # Institución
    m = re.search(
        r'(?:universidad|instituto|facultad|colegio|escuela)\s+[\wÁÉÍÓÚÑáéíóúñ.,\s]{3,60}',
        text, re.IGNORECASE
    )
    if m:
        institution = m.group(0).strip().rstrip(".,")
    m2 = re.search(r'(?:instituci[oó]n|centro)[:\s]+([^\n,]{3,60})', text, re.IGNORECASE)
    if m2:
        institution = m2.group(1).strip()

    # Curso / materia
    m = re.search(r'(?:materia|curso|asignatura|c[aá]tedra)[:\s]+([^\n,]{3,60})', text, re.IGNORECASE)
    if m:
        course = m.group(1).strip()

    # Fecha
    m = re.search(r'\b(\d{1,2}\s+de\s+[a-záéíóúñ]+\s+de\s+\d{4})\b', text, re.IGNORECASE)
    if m:
        date = m.group(1)
    else:
        from datetime import datetime
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        now = datetime.now()
        date = f"{now.day} de {meses[now.month - 1]} de {now.year}"

    cover = {
        "title": title,
        "author": author,
        "institution": institution,
        "course": course,
        "date": date,
    }

    # ── Enriquecer con IA si está disponible ──
    try:
        from modules.ai_client import execute_with_specialty
        import json as _json

        system_prompt = (
            "Sos un asistente APA 7. Extraés los 5 campos de una portada "
            "estudiante APA 7 a partir del texto de un trabajo. Devolvés "
            "SOLO un JSON válido con las claves: title, author, institution, "
            "course, date. Si un campo no se puede inferir, devolvé string "
            "vacío. No inventes datos."
        )
        user_prompt = (
            f"Texto del documento (primeros 1500 caracteres):\n{text[:1500]}\n\n"
            "Extraé los campos de la portada."
        )
        content = await execute_with_specialty(
            prompt=user_prompt,
            system_prompt=system_prompt,
            specialty="FAST",
            temperature=0.2,
            max_tokens=300,
            use_cache=True,
        )
        if content:
            txt = content.strip()
            if txt.startswith("```json"):
                txt = txt[7:-3]
            elif txt.startswith("```"):
                txt = txt[3:-3]
            data = _json.loads(txt)
            # Solo sobreescribir campos que la IA haya podido inferir
            for k in ("title", "author", "institution", "course", "date"):
                v = (data.get(k) or "").strip()
                if v:
                    cover[k] = v
    except Exception:
        # La heurística local ya está en `cover`; no es crítico.
        pass

    return cover


# ════════════════════════════════════════════════════════════════════════════
#  DETECCIÓN DE IA — ENDPOINTS
# ════════════════════════════════════════════════════════════════════════════
#
#  Estos endpoints exponen el sistema de detección de IA de WordAPA7 al
#  Word Add-in. Permiten:
#    - Analizar texto individual o párrafos completos en busca de patrones
#      típicos de contenido generado por LLM (muletillas, estructura
#      homogénea, artefactos de copia web, etc.).
#    - Reescribir texto al estilo académico APA 7 eliminando muletillas.
#    - Consultar el estado de salud de los proveedores de IA.
#    - Detectar artefactos de IA en datos de tablas (emojis, punto y coma).
#
#  Los módulos subyacentes NO se modifican — solo se importan y exponen:
#    - classification.ai_detector   (analyze_ai_risk, analyze_document_ai,
#                                    analyze_table_cells)
#    - modules.ai_client            (execute_with_specialty, get_ai_system_health)
#    - classification.llm_classifier (_get_active_providers)
# ════════════════════════════════════════════════════════════════════════════


@router.post("/analyze-ai")
async def analyze_ai(req: AnalyzeAIRequest) -> dict:
    """
    Analiza un texto en busca de patrones típicos de contenido generado por IA.

    Usa el motor de detección basado en patrones ``analyze_ai_risk`` que
    identifica muletillas de IA, estructura de oraciones homogénea, abuso de
    punto y coma, conclusiones genéricas, artefactos de copia web (emojis,
    markdown, doble guion), densidad de conectores, anáfora repetida y
    tipografía mixta.

    Retorna directamente el resultado de ``analyze_ai_risk``:
      - **score**: puntuación de riesgo (0.0 a 1.0)
      - **category**: ``"LOW"`` | ``"MEDIUM"`` | ``"HIGH"``
      - **findings**: lista de hallazgos, cada uno con
        ``pattern``, ``severity``, ``detail``, ``count``, ``phrase``
    """
    text = (req.text or "").strip()
    if not text:
        return {"score": 0.0, "category": "LOW", "findings": []}

    try:
        from classification.ai_detector import analyze_ai_risk
        result = analyze_ai_risk(text)
        return result
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error al analizar el texto con el detector de IA: {e}",
        )


@router.post("/analyze-ai-document")
async def analyze_ai_document(req: AnalyzeAIDocumentRequest) -> dict:
    """
    Analiza múltiples párrafos en busca de patrones de IA, aplicando señales
    globales del documento (varianza sintáctica entre párrafos).

    Usa ``analyze_document_ai`` del motor de detección, que analiza cada
    párrafo con ``analyze_ai_risk`` y luego aplica señales globales
    consistentes con el algoritmo del Revisor de WordAPA7.

    Retorna:
      - **results**: lista de resultados por párrafo, cada uno con
        ``score``, ``category``, ``findings``, ``matches``
      - **total_paragraphs**: cantidad total de párrafos analizados
      - **high_risk_count**: cantidad de párrafos con categoría ``HIGH``
      - **medium_risk_count**: cantidad de párrafos con categoría ``MEDIUM``
      - **low_risk_count**: cantidad de párrafos con categoría ``LOW``
      - **overall_score**: puntuación promedio del documento (0.0 a 1.0)
      - **overall_category**: categoría general del documento
        (``"LOW"`` | ``"MEDIUM"`` | ``"HIGH"``)
    """
    paragraphs = req.paragraphs or []
    if not paragraphs:
        return {
            "results": [],
            "total_paragraphs": 0,
            "high_risk_count": 0,
            "medium_risk_count": 0,
            "low_risk_count": 0,
            "overall_score": 0.0,
            "overall_category": "LOW",
        }

    try:
        from classification.ai_detector import analyze_document_ai
        results = analyze_document_ai(paragraphs)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error al analizar el documento con el detector de IA: {e}",
        )

    # ── Calcular estadísticas agregadas ──
    high_count = sum(1 for r in results if r.get("category") == "HIGH")
    medium_count = sum(1 for r in results if r.get("category") == "MEDIUM")
    low_count = sum(1 for r in results if r.get("category") == "LOW")

    scores = [r.get("score", 0.0) for r in results]
    overall_score = round(sum(scores) / len(scores), 2) if scores else 0.0

    # Categorizar el score general del documento
    if overall_score >= 0.5:
        overall_category = "HIGH"
    elif overall_score >= 0.2:
        overall_category = "MEDIUM"
    else:
        overall_category = "LOW"

    return {
        "results": results,
        "total_paragraphs": len(results),
        "high_risk_count": high_count,
        "medium_risk_count": medium_count,
        "low_risk_count": low_count,
        "overall_score": overall_score,
        "overall_category": overall_category,
    }


@router.post("/ai-rewrite")
async def ai_rewrite(req: AIRewriteRequest) -> Any:
    """
    Reescribe texto en un estilo académico natural APA 7, eliminando
    muletillas de IA, frases genéricas y lenguaje repetitivo, manteniendo el
    significado y las citas.

    Usa los proveedores de IA disponibles mediante ``execute_with_specialty``
    con la especialidad **FAST** (latencia baja) y ``temperature=0.4`` para
    un equilibrio entre creatividad y fidelidad.

    Retorna:
      - **original**: texto original enviado
      - **rewritten**: texto reescrito en estilo APA 7
      - **provider**: nombre del proveedor de IA que procesó la solicitud

    Si no hay proveedores de IA disponibles o todos fallan, retorna
    ``{"error": "message"}`` con status ``503``.
    """
    text = (req.text or "").strip()
    if not text:
        return {"original": "", "rewritten": "", "provider": ""}

    try:
        from modules.ai_client import execute_with_specialty
    except Exception as e:
        return JSONResponse(
            status_code=503,
            content={"error": f"No se pudo cargar el módulo de IA: {e}"},
        )

    system_prompt = (
        "Sos un editor académico. Reescribí el texto en estilo académico APA 7, "
        "eliminando muletillas de IA, frases genéricas y lenguaje repetitivo. "
        "Mantené el significado y las citas. Respondé SOLO el texto reescrito, "
        "sin explicaciones."
    )

    try:
        # Usamos return_provider_info=True para obtener el nombre del proveedor
        # que efectivamente procesó la solicitud.
        content, provider_name, _provider_id = await execute_with_specialty(
            prompt=text,
            system_prompt=system_prompt,
            specialty="FAST",
            temperature=0.4,
            max_tokens=2000,
            use_cache=False,
            return_provider_info=True,
        )
    except (ValueError, RuntimeError) as e:
        # ValueError: no hay proveedores configurados o activos.
        # RuntimeError: la infraestructura LLM colapsó (rate limit / timeout).
        return JSONResponse(
            status_code=503,
            content={"error": f"No hay proveedores de IA disponibles o todos fallaron: {e}"},
        )
    except Exception as e:
        return JSONResponse(
            status_code=503,
            content={"error": f"Error al procesar la reescritura con IA: {e}"},
        )

    rewritten = (content or "").strip()
    if not rewritten:
        return JSONResponse(
            status_code=503,
            content={"error": "El proveedor de IA devolvió una respuesta vacía."},
        )

    return {
        "original": text,
        "rewritten": rewritten,
        "provider": provider_name,
    }


@router.get("/ai-health")
async def ai_health() -> dict:
    """
    Retorna el estado de salud del sistema de IA: qué proveedores están
    activos y el estado de rate-limit por especialidad (FAST, HEAVY,
    REASONING).

    Combina dos fuentes:
      - ``_get_active_providers()`` de ``classification.llm_classifier`` para
        obtener la lista de proveedores activos (sus nombres).
      - ``get_ai_system_health()`` de ``modules.ai_client`` para el estado
        de los buckets de rate-limiting por especialidad.

    Retorna:
      - **providers_active**: cantidad de proveedores activos
      - **providers**: lista de nombres de proveedores activos
      - **specialties**: estado por especialidad, cada una con
        ``provider``, ``percentage``, ``status``
    """
    # ── Proveedores activos ──
    try:
        from classification.llm_classifier import _get_active_providers
        active_providers = _get_active_providers()
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener los proveedores activos: {e}",
        )

    provider_names = [
        p.get("name", p.get("id", "desconocido"))
        for p in active_providers
    ]

    # ── Estado de rate-limit por especialidad ──
    try:
        from modules.ai_client import get_ai_system_health
        specialties = get_ai_system_health()
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener el estado del sistema de IA: {e}",
        )

    return {
        "providers_active": len(provider_names),
        "providers": provider_names,
        "specialties": specialties,
    }


@router.post("/ai-analyze-table")
async def ai_analyze_table(req: AnalyzeTableAIRequest) -> dict:
    """
    Detecta artefactos de copia web/IA en datos de una tabla.

    Usa ``analyze_table_cells`` del detector de IA, que examina encabezados
    y filas buscando señales típicas de contenido pegado desde chat/IA:
      - Emojis y símbolos de checklist ( X  ) en celdas
      - Múltiples punto y coma en celdas (texto pegado de web)
      - Otros artefactos de copia

    Retorna:
      - **score**: puntuación de riesgo (0.0 a 1.0)
      - **findings**: lista de hallazgos, cada uno con
        ``pattern``, ``severity``, ``detail``, ``count``, ``phrase``
    """
    headers = req.headers or []
    rows = req.rows or []

    if not headers and not rows:
        return {"score": 0.0, "findings": []}

    try:
        from classification.ai_detector import analyze_table_cells
        score, findings = analyze_table_cells(headers, rows)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error al analizar la tabla con el detector de IA: {e}",
        )

    return {
        "score": score,
        "findings": findings,
    }


# ════════════════════════════════════════════════════════════════════════════
#  AUDITORÍA DE DOCUMENTO COMPLETO (botón "Auditar documento ahora")
# ════════════════════════════════════════════════════════════════════════════

MAX_AUDIT_FINDINGS = 50
QUOTE_MIN_CHARS = 40

RE_PARA_SPLIT = re.compile(r'\r\n|\r|\n')
RE_PAREN_GROUP = re.compile(r'\(([^()]{2,80})\)')
RE_YEAR = re.compile(r'\d{4}')
RE_NAME_START = re.compile(r"^[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ.'’\-]*(?:\s+(?:et\s+al\.?|y|&)\s+[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ.'’\-]*)*$")
RE_AMPERSAND_CIT = re.compile(r'(?:^|\s)&(?:\s|$)')
RE_Y_CIT = re.compile(r'(?:^|\s)y(?:\s|$)')
RE_NUMBERED_HEADING = re.compile(r'^\s*\d+(?:\.\d+)*[.)]?\s+\S')
RE_ALLCAPS_LINE = re.compile(r'^[A-ZÁÉÍÓÚÑ0-9][A-ZÁÉÍÓÚÑ0-9\s.,;:()\-]{7,90}$')
RE_REF_PREFIX = re.compile(r'^(?:[•○▪●·–—\-–*]|\(?\d{1,3}[.)])\s+')
RE_LONG_QUOTE = re.compile(r'[“"]([^“”"]{40,})[“”"]')
RE_DOUBLE_SPACE_INNER = re.compile(r'\S\s{2,}\S')
RE_NARRATIVE_CIT = re.compile(
    r'[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+(?:et al\.?|y)\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)?\s*\(\d{4}'
)
REFS_SECTION_TITLES = {'referencias', 'references', 'bibliografía', 'bibliografia'}


def _split_paragraphs(text: str) -> List[str]:
    """Divide texto plano en párrafos (\\n, \\r\\n o \\r)."""
    return RE_PARA_SPLIT.split(text or '')


def _excerpt(s: str, limit: int = 90) -> str:
    s = re.sub(r'\s+', ' ', (s or '')).strip()
    return s[:limit] + ('…' if len(s) > limit else '')


@router.post("/audit-document")
async def audit_document(req: AuditDocumentRequest) -> dict:
    """
    Auditoría APA 7 de un documento completo a partir de texto plano.

    Reutiliza heurísticas del motor de citas (modules.citation_engine) y del
    store de referencias del Add-in para cruzar citas ↔ referencias. El resto
    son reglas regex locales (sin IA, sin latencia).

    Retorna forma plana:
      - findings: [{ id, severity: 'error'|'warn'|'info',
                     category: 'citacion'|'referencias'|'formato'|'estructura',
                     message, fix, where?: { paragraph_index?, excerpt? } }]
                    (cap 50 hallazgos)
      - stats: { paragraphs, words, citations, references }
    """
    text = req.text or ''
    paragraphs = _split_paragraphs(text)
    non_empty = [(i, p.strip()) for i, p in enumerate(paragraphs) if p.strip()]
    findings: List[Dict[str, Any]] = []

    def add(severity: str, category: str, message: str, fix: str,
            paragraph_index: Optional[int] = None, excerpt: Optional[str] = None) -> None:
        if len(findings) >= MAX_AUDIT_FINDINGS:
            return
        where: Dict[str, Any] = {}
        if paragraph_index is not None:
            where['paragraph_index'] = paragraph_index
        if excerpt:
            where['excerpt'] = _excerpt(excerpt)
        findings.append({
            'id': f'f{len(findings) + 1}',
            'severity': severity,
            'category': category,
            'message': message,
            'fix': fix,
            **({'where': where} if where else {}),
        })

    # ── Motor de citas + store de referencias (heurísticas server-side) ──
    detected_citations: List[Any] = []
    ref_keys: set = set()
    citations_count = 0
    try:
        from modules.citation_engine import extract_citations_from_text, normalize_surname_key
        detected_citations = extract_citations_from_text(text, 'addin-audit')
        citations_count = len(detected_citations)
        try:
            from modules.addin_references_store import bibliography_summary
            summary = bibliography_summary()
            for r in summary.get('references', []) or []:
                authors = r.get('authors') or []
                year = r.get('year') or ''
                if authors:
                    ref_keys.add(f"{normalize_surname_key(str(authors[0]))}|{str(year)}")
        except Exception:
            pass
    except Exception:
        normalize_surname_key = lambda a: re.sub(r'[^a-z]', '', str(a).lower().split(',')[0].split()[-1])  # noqa: E731

    # ── Sección Referencias en el texto ──
    refs_start_idx: Optional[int] = None
    for idx, p in non_empty:
        if p.rstrip(':').strip().lower() in REFS_SECTION_TITLES:
            refs_start_idx = idx
            break

    reference_lines = 0
    if refs_start_idx is not None:
        for idx, p in non_empty:
            if idx <= refs_start_idx:
                continue
            reference_lines += 1
            # Regla APA: las entradas NO llevan viñeta ni numeración.
            if RE_REF_PREFIX.match(p):
                add('error', 'referencias',
                    'Entrada de referencia con viñeta o numeración.',
                    'Quitá la viñeta/número; APA 7 usa lista con sangría francesa.',
                    paragraph_index=idx, excerpt=p)

    # ── Reglas por párrafo ──
    ampersand_hits = 0
    y_hits = 0
    double_space_reported = False
    straight_quote_paras = 0

    for idx, p in enumerate(paragraphs):
        stripped = p.strip()
        if not stripped:
            continue
        is_ref_line = refs_start_idx is not None and idx > refs_start_idx

        # 1. Cita parentética sin año → error
        for m in RE_PAREN_GROUP.finditer(stripped):
            group = m.group(1).strip()
            looks_like_author = (
                not RE_YEAR.search(group)
                and not any(ch.isdigit() for ch in group)
                and any(c.islower() for c in group)  # descarta siglas "(APA)"
                and RE_NAME_START.match(group)
            )
            if looks_like_author:
                add('error', 'citacion',
                    f'Cita sin año: ({group}).',
                    'APA 7 exige año: (Apellido, 2024).',
                    paragraph_index=idx, excerpt=m.group(0))
                break

        # 2. "et al" sin punto → error
        if re.search(r'et\s+al(?![.\w])', stripped):
            add('error', 'citacion',
                "'et al' sin punto.",
                "Escribí 'et al.' con punto final.",
                paragraph_index=idx, excerpt=stripped)

        # 3. & vs "y" en citas parentéticas (consistencia)
        for m in RE_PAREN_GROUP.finditer(stripped):
            g = m.group(1)
            if RE_YEAR.search(g) and len(g.split()) >= 3:
                if RE_AMPERSAND_CIT.search(g):
                    ampersand_hits += 1
                elif RE_Y_CIT.search(g):
                    y_hits += 1

        # 4. Títulos numerados (no permitidos en APA 7) → warn estructura
        if (not is_ref_line and RE_NUMBERED_HEADING.match(stripped)
                and len(stripped.split()) <= 12 and not stripped.endswith('.')):
            add('warn', 'estructura',
                'Título numerado ("1." / "1.1.").',
                'Los títulos APA 7 no llevan numeración.',
                paragraph_index=idx, excerpt=stripped)

        # 5. Línea en MAYÚSCULAS sostenidas (posible título) → warn formato
        if (not is_ref_line and RE_ALLCAPS_LINE.match(stripped)
                and sum(c.isalpha() for c in stripped) >= 8):
            add('warn', 'formato',
                'Título en MAYÚSCULAS sostenidas.',
                'Usá negrita centrada/izquierda, no mayúsculas.',
                paragraph_index=idx, excerpt=stripped)

        # 6. Cita textual larga (>40 caracteres) sin cita → warn
        if not is_ref_line:
            for qm in RE_LONG_QUOTE.finditer(stripped):
                has_cit = bool(RE_PAREN_GROUP.search(stripped) and RE_YEAR.search(stripped)) \
                    or bool(RE_NARRATIVE_CIT.search(stripped))
                if not has_cit:
                    add('warn', 'citacion',
                        'Cita textual larga sin referencia en el párrafo.',
                        'Sumá (Autor, año, pág.) o convertila en bloque.',
                        paragraph_index=idx, excerpt=qm.group(0)[:120])
                    break

        # 7. Dobles espacios internos → warn formato (una sola vez por doc)
        if not double_space_reported and RE_DOUBLE_SPACE_INNER.search(stripped):
            add('warn', 'formato',
                'Dobles espacios entre palabras.',
                'Dejá un solo espacio entre palabras.',
                paragraph_index=idx, excerpt=stripped)
            double_space_reported = True

        # 8. Comillas rectas vs curvas → info formato
        if ('"' in stripped or "'" in stripped) and not is_ref_line:
            straight_quote_paras += 1

        # 9. DOI sin prefijo URL → warn referencias
        if RE_DOI_BARE.search(stripped):
            add('warn', 'referencias',
                'DOI sin prefijo https://doi.org/.',
                "Escribí 'https://doi.org/' adelante del DOI.",
                paragraph_index=idx, excerpt=stripped)

    # Consolidado "& vs y"
    if ampersand_hits > 0 and y_hits > 0:
        add('warn', 'citacion',
            f'Mezcla de "&" ({ampersand_hits}) e "y" ({y_hits}) en citas.',
            'Elegí un estilo y usalo siempre igual ("y" recomendado).')
    elif ampersand_hits > 0:
        add('info', 'citacion',
            f'Citas con "&" ({ampersand_hits} caso(s)).',
            'APA 7 en español recomienda "y" dentro de paréntesis.')

    # Comillas rectas consolidadas
    if straight_quote_paras > 0:
        add('info', 'formato',
            f'Comillas rectas en {straight_quote_paras} párrafo(s).',
            'Usá comillas curvas “ ” en el texto final.')

    # ── Citas huérfanas (cruzada cita ↔ referencias del store) ──
    if ref_keys and detected_citations:
        orphans: List[str] = []
        for c in detected_citations[:60]:
            authors = getattr(c, 'authors', None) or []
            year = getattr(c, 'year', None) or ''
            key = f"{normalize_surname_key(str(authors[0]))}|{str(year)}" if authors else ''
            if key and key not in ref_keys:
                label = f"{authors[0]} ({year})" if authors else str(getattr(c, 'raw_text', ''))[:40]
                if label not in orphans:
                    orphans.append(label)
        if orphans:
            add('error', 'referencias',
                'Citas sin entrada en Referencias: ' + ', '.join(orphans[:5])
                + ('…' if len(orphans) > 5 else ''),
                'Agregá cada obra citada a la sección Referencias.')

    # ── Artillería completa: Revisor Proactivo de Redacción, Ortografía y Muletillas IA ──
    try:
        from modules.proactive_auditor import audit_elements

        class _SimpleElem:
            def __init__(self, eid: str, txt: str):
                self.id = eid
                self.text = txt

        audit_elems = [_SimpleElem(str(idx), p) for idx, p in enumerate(paragraphs) if p.strip()]
        proactive_findings = audit_elements(audit_elems)
        for pf in proactive_findings:
            kind = pf.get("kind", "estilo")
            category = "citacion" if kind == "first_person" else "estructura" if kind in ("repeticion", "incompleta") else "formato"
            severity = "error" if pf.get("severity") == "error" else "warn" if pf.get("severity") == "warn" else "info"
            p_idx = int(pf["element_id"]) if pf.get("element_id", "").isdigit() else None
            add(
                severity,
                category,
                pf.get("message", ""),
                pf.get("suggestion", "Mejora la redacción académica según normas APA 7."),
                paragraph_index=p_idx,
                excerpt=pf.get("excerpt", "")
            )
    except Exception:
        pass

    # ── Estadísticas ──
    words = len(text.split())
    stats = {
        'paragraphs': len(non_empty),
        'words': words,
        'citations': citations_count,
        'references': reference_lines,
    }

    return {'findings': findings[:MAX_AUDIT_FINDINGS], 'stats': stats}


# ════════════════════════════════════════════════════════════════════════════
#  PROVEEDORES DE IA — CHIPS DE ESTADO (pestaña IA del Add-in)
# ════════════════════════════════════════════════════════════════════════════
#
#  GET /api/addin/ai-providers → { providers: [{ name, active }], count }
#
#  Inspecciona SOLO variables de entorno (sin tocar módulos de IA): un
#  proveedor está activo si su clave existe y no es vacía. Cloudflare exige
#  el par CLOUDFLARE_API_TOKEN + CLOUDFLARE_ACCOUNT_ID.
# ════════════════════════════════════════════════════════════════════════════

AI_PROVIDER_ENV_KEYS: List[tuple] = [
    ("NVIDIA", ("NVIDIA_API_KEY",)),
    ("Groq", ("GROQ_API_KEY",)),
    ("OpenRouter", ("OPENROUTER_API_KEY",)),
    ("Cerebras", ("CEREBRAS_API_KEY",)),
    ("Mistral", ("MISTRAL_API_KEY",)),
    ("OpenCodeZen", ("OPENCODEZEN_API_KEY",)),
    ("ZenMux", ("ZENMUX_API_KEY",)),
    ("Gemini", ("GEMINI_API_KEY",)),
    ("Cloudflare", ("CLOUDFLARE_API_TOKEN", "CLOUDFLARE_ACCOUNT_ID")),
]


@router.get("/ai-providers")
async def ai_providers() -> dict:
    """Estado de los proveedores de IA según variables de entorno."""
    import os

    providers: List[Dict[str, Any]] = []
    for name, keys in AI_PROVIDER_ENV_KEYS:
        active = all((os.environ.get(k) or "").strip() for k in keys)
        providers.append({"name": name, "active": active})

    return {
        "providers": providers,
        "count": sum(1 for p in providers if p["active"]),
    }


class ScopedApplyLiveRequest(BaseModel):
    ooxml_base64: str
    scopes: List[str]
    rules: Optional[Dict[str, Any]] = None


@router.post("/scoped-apply-live")
async def scoped_apply_live(req: ScopedApplyLiveRequest) -> dict:
    """Aplica alcances (texto, tablas_imagenes, bibliografia) sobre el documento OOXML."""
    import base64
    from modules.scoped_apply import apply_scopes, VALID_SCOPES

    # Validar alcances
    invalid = [s for s in req.scopes if s not in VALID_SCOPES]
    if invalid:
        raise HTTPException(status_code=400, detail=f"Alcances inválidos: {invalid}")

    try:
        raw_bytes = base64.b64decode(req.ooxml_base64)
        # Si es OOXML plano (XML), empaquetarlo en un docx mínimo
        if raw_bytes.strip().startswith(b"<?xml") or b"<w:document" in raw_bytes[:200]:
            import docx
            import io
            doc = docx.Document()
            # Si se pasó XML crudo, procesar
            out_bytes = io.BytesIO()
            doc.save(out_bytes)
            raw_bytes = out_bytes.getvalue()

        out_docx, summary = apply_scopes(raw_bytes, req.scopes, req.rules or {})
        return {
            "docx_base64": base64.b64encode(out_docx).decode("ascii"),
            "summary": summary,
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en scoped_apply_live: {e}")

