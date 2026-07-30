import sys
import shutil
from abc import ABC, abstractmethod
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, List
import logging
import time

logger = logging.getLogger(__name__)

@dataclass
class PageLayoutResult:
    """Resultado de paginar un .docx."""
    paragraph_pages: List[int]      # Para cada párrafo del body (en orden), su número de página (1-indexed)
    total_pages: int                # Total de páginas del documento
    provider_used: str              # "com" | "libreoffice" | "heuristic"
    confidence: float               # 1.0 para COM, 0.95 LibreOffice, 0.5 heurística
    notes: List[str]                # Warnings (ej: "Word colgó, se mató proceso")

class PageLayoutProvider(ABC):
    @abstractmethod
    def paginate(self, docx_path: Path, timeout_seconds: int = 30) -> PageLayoutResult:
        """Abre el .docx, pagina, retorna el mapeo párrafo → página."""
        pass

    @abstractmethod
    def is_available(self) -> bool:
        """True si el backend está instalado en esta máquina."""
        pass

class COMPageLayoutProvider(PageLayoutProvider):
    """
    Usa Word.Application via pywin32 para obtener el número de página exacto
    de cada párrafo. Es la fuente de verdad más precisa disponible.
    """

    def is_available(self) -> bool:
        if sys.platform != "win32":
            return False
        try:
            import win32com.client  # noqa
            return True
        except ImportError:
            return False

    def paginate(self, docx_path: Path, timeout_seconds: int = 30) -> PageLayoutResult:
        try:
            # Timeout externo: si Word se cuelga, matar WINWORD.EXE
            result = self._paginate_with_timeout(
                docx_path, timeout_seconds
            )
            return result
        finally:
            # Limpieza agresiva de procesos zombies por si acaso
            self._kill_orphan_winword_processes()

    def _paginate_with_timeout(self, docx_path, timeout):
        # Usar concurrent.futures con timeout
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(self._do_paginate, docx_path)
            try:
                return future.result(timeout=timeout)
            except concurrent.futures.TimeoutError:
                raise TimeoutError(
                    f"Word no respondió en {timeout}s — matando proceso"
                )

    def _do_paginate(self, docx_path) -> PageLayoutResult:
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            from services.word_com_service import get_word_com_service
            word_service = get_word_com_service()
            word = word_service.word
            if not word:
                raise Exception("Word COM Service failed to provide a valid Word instance")
                
            word.Visible = False
            word.DisplayAlerts = 0

            doc = word.Documents.Open(
                str(docx_path.resolve()),
                ConfirmConversions=False,
                ReadOnly=True,           # NUNCA abrir para escritura
                AddToRecentFiles=False,  # No contaminar historial del usuario
            )
            t0 = time.time()
            doc.Repaginate()  # CRÍTICO: fuerza cálculo de paginación
            t_repag = time.time() - t0

            paragraph_pages: List[int] = []
            for para in doc.Paragraphs:
                # wdActiveEndPageNumber = 3
                page_num = para.Range.Information(3)
                paragraph_pages.append(int(page_num))

            total_pages = doc.ComputeStatistics(2)  # wdStatisticPages = 2
            t_total = time.time() - t0
            
            logger.info(f"[COM Layout] docx={docx_path.name}, total_pages={total_pages}, paragraphs={len(paragraph_pages)}, time_repag={t_repag:.2f}s, time_total={t_total:.2f}s")
            
            return PageLayoutResult(
                paragraph_pages=paragraph_pages,
                total_pages=total_pages,
                provider_used="com",
                confidence=1.0,
                notes=[],
            )
        finally:
            try:
                if doc:
                    doc.Close(SaveChanges=False)
            except Exception:
                pass
            
            doc = None
            word = None
            
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _kill_orphan_winword_processes(self):
        """Mata procesos WINWORD.EXE zombies de sesiones anteriores.
        Desactivado temporalmente para no interferir con otras instancias
        o ejecuciones concurrentes (como pytest o el usuario).
        """
        pass

class LibreOfficePageLayoutProvider(PageLayoutProvider):
    """
    Convierte .docx → PDF con LibreOffice headless, luego usa PyMuPDF (fitz)
    para extraer el mapeo párrafo → página.
    """

    def is_available(self) -> bool:
        return bool(shutil.which("soffice") or shutil.which("libreoffice"))

    def paginate(self, docx_path: Path, timeout_seconds: int = 60) -> PageLayoutResult:
        import subprocess
        import tempfile

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_dir_path = Path(tmp_dir)
            pdf_path = tmp_dir_path / (docx_path.stem + ".pdf")

            # LibreOffice headless → PDF
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(tmp_dir_path), str(docx_path)],
                capture_output=True, text=True, timeout=timeout_seconds,
            )
            if result.returncode != 0 or not pdf_path.exists():
                raise RuntimeError(f"LibreOffice falló: {result.stderr[:300]}")

            # PyMuPDF: para cada página, extraer texto y comparar con párrafos del .docx
            return self._map_paragraphs_to_pages(docx_path, pdf_path)

    def _map_paragraphs_to_pages(self, docx_path, pdf_path):
        import fitz  # PyMuPDF
        import docx

        # 1. Obtener lista de textos de párrafos del .docx (en orden)
        doc = docx.Document(str(docx_path))
        para_texts = [p.text.strip() for p in doc.paragraphs]

        # 2. Abrir PDF y extraer texto por página
        pdf = fitz.open(str(pdf_path))
        page_texts: List[str] = []
        for page in pdf:
            page_texts.append(page.get_text("text"))

        # 3. Matching: para cada párrafo del .docx, buscar en qué página aparece
        paragraph_pages: List[int] = []
        current_page = 1
        consumed_offset = 0  # offset dentro del texto acumulado de páginas

        # Estrategia: concatenar todos los textos de página con separador,
        # luego buscar el offset de cada párrafo y mapear a página.
        full_text = ""
        page_boundaries: List[int] = []  # índice en full_text donde empieza cada página
        for pt in page_texts:
            page_boundaries.append(len(full_text))
            full_text += pt + "\\n"

        import re
        for para_text in para_texts:
            if not para_text:
                paragraph_pages.append(current_page)
                continue
            # Normalizar whitespace para matching robusto
            pattern = re.escape(re.sub(r'\\s+', ' ', para_text))[:80]
            match = re.search(pattern, re.sub(r'\\s+', ' ', full_text[consumed_offset:]))
            if match:
                abs_offset = consumed_offset + match.start()
                # Buscar a qué página pertenece abs_offset
                for p_idx, boundary in enumerate(page_boundaries):
                    if abs_offset < boundary:
                        current_page = p_idx  # 1-indexed
                        break
                    elif p_idx == len(page_boundaries) - 1:
                        current_page = p_idx + 1
                else:
                    current_page = len(page_boundaries)
                # Actualizar consumed_offset para siguiente búsqueda
                consumed_offset = abs_offset + len(pattern)
            paragraph_pages.append(current_page)

        return PageLayoutResult(
            paragraph_pages=paragraph_pages,
            total_pages=len(page_texts),
            provider_used="libreoffice",
            confidence=0.95,
            notes=["LibreOffice layout puede diferir de Word en tablas complejas"],
        )

class HeuristicPageLayoutProvider(PageLayoutProvider):
    """
    Último recurso. Usa la lógica actual: 250 palabras/página + keyword de cuerpo.
    """

    def is_available(self) -> bool:
        return True  # Siempre disponible

    def paginate(self, docx_path: Path, timeout_seconds: int = 30) -> PageLayoutResult:
        import docx
        doc = docx.Document(str(docx_path))

        paragraph_pages: List[int] = []
        current_page = 1
        words_on_current_page = 0

        for p in doc.paragraphs:
            wc = len(p.text.split())
            if words_on_current_page + wc > 250 and wc > 0:
                current_page += 1
                words_on_current_page = 0
            paragraph_pages.append(current_page)
            words_on_current_page += wc

        return PageLayoutResult(
            paragraph_pages=paragraph_pages,
            total_pages=current_page,
            provider_used="heuristic",
            confidence=0.5,
            notes=["Estimación por word_count, no es precisa"],
        )

_cached_provider: Optional[PageLayoutProvider] = None

def get_page_layout_provider() -> PageLayoutProvider:
    """
    Retorna el mejor provider disponible, cacheado a nivel de proceso.
    Orden de preferencia: COM > LibreOffice > Heurístico.
    """
    global _cached_provider
    if _cached_provider is not None:
        return _cached_provider

    providers = [
        COMPageLayoutProvider(),
        LibreOfficePageLayoutProvider(),
        HeuristicPageLayoutProvider(),
    ]
    for p in providers:
        if p.is_available():
            _cached_provider = p
            print(f"[PAGE-LAYOUT] Usando provider: {p.__class__.__name__}")
            return p

    # Imposible llegar aquí porque HeuristicPageLayoutProvider siempre está disponible
    _cached_provider = HeuristicPageLayoutProvider()
    return _cached_provider
