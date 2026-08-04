"""
WordAPA7 — Motor de Estilos APA 7

Aplica reglas estrictas APA 7 a nivel de tipografia, parrafos y titulos (Nivel 1 a 5).
Normaliza la fuente en todo el documento para evitar mezclas indeseadas.
"""

from typing import Optional

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from models import APARuleSet


def set_run_font(run, font_family: str, font_size_pt: float) -> None:
    """
    Aplica fuente y tamaño de forma robusta a nivel de run, incluyendo
    atributos rFonts (ascii, hAnsi, eastAsia, cs) que python-docx no setea
    por defecto. Esto garantiza que Word no herede la fuente del estilo
    de párrafo o del numbering definition.
    """
    run.font.name = font_family
    run.font.size = Pt(font_size_pt)
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), font_family)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(font_size_pt * 2)))


def apply_page_setup(doc: docx.Document, rules: APARuleSet, preserve_landscape: bool = True) -> None:
    """
    Aplica margenes APA 7 (2.54 cm / 1 in) respetando secciones en landscape si existen.

    Tambien normaliza el tamaño de página a Carta (8.5 x 11 in, w=12240 h=15840 twips)
    en secciones portrait — evita que documentos A4/Legal/tabloide salgan con hojas
    demasiado anchas. Las secciones landscape conservan su orientación.
    """
    margin_inches = rules.margins_cm / 2.54
    LETTER_W_TWIPS = "12240"
    LETTER_H_TWIPS = "15840"

    for section in doc.sections:
        is_landscape: bool = False
        try:
            pg_sz = section._sectPr.find(qn("w:pgSz"))
            if pg_sz is not None:
                orient = pg_sz.attrib.get(qn("w:orient"))
                if orient == "landscape":
                    is_landscape = True
        except Exception:
            pass

        # Normalizar tamaño de página a Carta en portrait
        if not (preserve_landscape and is_landscape):
            try:
                pg_sz = section._sectPr.find(qn("w:pgSz"))
                if pg_sz is None:
                    pg_sz = OxmlElement('w:pgSz')
                    section._sectPr.append(pg_sz)
                pg_sz.set(qn("w:w"), LETTER_W_TWIPS)
                pg_sz.set(qn("w:h"), LETTER_H_TWIPS)
                pg_sz.attrib.pop(qn("w:orient"), None)
            except Exception:
                pass

        if not (preserve_landscape and is_landscape):
            section.top_margin = Inches(margin_inches)
            section.bottom_margin = Inches(margin_inches)
            section.left_margin = Inches(margin_inches)
            section.right_margin = Inches(margin_inches)


def normalize_all_fonts(
    doc: docx.Document,
    rules: Optional[APARuleSet] = None,
    target_font: Optional[str] = None,
    target_size_pt: Optional[int] = None,
    skip_body_paragraphs: int = 0,
) -> None:
    """
    Normaliza TODAS las fuentes en el documento a la fuente APA.

    REGLA DURA — se aplica a:
    - Parrafos normales (salvo los primeros `skip_body_paragraphs`, que son la portada)
    - Runs dentro de parrafos (fuentes mixtas pegadas de internet)
    - Texto dentro de celdas de tabla (ALL cells)
    - Encabezados y pies de pagina

    EXCEPCION: Preserva negrita/italica intencional (solo limpia font_name y font_size)
    NO AFECTA: Ecuaciones OMML, objetos OLE, texto en imagenes
    """
    target_font_pt = Pt(target_size_pt)

    # Normalizar parrafos del cuerpo (saltando la portada si se indica)
    for para in doc.paragraphs[skip_body_paragraphs:]:
        _normalize_paragraph_runs(para, target_font, target_font_pt)

    # Normalizar celdas de tabla
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _normalize_paragraph_runs(para, target_font, target_font_pt)

    # Normalizar encabezados y pies de pagina
    for section in doc.sections:
        # Header
        for para in section.header.paragraphs:
            _normalize_paragraph_runs(para, target_font, target_font_pt)
        # Footer
        for para in section.footer.paragraphs:
            _normalize_paragraph_runs(para, target_font, target_font_pt)

    # Normalizar estilos de la plantilla directamente en el XML de styles.xml
    if rules is None:
        rules = APARuleSet()
    update_docx_styles_xml(doc, rules)


def update_docx_styles_xml(doc: docx.Document, rules: APARuleSet) -> None:
    """
    Modifica directamente las definiciones XML en styles.xml de la plantilla/documento:
    - Forzar fuente Times New Roman (12pt / sz=24)
    - Color negro (000000)
    - Interlineado doble (w:line="480", lineRule="auto")
    - Estilos Heading 1 a 5 y Normal / List Bullet
    """
    font_name = rules.font_family
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    try:
        styles_element = doc.styles._element

        # 1. Modificar docDefaults
        docDefaults = styles_element.find(f".//{{{w_ns}}}docDefaults")
        if docDefaults is not None:
            rPr = docDefaults.find(f".//{{{w_ns}}}rPr")
            if rPr is not None:
                rFonts = rPr.find(f"{{{w_ns}}}rFonts")
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)

                sz = rPr.find(f"{{{w_ns}}}sz")
                if sz is None:
                    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="24"/>')
                    rPr.append(sz)
                else:
                    sz.set(qn("w:val"), "24")

            pPr = docDefaults.find(f".//{{{w_ns}}}pPr")
            if pPr is not None:
                spacing = pPr.find(f"{{{w_ns}}}spacing")
                if spacing is None:
                    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:line="480" w:lineRule="auto" w:before="0" w:after="0"/>')
                    pPr.append(spacing)
                else:
                    spacing.set(qn("w:line"), "480")
                    spacing.set(qn("w:lineRule"), "auto")

        # 2. Actualizar cada estilo individual
        heading_configs = {
            "Normal": {"color": "000000", "sz": "24", "line": "480", "bold": False, "italic": False},
            "Heading 1": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": False, "jc": "center"},
            "Heading 2": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": False, "jc": "left"},
            "Heading 3": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": True, "jc": "left"},
            "Heading 4": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": False, "jc": "left"},
            "Heading 5": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": True, "jc": "left"},
            "Título 1": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": False, "jc": "center"},
            "Título 2": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": False, "jc": "left"},
            "Título 3": {"color": "000000", "sz": "24", "line": "480", "bold": True, "italic": True, "jc": "left"},
            "List Bullet": {"color": "000000", "sz": "24", "line": "480"},
            "List Paragraph": {"color": "000000", "sz": "24", "line": "480"},
        }

        for style in styles_element.findall(f".//{{{w_ns}}}style"):
            name_el = style.find(f"{{{w_ns}}}name")
            name_val = name_el.attrib.get(f"{{{w_ns}}}val") if name_el is not None else ""

            cfg = heading_configs.get(name_val)
            if cfg:
                rPr = style.get_or_add_rPr()
                pPr = style.get_or_add_pPr()

                # Fuente
                rFonts = rPr.find(f"{{{w_ns}}}rFonts")
                if rFonts is None:
                    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                    rPr.append(rFonts)
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)

                # Color
                color = rPr.find(f"{{{w_ns}}}color")
                if color is None:
                    color = parse_xml(f'<w:color {nsdecls("w")} w:val="{cfg["color"]}"/>')
                    rPr.append(color)
                else:
                    color.set(qn("w:val"), cfg["color"])

                # Tamaño
                sz = rPr.find(f"{{{w_ns}}}sz")
                if sz is None:
                    sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{cfg["sz"]}"/>')
                    rPr.append(sz)
                else:
                    sz.set(qn("w:val"), cfg["sz"])

                # Negrita / Cursiva
                if cfg.get("bold"):
                    if rPr.find(f"{{{w_ns}}}b") is None:
                        rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
                if cfg.get("italic"):
                    if rPr.find(f"{{{w_ns}}}i") is None:
                        rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))

                # Interlineado y alineacion
                if "line" in cfg:
                    spacing = pPr.find(f"{{{w_ns}}}spacing")
                    if spacing is None:
                        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:line="{cfg["line"]}" w:lineRule="auto" w:before="0" w:after="0"/>')
                        pPr.append(spacing)
                    else:
                        spacing.set(qn("w:line"), cfg["line"])
                        spacing.set(qn("w:lineRule"), "auto")

                # Tipografía avanzada: Widow Control y Keep With Next (Fase 8)
                widow = pPr.find(f"{{{w_ns}}}widowControl")
                if widow is None:
                    pPr.append(parse_xml(f'<w:widowControl {nsdecls("w")}/>'))

                if name_val.startswith("Heading") or name_val.startswith("Título"):
                    keep_next = pPr.find(f"{{{w_ns}}}keepNext")
                    if keep_next is None:
                        pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))

                if "jc" in cfg:
                    jc = pPr.find(f"{{{w_ns}}}jc")
                    if jc is None:
                        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="{cfg["jc"]}"/>')
                        pPr.append(jc)
                    else:
                        jc.set(qn("w:val"), cfg["jc"])

    except Exception as e:
        print(f"[WARN] Error actualizando styles.xml: {e}")


def _normalize_paragraph_runs(para, target_font: str, target_size) -> None:
    """Normaliza fuente y tamano en todos los runs de un parrafo."""
    for run in para.runs:
        run.font.name = target_font
        run.font.size = target_size


def format_heading_paragraph(p, level: int, text: str, rules: APARuleSet, preserve_text: bool = False) -> None:
    """
    Formatea un titulo estrictamente segun los 5 niveles APA 7:
    - Nivel 1: Centrado, Negrita, Caso Titulo
    - Nivel 2: Izquierda, Negrita, Caso Titulo
    - Nivel 3: Izquierda, Negrita y Cursiva, Caso Titulo
    - Nivel 4: Sangria 1.27cm, Negrita, Termina en punto. Texto en la misma linea
    - Nivel 5: Sangria 1.27cm, Negrita y Cursiva, Termina en punto. Texto en la misma linea
    """
    if not preserve_text:
        p.text = ""  # Limpiar runs
    # ponytail: clear inherited style to prevent double indentation (style + direct)
    try:
        p.style = None
    except Exception:
        pass
    p.paragraph_format.line_spacing = rules.line_spacing
    p.paragraph_format.space_before = Pt(rules.space_before_pt)
    p.paragraph_format.space_after = Pt(rules.space_after_pt)

    if preserve_text:
        for r in p.runs:
            set_run_font(r, rules.font_family, rules.font_size_pt)
        return

    if level == 1:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(text)
        run.bold = True
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)

    elif level == 2:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(text)
        run.bold = True
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)

    elif level == 3:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)

    elif level == 4:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(rules.paragraph_indent_cm / 2.54)
        formatted_text: str = text if text.endswith(".") else text + "."
        run = p.add_run(formatted_text + " ")
        run.bold = True
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)

    elif level == 5:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Inches(rules.paragraph_indent_cm / 2.54)
        formatted_text = text if text.endswith(".") else text + "."
        run = p.add_run(formatted_text + " ")
        run.bold = True
        run.italic = True
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)


def format_normal_paragraph(p, text: str, rules: APARuleSet, preserve_text: bool = False) -> None:
    """
    Formatea un parrafo normal APA 7:
    - Fuente uniforme (ej. Times New Roman 12pt)
    - Alineado a la izquierda sin justificar
    - Interlineado 2.0 (doble)
    - Sangria en primera linea de 1.27 cm (0.5 in)
    - Sin espacio entre parrafos (0 pt antes y despues)
    Si preserve_text es True, no destruye el XML interno (usado para proteger ecuaciones y campos).
    """
    if not preserve_text:
        p.text = ""
    # ponytail: clear inherited style to prevent double indentation
    try:
        p.style = None
    except Exception:
        pass
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = rules.line_spacing
    p.paragraph_format.first_line_indent = Inches(rules.paragraph_indent_cm / 2.54)
    p.paragraph_format.space_before = Pt(rules.space_before_pt)
    p.paragraph_format.space_after = Pt(rules.space_after_pt)

    if not preserve_text:
        run = p.add_run(text)
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        for r in p.runs:
            set_run_font(r, rules.font_family, rules.font_size_pt)


def format_block_quote(p, text: str, rules: APARuleSet) -> None:
    """
    Formatea una cita en bloque (mas de 40 palabras):
    - Todo el bloque sangrado 1.27 cm desde la izquierda
    - Interlineado 2.0
    - Sin comillas alrededor
    """
    p.text = ""
    # ponytail: clear inherited style to prevent double indentation
    try:
        p.style = None
    except Exception:
        pass
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(rules.paragraph_indent_cm / 2.54)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = rules.line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    set_run_font(run, rules.font_family, rules.font_size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)


def normalize_global_body_spacing(doc: docx.Document, rules: APARuleSet, cover_paragraph_count: int = 0) -> None:
    """
    Pasada final global: Forzar interlineado doble (w:line=480, lineRule=auto) y 0pt antes/despues
    en el 100% de los parrafos del cuerpo fuera de la portada y celdas de tabla.
    """
    paragraphs = doc.paragraphs[cover_paragraph_count:]
    for p in paragraphs:
        if not p.text.strip():
            continue
        p.paragraph_format.line_spacing = rules.line_spacing
        p.paragraph_format.space_before = Pt(rules.space_before_pt)
        p.paragraph_format.space_after = Pt(rules.space_after_pt)

        # Limpiar negritas residuales en párrafos de cuerpo (no headings)
        style_name = (p.style.name or '') if p.style else ''
        if 'Heading' not in style_name and 'Título' not in style_name:
            all_runs_bold = all(r.bold is True for r in p.runs if r.text.strip())
            if all_runs_bold and p.runs:
                for run in p.runs:
                    if run.bold is True:
                        run.bold = None  # Heredar del estilo
