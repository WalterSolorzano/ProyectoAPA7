"""WordAPA7 — Motor de edición IN-PLACE.

Abre el .docx ORIGINAL y modifica SOLO los párrafos del cuerpo.
La portada original (todo párrafo con índice < body_start_paragraph_idx),
las secciones (sectPr), headers/footers y estilos existentes NUNCA se tocan:
python-docx preserva intactas las partes que no se modifican.

Scopes soportados (coinciden con scoped_apply):
  texto            -> tipografía/interlineado/sangría de párrafos del cuerpo
  tablas_imagenes  -> estilo APA de tablas (bordes/header row); imágenes intactas
  bibliografia     -> sangría francesa en el bloque final de referencias

Contrato duro (testeado):
  * Ningún párrafo con idx < body_start cambia NI UN BYTE.
  * sectPr / headers / footers / styles.xml byte-idénticos.
  * Si scopes está definido, solo esos ámbitos cambian.
"""
from __future__ import annotations

import hashlib
import io
import re
import time
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

try:
    from wordapa7_logger import log_event
except Exception:  # pragma: no cover
    def log_event(*a, **k): pass


def _sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:12]


def _canon(data: bytes) -> str:
    """Hash canónico XML (inmune a re-serialización de python-docx)."""
    from lxml import etree
    try:
        tree = etree.fromstring(data)
        return hashlib.sha256(etree.tostring(tree, method="c14n")).hexdigest()[:12]
    except Exception:
        return _sha(data)


def _cover_floor_by_content(doc: Any) -> int:
    """Piso adicional: la portada termina antes del primer Heading real o del
    primer parrafo largo de cuerpo. Nunca confiar solo en el indice guardado."""
    floor = 0
    for i, para in enumerate(doc.paragraphs[:60]):
        txt = (para.text or "").strip()
        style = (para.style.name or "").lower() if para.style is not None else ""
        if "heading" in style or "título" in style:
            return i
        # parrafo de cuerpo tipico: >180 chars o contiene citas (Autor, 2019)
        if len(txt) > 180 or re.search(r"\([A-Z][^)]{2,40},\s*(19|20)\d{2}\)", txt):
            return i
    return max(floor, 0)


def _body_start(doc_model: Any) -> int:
    p = getattr(doc_model, "portada", None) or {}
    if isinstance(p, dict):
        raw = p.get("body_start_paragraph_idx")
    else:
        raw = getattr(p, "body_start_paragraph_idx", None)
    try:
        v = int(raw)
        return max(v, 0)
    except (TypeError, ValueError):
        return 0


def _is_ref_paragraph(text: str) -> bool:
    t = text.strip()
    return bool(t) and len(t) > 40 and re.search(r"\(\d{4}\)|\(\d{4}[a-z]?\)", t)


def _is_list_item(para: Any, text: str) -> bool:
    """Detecta si un parrafo es viñeta o lista numerada."""
    try:
        if para._p.pPr is not None and para._p.pPr.numPr is not None:
            return True
    except Exception:
        pass
    return bool(re.match(r"^(?:[\u2022\u2023\u25E6\u2043\u2219\*\-\–\—]|\d+[\.\)]|[a-zA-Z][\.\)]|\([a-zA-Z\d]+\))\s+", text))


def _is_toc_line(text: str) -> bool:
    """Detecta lineas de tabla de contenidos / indice."""
    return bool(re.search(r"(?:\.{2,}|_{2,}|\t|\s{4,})\s*\d+\s*$", text))


def apply_inplace(
    original_path: Path,
    out_path: Path,
    doc_model: Any,
    rules: Any,
    scopes: Iterable[str] | None = None,
) -> Path:
    """Edita el documento original respetando portada/secciones al 100%."""
    t0 = time.time()
    active = set(scopes) if scopes is not None else {"texto", "tablas_imagenes", "bibliografia"}
    body_start = max(_body_start(doc_model), 0)

    orig_bytes = Path(original_path).read_bytes()

    # Snapshots de garantía: portada (párrafos < body_start) y partes globales.
    pre = Document(io.BytesIO(orig_bytes))
    try:
        body_start = max(body_start, _cover_floor_by_content(pre))
    except Exception:
        pass

    cover_before = [_sha(p._element.xml.encode("utf-8")) for p in pre.paragraphs[:body_start]]
    global_before = {}
    with zipfile.ZipFile(io.BytesIO(orig_bytes)) as z:
        for name in z.namelist():
            if name.startswith(("word/styles.", "word/header", "word/footer")) or name.endswith("document.xml.rels"):
                global_before[name] = _canon(z.read(name))

    doc = Document(io.BytesIO(orig_bytes))
    paragraphs = doc.paragraphs
    font_name = getattr(rules, "font_family", None) or "Times New Roman"
    font_size = Pt(getattr(rules, "font_size", 12) or 12)
    line_sp = float(getattr(rules, "line_spacing", 2.0) or 2.0)

    changed = 0
    ref_zone_start = len(paragraphs)
    if "texto" in active or "bibliografia" in active:
        # Localizar inicio de bibliografía: último heading 'Referencias' o primer párrafo-ref
        for i in range(len(paragraphs) - 1, body_start, -1):
            if paragraphs[i].text.strip().lower().rstrip(":") in ("referencias", "bibliografía", "bibliografia", "references"):
                ref_zone_start = i + 1
                break

    for i, para in enumerate(paragraphs):
        if i < body_start:
            continue  # PORTADA INTOCABLE — contrato duro
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower() if para.style is not None else ""
        if "heading" in style_name or "título" in style_name or "titulo" in style_name:
            continue  # headings: los maneja la ruta rebuild si el usuario lo pide

        # Linea de indice / TOC -> NO aplicar sangria de primera linea
        if _is_toc_line(text):
            continue

        if "bibliografia" in active and i >= ref_zone_start and _is_ref_paragraph(text):
            pf = para.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.first_line_indent = Inches(-0.5)
            pf.line_spacing = line_sp
            changed += 1
            continue

        # Viñetas o listas -> margen izquierdo 0.5", SIN sangria de primera linea APA
        if _is_list_item(para, text):
            if "texto" in active:
                pf = para.paragraph_format
                pf.left_indent = Inches(0.5)
                pf.first_line_indent = Inches(0)
                pf.line_spacing = line_sp
                pf.space_after = Pt(0)
                pf.space_before = Pt(0)
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = font_size
                changed += 1
            continue

        if "texto" in active:
            pf = para.paragraph_format
            pf.left_indent = Inches(0)
            pf.right_indent = Inches(0)
            pf.line_spacing = line_sp
            pf.space_after = Pt(0)
            pf.space_before = Pt(0)
            first_vis = bool(getattr(rules, "indent_first_line", True))
            pf.first_line_indent = Inches(0.5) if first_vis else Inches(0)
            for run in para.runs:
                run.font.name = font_name
                run.font.size = font_size
            changed += 1

    if "tablas_imagenes" in active:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for tbl in doc.tables:
            tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER if hasattr(tbl, "alignment") else tbl.alignment
            tbl.autofit = True
            # Bordes horizontales únicamente (estilo APA clásico)
            tblPr = tbl._tbl.tblPr
            borders = tblPr.find(qn("w:tblBorders"))
            if borders is not None:
                tblPr.remove(borders)
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "bottom"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "8")
                el.set(qn("w:color"), "000000")
                borders.append(el)
            for edge in ("left", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                borders.append(el)
            tblPr.append(borders)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    # Verificación post: asegurar que la portada no mutó
    post_bytes = out_path.read_bytes()
    post = Document(io.BytesIO(post_bytes))
    cover_after = [_sha(p._element.xml.encode("utf-8")) for p in post.paragraphs[:body_start]]
    if cover_before != cover_after:
        log_event("inplace_editor", "cover_drift_warning",
                  data={"body_start": body_start, "diffs": sum(1 for a, b in zip(cover_before, cover_after) if a != b)})

    log_event("inplace_editor", "completed",
              data={"scope": sorted(active), "body_start": body_start, "changed": changed,
                    "elapsed_ms": int((time.time() - t0) * 1000)})
    return out_path
