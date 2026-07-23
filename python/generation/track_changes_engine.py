"""
WordAPA7 — Motor de Track Changes OOXML (Comparacion de Cambios)

Genera un archivo .docx con marcas de revision de Word (<w:ins> y <w:del>)
para mostrar las modificaciones realizadas por el sistema APA 7 sobre el documento original.
"""

from datetime import datetime
from pathlib import Path

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from models import DocumentModel, ElementType, ElementModel, APARuleSet


def _make_ins_run(text: str, author: str = "WordAPA7") -> OxmlElement:
    """
    Crea un elemento <w:ins> ... </w:ins> con el texto insertado.
    Produce OOXML valido que Word reconoce como Track Change.
    """
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"))

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rPr_elements = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'  <w:color w:val="107C10"/>'
        f'</w:rPr>'
    )
    r.append(rPr_elements)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _make_del_run(text: str, author: str = "WordAPA7") -> OxmlElement:
    """
    Crea un elemento <w:del> ... </w:del> con el texto eliminado.
    Produce OOXML valido que Word reconoce como Track Change.
    """
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), "2")
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"))

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rPr_elements = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'  <w:color w:val="A4262C"/>'
        f'  <w:strike w:val="true"/>'
        f'</w:rPr>'
    )
    r.append(rPr_elements)
    del_text = OxmlElement("w:delText")
    del_text.set(qn("xml:space"), "preserve")
    del_text.text = text
    r.append(del_text)
    del_elem.append(r)
    return del_elem


def create_tracked_changes_docx(
    doc_model: DocumentModel,
    output_filepath: Path | str,
    rules: APARuleSet | None = None,
) -> Path:
    """
    Genera un documento Word con control de cambios OOXML real (<w:ins> / <w:del>).
    Word muestra estos cambios como revisiones reales que se pueden aceptar o rechazar.
    """
    if rules is None:
        rules = APARuleSet()

    out_path = Path(output_filepath)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()

    # Configurar margenes
    margin_inches: float = rules.margins_cm / 2.54
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

    # Titulo del documento de comparacion
    p_title = doc.add_paragraph()
    p_title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("Control de Cambios — Formato APA 7 Aplicado")
    r_t.bold = True
    r_t.font.name = rules.font_family
    r_t.font.size = Pt(14)

    # Subtitulo
    p_sub = doc.add_paragraph()
    p_sub.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(
        "Las marcas en rojo (tachado) muestran texto original. "
        "Las marcas en verde muestran texto corregido con APA 7."
    )
    r_sub.font.name = rules.font_family
    r_sub.font.size = Pt(10)
    r_sub.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Espacio

    ins_id_counter: int = 0
    del_id_counter: int = 0

    for item in doc_model.elements:
        # Obtener texto del elemento
        if isinstance(item, dict):
            elem = ElementModel.model_validate(item)
        else:
            elem = item

        text: str = (elem.text or "").strip()
        orig_text: str = (elem.original_text or "").strip()
        if not orig_text:
            orig_text = text

        if not text and not orig_text:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = rules.line_spacing
        p.paragraph_format.first_line_indent = Inches(rules.paragraph_indent_cm / 2.54)

        # Heading handling
        if elem.type == ElementType.HEADING:
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(12)

        # Si el texto cambio (fue modificado por el sistema)
        if orig_text != text:
            # Mostrar eliminado con tag OOXML <w:del>
            del_id_counter += 1
            del_elem = _make_del_run(
                f"[Original] {orig_text}", f"WordAPA7-{del_id_counter}"
            )
            del_elem.set(qn("w:id"), str(del_id_counter))
            p._element.append(del_elem)

            # Separador
            sep_run = OxmlElement("w:r")
            sep_br = OxmlElement("w:br")
            sep_run.append(sep_br)
            p._element.append(sep_run)

            # Mostrar insertado con tag OOXML <w:ins>
            ins_id_counter += 1
            ins_elem = _make_ins_run(
                f"[APA 7] {text}", f"WordAPA7-{ins_id_counter}"
            )
            ins_elem.set(qn("w:id"), str(ins_id_counter))
            p._element.append(ins_elem)
        else:
            # Sin cambios: texto normal
            r = p.add_run(text)
            r.font.name = rules.font_family
            r.font.size = Pt(rules.font_size_pt)

    doc.save(str(out_path))
    return out_path
