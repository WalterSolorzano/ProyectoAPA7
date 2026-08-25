"""
WordAPA7 — Generador Final de Documentos DOCX (Con Modificación Directa sobre Original)

Abre el archivo original.docx de la sesión y le aplica los estilos APA 7
sin destruir imágenes, tablas, shapes, logos, cuadros de texto o fórmulas.
"""

from copy import deepcopy
from pathlib import Path
from typing import List

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from models import (
    APARuleSet,
    DocumentModel,
    ElementModel,
    ElementType,
    PortadaData,
    ReferenciaModel,
)
from modules.cover_designer import CoverTemplate, apply_cover_to_document, list_cover_templates
from modules.portada_module import format_apa_portada
from modules.referencias_module import format_apa_referencias_section
from parsing.pre_classifier import REGEX_FIGURE_CAPTION

from generation.bullet_engine import format_bullet_item, format_numbered_item
from generation.document_structure import setup_apa_header
from generation.image_handler import format_apa_figure
from generation.style_engine import (
    apply_page_setup,
    format_block_quote,
    format_heading_paragraph,
    format_normal_paragraph,
    set_run_font,
)
from generation.table_engine import (
    fit_table_to_page,
    format_apa_table,
    set_table_borders,
    validate_table_widths,
)


def _to_roman(num: int) -> str:
    """Convierte entero (1-50) a numeral romano en mayúsculas."""
    if num <= 0:
        return str(num)
    val = [50, 40, 10, 9, 5, 4, 1]
    syms = ["L", "XL", "X", "IX", "V", "IV", "I"]
    roman = ""
    n = num
    for i, v in enumerate(val):
        while n >= v:
            roman += syms[i]
            n -= v
    return roman


def _strip_inline_footnote_markers(text: str) -> str:
    """Elimina los marcadores inline '(nota N)' que el parser inserta en el texto."""
    import re
    if not text:
        return text or ""
    return re.sub(r'\s*\(nota\s+\d+\)', '', text).strip()


def _detect_heading_numbering_style(heading_text: str) -> str:
    """
    Detecta si el texto del heading usa numeración romana (I., II., III.)
    o decimal (1., 2., 3.) basado en el prefijo original.
    Retorna 'roman' o 'decimal'.
    """
    import re
    text_stripped = heading_text.strip()
    first_word = text_stripped.split()[0] if text_stripped else ""
    if re.match(r'^(?:X{0,3})(?:I[XV]|V?I{1,3})\.$', first_word):
        return 'roman'
    return 'decimal'


def _build_heading_prefix(counters: dict[int, int], level: int, numbering_style: str = 'decimal') -> str:
    """Construye el prefijo numérico para un heading según su nivel y estilo.

    Estilos:
    - 'decimal': 1., 1.1., 1.1.1. (numeros arabigos)
    - 'roman': I., II., III. para nivel 1, seguido de decimal para subniveles
    - 'none': sin prefijo numerico
    """
    # Solo numerar H1 y H2. Niveles 3+ no llevan numeracion (APA 7 no lo requiere
    # y resulta visualmente cargado con demasiados digitos).
    if numbering_style == 'none' or level > 2:
        return ""

    # Safety check: a sub-heading (level > 1) must have a numbered parent.
    # If the parent level counter is 0 (no prior heading at that level), don't
    # number this heading at all — a sub-heading without a parent heading
    # should not be numbered. This prevents bogus prefixes like "0.1." from
    # appearing when a level 2 heading shows up before any level 1 heading.
    if level > 1 and counters.get(level - 1, 0) == 0:
        return ""

    parts = []
    for l in range(1, level + 1):
        c = counters.get(l, 0)
        if c > 0:
            if numbering_style == 'roman' and l == 1:
                parts.append(_to_roman(c))
            else:
                parts.append(str(c))
    if parts:
        return ".".join(parts) + ". "
    return ""


def _render_equation_number(number: str, number_format: str) -> str:
    """Renderiza el número de ecuación según el formato elegido.

    Formatos soportados (el token {n} se reemplaza por el número):
      "(1)"      → (1)
      "[1]"      → [1]
      "1."       → 1.
      "(1.1)"    → (1.1)
      "Ecuación {n}" → Ecuación 1
    Si el formato contiene "{n}", se usa como plantilla; si no, se
    asume el número envuelto en paréntesis (convención APA/tesis).
    """
    number = str(number or "1").strip()
    if "{n}" in number_format:
        return number_format.replace("{n}", number)
    # Formatos envolventes sin token: insertar el número en el lugar de los dígitos
    low = number_format.lower()
    if low in ("(1)", "[1]", "1.", "(1.1)"):
        template = {
            "(1)": "({})",
            "[1]": "[{}]",
            "1.": "{}.",
            "(1.1)": "({})",
        }[number_format]
        return template.format(number)
    return f"({number})"


def _apply_equation_number(p, display: str, eq_cfg, doc=None) -> None:
    """Agrega el número de ecuación al párrafo con un tab stop derecho.

    - Inserta un carácter tab antes del número para que quede alineado al
      margen derecho.
    - El run del número usa la fuente configurada (sin tocar el XML OMML).
    """
    from docx.enum.text import WD_TAB_ALIGNMENT

    # Tab stop derecho pegado al margen derecho de la sección activa
    # (usable − offset); fallback 6.3" APA si no hay dimensiones.
    tab_position = 6.3
    try:
        usable_in = _usable_width_inches(doc.sections[0] if (doc and doc.sections) else None)
        if usable_in:
            tab_position = max(3.0, usable_in - 0.2)
    except Exception:
        pass
    tabs = p.paragraph_format.tab_stops
    try:
        tabs.add_tab_stop(Inches(tab_position), WD_TAB_ALIGNMENT.RIGHT)
    except Exception:
        pass
    try:
        r = p.add_run("\t" + display)
        set_run_font(r, eq_cfg.font_name or "Times New Roman", eq_cfg.font_size_pt or 12.0)
    except Exception:
        pass


def _generate_toc_from_headings(
    doc: docx.Document,
    doc_model: DocumentModel,
    rules: APARuleSet,
    toc_paragraph,
) -> None:
    """
    Genera un índice / tabla de contenidos estático con dot leaders
    basado en los headings detectados en el documento.

    Formato:
        Introducción ................................. 3
            Antecedentes .............................. 4
            Objetivos ................................. 5
        Marco Teórico ................................ 7

    Si el documento tiene un TOC nativo (ElementType.TOC), se reemplaza
    por este contenido generado.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    # Recolectar headings del modelo con sus niveles
    toc_headings: list[tuple[int, str]] = []
    for elem in doc_model.elements:
        if elem.type == ElementType.HEADING and elem.heading_level and elem.text:
            clean_text = elem.text.strip()
            # Saltar headings de índice para evitar recursión
            lower = clean_text.lower()
            if any(kw in lower for kw in ["índice", "indice", "tabla de contenido", "contenido"]):
                continue
            toc_headings.append((elem.heading_level, clean_text))

    if not toc_headings:
        # Sin headings: insertar marcador de posición
        toc_paragraph.text = ""
        r = toc_paragraph.add_run("[Tabla de Contenidos — No se detectaron títulos en el documento]")
        r.italic = True
        r.font.name = rules.font_family
        r.font.size = Pt(rules.font_size_pt)
        return

    # Limpiar el párrafo TOC original
    toc_paragraph.text = ""
    toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Título del índice (opcional pero bueno para la UI)
    r_title = toc_paragraph.add_run("Índice / Tabla de Contenidos\n")
    r_title.bold = True
    r_title.font.name = rules.font_family
    r_title.font.size = Pt(rules.font_size_pt)

    new_p = OxmlElement('w:p')
    toc_paragraph._element.addnext(new_p)
    toc_field_p = docx.text.paragraph.Paragraph(new_p, doc)

    # Inyectar el campo dinámico de Word
    run1 = toc_field_p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)

    run2 = toc_field_p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = toc_field_p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = toc_field_p.add_run("[El índice se actualizará automáticamente al abrir en Word. Si no se actualiza, presione Clic Derecho -> Actualizar Campos]")
    run4.font.name = rules.font_family
    run4.font.size = Pt(10)
    run4.italic = True

    run5 = toc_field_p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


# ─── HELPER FUNCTIONS ────────────────────────────────────────────────────────

def _strip_existing_numbering(text: str) -> str:
    """Remove existing numbering prefix from heading text to avoid double prefixes.

    Handles patterns like 'I. Title', 'II. Title', 'A. Title', '1. Title',
    '0.1 Title', '1.1 Title', '1.1.1 Title' (multi-level decimal numbering
    with or without a trailing dot).
    """
    if not text:
        return text
    import re
    t = text.strip()
    # Roman prefixes: I. II. III. IV. V. VI. VII. VIII. IX. X.
    m = re.match(
        r'^(M{0,4}(?:CM|CD|D?C{0,3})(?:XC|XL|L?X{0,3})(?:IX|IV|V?I{0,3}))\.\s+(.*)',
        t, re.IGNORECASE
    )
    if m:
        return m.group(2)
    # Multi-level decimal prefixes: 0.1, 1.1, 1.1.1, etc. followed by a space.
    # The optional trailing dot also covers build-style prefixes like "1.1. ".
    # Checked before the single-level decimal so that '1.1 Title' is stripped as
    # a whole rather than leaving the '.1 Title' remainder behind.
    m_multi = re.match(r'^(\d+(?:\.\d+)+)\.?\s+(.*)', t)
    if m_multi:
        return m_multi.group(2)
    # Decimal prefixes: 1. 2. 10.
    m2 = re.match(r'^(\d+)\.\s+(.*)', t)
    if m2:
        return m2.group(2)
    # Letter prefixes: A. B. a. b.
    m3 = re.match(r'^([A-Za-z])\.\s+(.*)', t)
    if m3:
        return m3.group(2)
    return text


# ponytail: removed _find_matching_paragraph_idx — text fingerprint matching
# was inserting page breaks at wrong positions when model text ≠ paragraph text.
# doc_model.elements order ALREADY matches doc.paragraphs[cover_paragraph_count:] order
# (both come from the same DOCX). Sequential access is correct and deterministic.

def _is_table_too_wide(table_info) -> tuple[bool, bool]:
    """Determine if a table needs landscape orientation.

    Returns (too_many_columns, too_wide_text).
    - too_many_columns: True if > 6 columns
    - too_wide_text: True if any cell text exceeds ~80 chars
    """
    if not table_info:
        return False, False
    col_count = 0
    if table_info.headers:
        col_count = len(table_info.headers)
    elif table_info.rows and table_info.rows:
        col_count = len(table_info.rows[0])

    too_many = col_count > 6
    too_wide = False
    if table_info.rows:
        for row in table_info.rows:
            for cell_text in row:
                if len(cell_text) > 100:
                    too_wide = True
                    break
            if too_wide:
                break
    return too_many, too_wide


def _detect_existing_figure_caption(img_paragraph, all_paragraphs: list) -> str | None:
    """Detecta si ya existe un caption de figura en el texto del parrafo de
    la imagen o en el parrafo inmediatamente anterior. Retorna el texto del
    caption encontrado o None. Usa la MISMA semantica de deteccion que el
    parser (REGEX_FIGURE_CAPTION: Figura|Figure|Fig. + digito) para que las
    etiquetas detectadas en parseo no queden huerfanas aqui."""
    fig_pattern = REGEX_FIGURE_CAPTION
    # Revisar el propio parrafo de la imagen
    img_text = img_paragraph.text.strip()
    if fig_pattern.match(img_text):
        return img_text
    # Revisar el parrafo inmediatamente anterior
    prev_elem = img_paragraph._element.getprevious()
    if prev_elem is not None:
        for p in all_paragraphs:
            if p._element is prev_elem:
                prev_text = p.text.strip()
                if fig_pattern.match(prev_text):
                    return prev_text
                break
    return None


def _update_existing_caption_text(img_paragraph, new_caption: str, rules: APARuleSet) -> None:
    """Actualiza el texto del caption existente en el parrafo (sin duplicar el parrafo)."""
    from docx.shared import Pt
    # Buscar si el caption esta en el parrafo de la imagen o en el anterior
    for para in [img_paragraph]:
        text = para.text.strip()
        if text:
            # Preservar el prefijo "Figura X." si existe, solo reemplazar la descripcion
            import re
            m = re.match(r'^(.*?\d+[\.:)]\s*)', text)
            prefix = m.group(1) if m else ""
            para.text = ""
            r = para.add_run(prefix + new_caption)
            r.italic = True
            r.font.name = rules.font_family
            r.font.size = Pt(rules.font_size_pt)
            r.font.color.rgb = RGBColor(0, 0, 0)
            return


def _wrap_in_landscape_section(doc: docx.Document, element) -> None:
    """Wrap an XML element in landscape section breaks."""
    parent = element.getparent()
    if parent is None:
        return

    from copy import deepcopy

    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    try:
        children = list(parent)
        idx = children.index(element)
    except (ValueError, AttributeError):
        return

    # OOXML section break semantics:
    # w:sectPr inside a paragraph's w:pPr defines the properties of the section
    # that ENDS at that paragraph. The NEXT content belongs to a new section
    # whose properties are defined by the NEXT section break (or body-level sectPr).
    #
    # For a landscape sandwich we need:
    #   [portrait content] → before_p (PORTRAIT sectPr) → [table] → after_p (LANDSCAPE sectPr) → [portrait content]

    # Clonar el sectPr ACTUAL (pgSz/pgMar/header/footer refs) para no
    # re-dimensionar documentos A4 a Letter dentro del sándwich.
    current_sect = None
    try:
        body = doc._element.body
        sect_prs = body.findall(qn('w:sectPr'))
        if sect_prs:
            current_sect = deepcopy(sect_prs[-1])
    except Exception:
        current_sect = None

    if current_sect is not None:
        portrait_sect = current_sect
        landscape_sect = deepcopy(current_sect)
        pg_sz = landscape_sect.find(qn('w:pgSz'))
        if pg_sz is None:
            pg_sz = OxmlElement('w:pgSz')
            landscape_sect.insert(0, pg_sz)
        try:
            w_val = int(pg_sz.get(qn('w:w'), '12240'))
            h_val = int(pg_sz.get(qn('w:h'), '15840'))
        except ValueError:
            w_val, h_val = 12240, 15840
        pg_sz.set(qn('w:w'), str(max(w_val, h_val)))
        pg_sz.set(qn('w:h'), str(min(w_val, h_val)))
        pg_sz.set(qn('w:orient'), 'landscape')
    else:
        # Fallback Letter cuando no hay sectPr clonable
        portrait_sect = parse_xml(
            f'<w:sectPr {nsdecls("w")}>'
            f'<w:pgSz w:w="12240" w:h="15840" w:orient="portrait"/>'
            f'</w:sectPr>'
        )
        landscape_sect = parse_xml(
            f'<w:sectPr {nsdecls("w")}>'
            f'<w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/>'
            f'</w:sectPr>'
        )

    # Section break BEFORE the table: ends the preceding PORTRAIT section
    before_p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr/></w:p>')
    before_p.find(qn('w:pPr')).append(portrait_sect)

    # Section break AFTER the table: ends the LANDSCAPE section that contains the table
    after_p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr/></w:p>')
    after_p.find(qn('w:pPr')).append(landscape_sect)

    # Insert before the element
    parent.insert(idx, before_p)
    # Insert after the element (idx + 2 because before_p shifted by 1)
    parent.insert(idx + 2, after_p)


def _usable_width_inches(section) -> float | None:
    """Ancho utilizable de la sección (page − márgenes) en pulgadas, o None."""
    try:
        if (section is None or section.page_width is None
                or section.left_margin is None or section.right_margin is None):
            return None
        return float((section.page_width - section.left_margin - section.right_margin) / 914400)
    except Exception:
        return None


def _usable_width_cm(section) -> float | None:
    """Ancho utilizable de la sección en centímetros, o None."""
    usable_in = _usable_width_inches(section)
    return usable_in * 2.54 if usable_in is not None else None


def _apply_image_design_style(
    p,
    img_elem: ElementModel,
    rules: APARuleSet,
) -> None:
    """Apply design_style formatting to an image paragraph.

    Styles:
    - 'standard' — centered, caption below (APA default)
    - 'sidebar' — right-floating, text wraps around
    - 'scientific' — thin black border, caption below with auto-numbering
    - 'corner' — corner-positioned with text flowing around
    """
    design_style = 'standard'
    if img_elem.image_info:
        design_style = getattr(img_elem.image_info, 'design_style', 'standard')

    if design_style == 'standard':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    elif design_style == 'sidebar':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Tight wrapping via XML
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        for drawing in p._element.findall(qn('w:drawing')):
            for anchor in drawing.findall(qn('wp:anchor')):
                anchor.set('wrapSquare', '')
                anchor.set('wrapDistanceLeft', '914400')
                anchor.set('wrapDistanceTop', '0')

    elif design_style == 'scientific':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Add thin black border via paragraph shading / border
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        pPr = p._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is None:
            from docx.oxml import OxmlElement
            pPr = OxmlElement('w:pPr')
            p._element.insert(0, pPr)
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="4" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="4" w:space="4" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="4" w:space="4" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    elif design_style == 'corner':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Floating positioning via XML manipulation
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        for drawing in p._element.findall(qn('w:drawing')):
            for inline in drawing.findall(qn('wp:inline')):
                # Convert inline to anchor (floating)
                anchor_xml = (
                    '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
                    'behindDoc="0" layoutInCell="1" allowOverlap="1" locked="0" '
                    'relativeHeight="251658240" simplePos="0">'
                    '<wp:simplePos x="0" y="0"/>'
                    '<wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>'
                    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
                    '</wp:anchor>'
                )
                # This is a simplified approach; full anchor conversion is complex
                pass

    p.paragraph_format.first_line_indent = Inches(0)


# ─── REHIDRATACIÓN DE ANCLAS / ENLACES / NOTAS (POST-PROCESO) ────────────────

def rehydrate_hyperlinks(p, element, doc) -> None:
    """Best-effort: re-emite hyperlinks reales (`w:hyperlink` + relación externa).

    Tras formatear el párrafo, busca en sus runs el texto de cada hyperlink
    almacenado en `element.hyperlinks` ([{"text", "url"}]) y envuelve la porción
    coincidente en un `w:hyperlink` con rId externo. Divide el run en
    (antes, enlace, después) cuando el enlace aparece dentro de un run mayor.
    Si el texto ya vive dentro de un `w:hyperlink` preservado del original, no lo toca.
    """
    try:
        hyperlinks = getattr(element, 'hyperlinks', None) or []
        if not hyperlinks:
            return
        for hl in hyperlinks:
            if not isinstance(hl, dict):
                continue
            hl_text = str(hl.get('text') or '').strip()
            url = str(hl.get('url') or '')
            if not hl_text or not url:
                continue
            try:
                r_id = doc.part.relate_to(url, RT.HYPERLINK, is_external=True)
            except Exception:
                continue
            for run in list(p.runs):
                try:
                    run_text = ''.join(t.text or '' for t in run._r.findall(qn('w:t')))
                except Exception:
                    continue
                if not run_text or hl_text not in run_text:
                    continue
                # Evitar re-envolver si ya es un hyperlink (preservado del original)
                if run._r.getparent() is not None and run._r.getparent().tag == qn('w:hyperlink'):
                    continue
                try:
                    before, sep, after = run_text.partition(hl_text)
                    if not sep:
                        continue
                    parent = run._r.getparent()
                    if parent is None:
                        continue
                    idx = list(parent).index(run._r)

                    def _clone_run(part_text: str):
                        new_r = OxmlElement('w:r')
                        for child in list(run._r):
                            try:
                                local = child.tag.split('}')[-1] if '}' in str(child.tag) else str(child.tag)
                            except AttributeError:
                                continue
                            if local in ('t', 'tab', 'br', 'cr', 'noBreakHyphen'):
                                continue
                            new_r.append(deepcopy(child))
                        t = OxmlElement('w:t')
                        t.set(qn('xml:space'), 'preserve')
                        t.text = part_text
                        new_r.append(t)
                        return new_r

                    if before:
                        parent.insert(idx, _clone_run(before))
                        idx += 1
                    link_run = _clone_run(sep)
                    # Estilo de enlace típico de Word (subrayado + azul)
                    try:
                        link_rPr = link_run.find(qn('w:rPr'))
                        if link_rPr is None:
                            link_rPr = OxmlElement('w:rPr')
                            link_run.insert(0, link_rPr)
                        if link_rPr.find(qn('w:u')) is None:
                            u = OxmlElement('w:u')
                            u.set(qn('w:val'), 'single')
                            link_rPr.append(u)
                        color = link_rPr.find(qn('w:color'))
                        if color is None:
                            color = OxmlElement('w:color')
                            link_rPr.append(color)
                        color.set(qn('w:val'), '0563C1')
                    except Exception:
                        pass
                    h = OxmlElement('w:hyperlink')
                    h.set(qn('r:id'), r_id)
                    h.append(link_run)
                    parent.insert(idx, h)
                    idx += 1
                    if after:
                        parent.insert(idx, _clone_run(after))
                    parent.remove(run._r)
                except Exception:
                    continue
                break
    except Exception:
        pass


def rehydrate_bookmarks(p, element) -> None:
    """Best-effort: inserta `w:bookmarkStart`/`w:bookmarkEnd` alrededor del
    primer run del párrafo usando los nombres/ids de `element.bookmarks`."""
    try:
        bookmarks = getattr(element, 'bookmarks', None) or []
        if not bookmarks:
            return
        runs = list(p.runs)
        if not runs:
            return
        anchor_run = runs[0]._r
        parent = anchor_run.getparent()
        if parent is None:
            return
        for bm in bookmarks:
            try:
                if not isinstance(bm, dict):
                    continue
                name = str(bm.get('name') or '')
                if not name:
                    continue
                try:
                    bm_id = int(bm.get('id', 0))
                except (TypeError, ValueError):
                    bm_id = 0
                idx = list(parent).index(anchor_run)
                start = OxmlElement('w:bookmarkStart')
                start.set(qn('w:id'), str(bm_id))
                start.set(qn('w:name'), name)
                parent.insert(idx, start)
                end = OxmlElement('w:bookmarkEnd')
                end.set(qn('w:id'), str(bm_id))
                parent.insert(idx + 1, end)
            except Exception:
                continue
    except Exception:
        pass


def rehydrate_footnotes(p, element, footnotes) -> None:
    """Best-effort: notas al pie como endnotes simples.

    Si `element.footnote_ids` referencia notas en `footnotes` ([{"id", "text"}]),
    agrega un marcador superíndice "[N]" al párrafo y, después de él, un párrafo
    de fuente pequeña "N <texto>".
    """
    try:
        footnote_ids = getattr(element, 'footnote_ids', None) or []
        if not footnote_ids:
            return
        id_to_text: dict = {}
        for fn in (footnotes or []):
            if isinstance(fn, dict):
                try:
                    id_to_text[str(fn.get('id'))] = str(fn.get('text') or '')
                except Exception:
                    continue
        anchor = p._element
        for fn_id in footnote_ids:
            try:
                fid = int(fn_id)
            except (TypeError, ValueError):
                continue
            try:
                marker = p.add_run(f"[{fid}]")
                marker.font.superscript = True
                marker.font.size = Pt(10)
            except Exception:
                pass
            text = id_to_text.get(str(fid), "")
            if not text:
                continue
            try:
                new_p_xml = OxmlElement('w:p')
                anchor.addnext(new_p_xml)
                anchor = new_p_xml
                para = Paragraph(new_p_xml, p._parent)
                r = para.add_run(f"{fid} {text}")
                r.font.size = Pt(10)
                r.font.name = p.runs[0].font.name if p.runs and p.runs[0].font.name else None
            except Exception:
                pass
    except Exception:
        pass


def generate_apa7_docx(
    doc_model: DocumentModel,
    output_filepath: Path | str,
    rules: APARuleSet | None = None,
    portada: PortadaData | None = None,
    references: List[ReferenciaModel] | None = None,
    remove_cover_paragraphs: bool = False,
) -> Path:
    """
    Genera el archivo .docx final aplicando todas las reglas APA 7.
    Modifica el archivo original.docx de la sesión conservando shapes, logos y layout intactos.
    """
    if rules is None:
        rules = APARuleSet()

    # Check if we should use the layered (from-scratch) generator
    import os
    use_layered = os.getenv("WORDAPA7_LAYERED_GEN", "0") == "1"

    if use_layered:
        from generation.layered_generator import generate_apa7_from_scratch
        return generate_apa7_from_scratch(doc_model, output_filepath, rules, portada, references)

    out_path = Path(output_filepath)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # 1. Intentar cargar el archivo original.docx de la sesión
    session_id = doc_model.session_id
    original_file = out_path.parent / "original.docx"

    if not original_file.exists() and doc_model.file_name:
        fname_orig = out_path.parent / doc_model.file_name
        if fname_orig.exists():
            original_file = fname_orig

    if not original_file.exists() and session_id:
        alt_orig = Path("storage") / "sessions" / session_id / "original.docx"
        if alt_orig.exists():
            original_file = alt_orig

    if original_file.exists():
        doc = docx.Document(original_file)
    else:
        # Fallback si no existe original (ej: pruebas unitarias sintǸticas)
        doc = docx.Document()

    # H20: si el original ya trae índice, actualizarlo (campo TOC) y no duplicar.
    try:
        from generation.document_structure import refresh_or_flag_existing_toc
        _toc_kind = refresh_or_flag_existing_toc(doc)
        if _toc_kind:
            print(f"[TOC] Índice existente detectado ({_toc_kind}): se actualiza, no se duplica")
    except Exception as _e:
        print(f"[WARN] [TOC] revisión de índice omitida: {_e}")

    # 2. Configuración de página y márgenes
    apply_page_setup(doc, rules, preserve_landscape=doc_model.has_landscape_sections)

    # 2.5 Reescritura dura de los estilos base en styles.xml (Normal, Headings 1-5, Listas)
    from generation.style_engine import update_docx_styles_xml
    update_docx_styles_xml(doc, rules)

    # 3. Encabezado y número de página
    running_head: str = portada.running_head if portada else ""
    setup_apa_header(doc, doc_model.apa_format, running_head, rules)

    # 3.5 Cover Designer: Si la portada tiene una plantilla seleccionada, anteponerla AHORA
    # (antes de que el sistema procese cualquier portada original)
    cover_template_applied = False
    if doc_model.portada:
        cover_template_id = None
        if isinstance(doc_model.portada, dict):
            cover_template_id = doc_model.portada.get('cover_template_id')
        elif hasattr(doc_model.portada, 'get'):
            cover_template_id = doc_model.portada.get('cover_template_id')

        if cover_template_id:
            print(f"[COVER-DESIGNER] Buscando plantilla '{cover_template_id}'...")
            try:
                base_dir = Path(__file__).resolve().parent.parent
                templates = list_cover_templates(base_dir)
                print(f"[COVER-DESIGNER] {len(templates)} plantillas disponibles: {[t.name for t in templates]}")
                selected_template: CoverTemplate | None = None
                for t in templates:
                    if t.name == cover_template_id:
                        selected_template = t
                        break

                if selected_template:
                    print(f"[COVER-DESIGNER] Plantilla '{cover_template_id}' encontrada (source_type={selected_template.source_type}). Aplicando...")
                    # Preparar datos de portada para plantillas builtin
                    p_title = portada.title if portada else ""
                    p_author = portada.author if portada else ""
                    p_institution = portada.institution if portada else ""
                    p_course = portada.course if portada else ""
                    p_instructor = portada.instructor if portada else ""
                    p_date = portada.date if portada else ""

                    paragraphs_added = apply_cover_to_document(
                        doc, selected_template, base_dir,
                        portada_title=p_title,
                        portada_author=p_author,
                        portada_institution=p_institution,
                        portada_course=p_course,
                        portada_instructor=p_instructor,
                        portada_date=p_date,
                    )
                    cover_template_applied = True
                    print(f"[COVER-DESIGNER]  Portada '{cover_template_id}' aplicada. Parrafos agregados: {paragraphs_added}")
                else:
                    print(f"[COVER-DESIGNER] X Plantilla '{cover_template_id}' NO encontrada en la lista de templates")
                    print(f"[COVER-DESIGNER] Templates disponibles: {[t.name for t in templates]}")
            except Exception as e:
                import traceback
                print(f"[COVER-DESIGNER] X Error aplicando cover designer: {e}")
                traceback.print_exc()

    # 4. Portada sintética (si no existe archivo original o el usuario la solicitó explícitamente)
    # cover_mode: 'keep_original' = 100% intacta, 'keep_design_update_data' = edicion in-place, 'generate_apa7_template' = nueva
    cover_mode = 'keep_original'  # default seguro
    force_skip_cover = False
    if portada:
        force_skip_cover = getattr(portada, 'force_skip_cover', False)
        cover_mode = getattr(portada, 'cover_mode', None) or (
            'keep_original' if getattr(portada, 'use_original_cover', True) else 'generate_apa7_template'
        )
    if not original_file.exists():
        cover_mode = 'generate_apa7_template'

    use_orig_cover = cover_mode != 'generate_apa7_template'

    # Si se aplico una portada via Cover Designer, forzar modo "preservar original"
    # para que el sistema NO genere otra portada sintetica y saltee los parrafos originales
    if cover_template_applied:
        use_orig_cover = True
        force_skip_cover = True

    # Al reemplazar la portada original por una sintética (APA/UNI), la portada
    # original debe ELIMINARSE: de lo contrario sus párrafos quedan DESPUÉS de la
    # nueva portada y se reformatean como cuerpo (el "me destrozó la portada").
    if not use_orig_cover:
        orig_cover_count = 0
        if doc_model.portada:
            raw_idx = doc_model.portada.get('body_start_paragraph_idx')
            if raw_idx is not None:
                try:
                    orig_cover_count = int(raw_idx)
                except (ValueError, TypeError):
                    orig_cover_count = 0
        if orig_cover_count <= 0:
            orig_cover_count = 0
            for item in doc_model.elements:
                elem = ElementModel.model_validate(item) if isinstance(item, dict) else item
                if elem.is_cover_section or elem.type == ElementType.PORTADA_BLOCK:
                    orig_cover_count += 1
                elif elem.type != ElementType.IMAGE and elem.type != ElementType.TABLE:
                    break
        if orig_cover_count > 0:
            cover_paras = doc.paragraphs[:min(orig_cover_count, len(doc.paragraphs))]
            for p in cover_paras:
                el = p._element
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)

    paragraphs_before_body = 0
    if not use_orig_cover and portada:
        is_uni_profile = (cover_mode == 'generate_uni_cover')

        if is_uni_profile:
            try:
                import re as _re

                from modules.portada_uni import generate_uni_cover
                autores_parsed = []
                if portada.author:
                    from modules.portada_normalize import normalize_author_lines
                    portada.author = normalize_author_lines(portada.author)
                    # Formato del frontend: "Br. Nombre Apellido | Carnet: 2023-XXXX"
                    # o dos lineas "Br. Nombre" + "Carnet: 2023-XXXX".
                    author_lines = [l.strip() for l in portada.author.split('\n') if l.strip()]
                    pending_nombre = ""
                    for l in author_lines:
                        l_lower = l.lower()
                        if l_lower.startswith("carnet:") or _re.match(r'^\d{4}-\d+', l):
                            carnet_val = l.split(':', 1)[-1].strip() if ':' in l else l.strip()
                            if pending_nombre:
                                autores_parsed.append({"nombre": pending_nombre, "carnet": carnet_val})
                                pending_nombre = ""
                            elif autores_parsed:
                                autores_parsed[-1]["carnet"] = carnet_val
                            else:
                                autores_parsed.append({"nombre": l, "carnet": ""})
                            continue
                        if '|' in l:
                            parts = [p.strip() for p in l.split('|')]
                            nombre = parts[0]
                            carnet = ""
                            for part in parts[1:]:
                                if part.lower().startswith('carnet:'):
                                    carnet = part.split(':', 1)[-1].strip()
                            autores_parsed.append({"nombre": nombre, "carnet": carnet})
                            continue
                        if pending_nombre:
                            autores_parsed.append({"nombre": pending_nombre, "carnet": ""})
                        pending_nombre = l
                    if pending_nombre:
                        autores_parsed.append({"nombre": pending_nombre, "carnet": ""})
                    # Deduplicar por nombre (el parser de textboxes no debe repetirlos, pero por seguridad)
                    seen_aut = set()
                    seen_carnets = set()
                    uniq_aut = []
                    for a in autores_parsed:
                        key = _re.sub(r'\s+', ' ', a["nombre"]).strip().lower()
                        if key in seen_aut:
                            continue
                        carnet = (a.get("carnet") or "").strip().lower()
                        if carnet and len(uniq_aut) > 0:
                            # Mismo carnet = misma persona aunque el nombre varíe
                            if carnet in seen_carnets:
                                continue
                            seen_carnets.add(carnet)
                        seen_aut.add(key)
                        uniq_aut.append(a)
                    autores_parsed = uniq_aut

                paragraphs_before_body = generate_uni_cover(
                    doc,
                    titulo=portada.title or "Sin Título",
                    asignatura=portada.course or "",
                    autores=autores_parsed,
                    tutor=portada.instructor or "",
                    grupo=getattr(portada, 'grupo', '') or "",
                    fecha=portada.date or "",
                    departamento=getattr(portada, 'departamento', '') or "",
                )
            except Exception as err:
                print(f"[WARN] Error creando portada UNI: {err}")
                paragraphs_before_body = format_apa_portada(doc, portada, rules)
        else:
            paragraphs_before_body = format_apa_portada(doc, portada, rules)
    for table in doc.tables:
        _tbl_style = getattr(rules, 'table_border_style', None)
        set_table_borders(table, _tbl_style.value if _tbl_style else "apa")

    # 6. Mapear y aplicar estilos sobre los párrafos existentes del documento
    elem_map: dict[str, ElementModel] = {}
    for item in doc_model.elements:
        elem = ElementModel.model_validate(item) if isinstance(item, dict) else item
        if elem.id:
            elem_map[elem.id] = elem

    numbered_counters: dict[int, int] = {1: 0, 2: 0, 3: 0}
    last_numbered_level: int = 0
    # Contadores separados para auto-numeración de headings
    heading_counters: dict[int, int] = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    # Contador de ecuaciones (auto-numeración cuando show_number y sin número explícito)
    _equation_counter: int = 0
    # Estilos de numeración por nivel configurables desde APARuleSet
    heading_numbering_styles: dict[int, str] = getattr(rules, 'heading_numbering_styles', {1: 'decimal', 2: 'decimal', 3: 'decimal', 4: 'decimal', 5: 'decimal'})

    # Iterar párrafos existentes o crear nuevos si no existen suficientes
    existing_paragraphs = list(doc.paragraphs)

    # ─── Calcular cover_paragraph_count de forma robusta ──────────────────────
    cover_paragraph_count = 0
    if use_orig_cover:
        # Señal 0 (PRIMARIA): Usar body_start_paragraph_idx del analisis XML
        # Este es el valor mas preciso porque se calcula directamente del XML raw
        # y no pasa por la interpretacion de python-docx
        raw_body_idx = None
        if doc_model.portada:
            raw_body_idx = doc_model.portada.get('body_start_paragraph_idx')
            if raw_body_idx is not None:
                try:
                    raw_body_idx = int(raw_body_idx)
                except (ValueError, TypeError):
                    raw_body_idx = None

        if raw_body_idx is not None and raw_body_idx > 0:
            cover_paragraph_count = raw_body_idx
        else:
            # Señal 1: usar is_cover_section de los elementos del modelo (desde pre_classifier)
            cover_element_count = 0
            for item in doc_model.elements:
                elem = ElementModel.model_validate(item) if isinstance(item, dict) else item
                if elem.is_cover_section or elem.type == ElementType.PORTADA_BLOCK:
                    cover_element_count += 1
                elif elem.type != ElementType.IMAGE and elem.type != ElementType.TABLE:
                    break

            # Si force_skip_cover esta activo, confiar exclusivamente en cover_element_count
            if force_skip_cover:
                cover_paragraph_count = max(1, cover_element_count) if cover_element_count > 0 else 0
            else:
                # Señal 2: hacer coincidir textos del primer cuerpo con parrafos reales del docx
                first_body_texts = []
                for item in doc_model.elements:
                    elem = ElementModel.model_validate(item) if isinstance(item, dict) else item
                    if not (elem.is_cover_section or elem.type == ElementType.PORTADA_BLOCK):
                        if elem.text and elem.text.strip():
                            first_body_texts.append(elem.text.strip()[:80].lower())
                            if len(first_body_texts) >= 3:
                                break

                found_body_start = False
                if first_body_texts:
                    for para_idx, p in enumerate(existing_paragraphs):
                        p_text = p.text.strip()[:80].lower()
                        if p_text:
                            for fb_text in first_body_texts:
                                if fb_text and (p_text.startswith(fb_text[:40]) or fb_text.startswith(p_text[:40])):
                                    cover_paragraph_count = para_idx
                                    found_body_start = True
                                    break
                            if found_body_start:
                                break

                if not found_body_start and cover_element_count > 0:
                    cover_paragraph_count = max(1, cover_element_count)

                if not found_body_start and cover_element_count == 0:
                    img_count_at_start = 0
                    for para_idx, p in enumerate(existing_paragraphs[:10]):
                        if (p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                                or p._element.findall('.//{urn:schemas-microsoft-com:vml}imagedata')):
                            img_count_at_start = para_idx + 1
                    cover_paragraph_count = img_count_at_start
    else:
        # Al reemplazar portada con la plantilla UNI, los parrafos de cuerpo empiezan despues de la nueva portada
        cover_paragraph_count = paragraphs_before_body

    # SAFETY NET: Si cover_paragraph_count sigue siendo 0 pero el modelo
    # tiene elementos marcados como is_cover_section, usar ese count.
    # Esto previene que la portada original sea reformateada como cuerpo.
    if cover_paragraph_count == 0 and use_orig_cover:
        cover_from_model = sum(
            1 for item in doc_model.elements
            if (ElementModel.model_validate(item) if isinstance(item, dict) else item).is_cover_section
        )
        if cover_from_model > 0:
            cover_paragraph_count = cover_from_model
            print(f"[COVER-SAFETY] cover_paragraph_count ajustado desde modelo: {cover_paragraph_count}")

    p_idx = cover_paragraph_count

    existing_tables = list(doc.tables)
    table_count_processed = 0

    # Construir fingerprint de tablas existentes (hash de primera celda) para matching robusto
    table_fingerprints: list[tuple[int, str]] = []
    for tbl_idx, tbl in enumerate(existing_tables):
        first_cell_text = ""
        try:
            if tbl.rows and tbl.rows[0].cells:
                first_cell_text = tbl.rows[0].cells[0].text.strip()[:100].lower()
        except Exception:
            pass
        table_fingerprints.append((tbl_idx, first_cell_text))
    used_table_indices: set[int] = set()

    # Coleccionar imágenes existentes en el documento para formateo in-place
    existing_drawings = []
    for p in doc.paragraphs:
        if p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip') or p._element.findall('.//{urn:schemas-microsoft-com:vml}imagedata'):
            existing_drawings.append(p)
    image_count_processed = 0

    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    in_cover = cover_paragraph_count > 0
    last_was_heading1 = False

    def insert_page_break_before(oxml_element):
        """Insert a page break paragraph before the given element (robust version)."""
        parent = oxml_element.getparent()
        if parent is None:
            return
        try:
            parent_children = list(parent)
            idx = parent_children.index(oxml_element)
            new_p_xml = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>')
            parent.insert(idx, new_p_xml)
        except (ValueError, AttributeError):
            new_p_xml = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>')
            oxml_element.addprevious(new_p_xml)

    def _clone_current_section_settings(docx_doc):
        """Clona la configuracion de la seccion actual (pgSz, pgMar, cols, docGrid)."""
        try:
            from copy import deepcopy
            body = docx_doc._element.body
            sect_prs = body.findall(qn('w:sectPr'))
            if sect_prs:
                return deepcopy(sect_prs[-1])
        except Exception:
            pass
        return None

    def insert_section_break_before(oxml_element, docx_doc):
        """Insert a section break paragraph before the given element (robust version).

        En OOXML, un `w:sectPr` dentro del `w:pPr` de un párrafo cierra la sección
        actual e inicia una nueva con la misma configuración (equivalente a un salto
        de sección "página siguiente" con herencia de página/márgenes).
        """
        new_p = None
        try:
            sect_pr = _clone_current_section_settings(docx_doc)
            if sect_pr is None:
                sect_pr = OxmlElement('w:sectPr')
            new_p = OxmlElement('w:p')
            p_pr = OxmlElement('w:pPr')
            p_pr.append(sect_pr)
            new_p.append(p_pr)
            parent = oxml_element.getparent()
            if parent is None:
                return
            parent_children = list(parent)
            idx = parent_children.index(oxml_element)
            parent.insert(idx, new_p)
        except (ValueError, AttributeError):
            try:
                if new_p is not None:
                    oxml_element.addprevious(new_p)
            except Exception:
                pass
        except Exception:
            pass

    def _has_force_page_break_marker(text: str) -> bool:
        """Verifica si el texto contiene el marcador [FORCE_PAGE_BREAK]."""
        return text and '[FORCE_PAGE_BREAK]' in text

    def _extract_numbering_style_marker(text: str) -> str:
        """Extrae el marcador de estilo de numeración: [ROMAN] o [DECIMAL]. Retorna 'decimal' por defecto."""
        if not text:
            return 'decimal'
        upper = text.upper()
        if '[ROMAN]' in upper:
            return 'roman'
        return 'decimal'

    def _strip_markers(text: str) -> str:
        """Elimina marcadores internos como [FORCE_PAGE_BREAK], [ROMAN], [DECIMAL] del texto."""
        if not text:
            return text or ""
        import re
        result = re.sub(r'\s*\[FORCE_PAGE_BREAK\]\s*', '', text)
        result = re.sub(r'\s*\[ROMAN\]\s*', '', result, flags=re.IGNORECASE)
        result = re.sub(r'\s*\[DECIMAL\]\s*', '', result, flags=re.IGNORECASE)
        return result.strip()

    def _should_insert_page_break(elem_type, heading_level: int, text: str, original_text: str = "") -> bool:
        """
        Decide si se debe insertar page break antes de este elemento.
        - Heading 1: SIEMPRE page break (Bug 1 fix: todo H1 en pagina nueva)
        - Elementos con [FORCE_PAGE_BREAK]: SIEMPRE page break
        """
        if _has_force_page_break_marker(original_text or text):
            return True
        if elem_type == ElementType.HEADING and heading_level == 1:
            return True
        return False

    _landscape_threshold_cm = 16.0
    try:
        _usable_cm = _usable_width_cm(doc.sections[0] if doc.sections else None)
        if _usable_cm:
            _landscape_threshold_cm = max(10.0, _usable_cm - 0.5)
    except Exception:
        pass

    for item in doc_model.elements:
        elem = ElementModel.model_validate(item) if isinstance(item, dict) else item
        elem_type = elem.type
        elem_level: int = elem.heading_level if elem.heading_level else 1
        list_lvl: int = elem.list_level if elem.list_level else 1

        if elem_type == ElementType.EMPTY:
            continue

        # No formatear párrafos pertenecientes a la sección de portada (pero avanzar contadores de elementos de portada)
        if elem.is_cover_section or elem_type == ElementType.PORTADA_BLOCK:
            if elem_type == ElementType.IMAGE and image_count_processed < len(existing_drawings):
                image_count_processed += 1
            elif elem_type == ElementType.TABLE:
                # Avanzar contador de tabla usando fingerprint matching si es cover
                if table_count_processed < len(existing_tables):
                    curr_tbl = existing_tables[table_count_processed]
                    table_count_processed += 1
                    used_table_indices.add(table_count_processed - 1)
            continue

        # ── CASO ESPECIAL: TABLA EXISTENTE (FORMATO ATÓMICO IN-PLACE + ETIQUETA PEGUERA) ──
        if elem_type == ElementType.TABLE:
            # 🆕 LANDSCAPE: Check if table needs landscape orientation
            table_needs_landscape = False
            if elem.table_info:
                too_many, too_wide = _is_table_too_wide(elem.table_info)
                table_needs_landscape = too_many or too_wide
            # Intentar matching por contenido (fingerprint de primera celda) en vez de solo índice
            matched_tbl_idx = -1
            elem_first_cell = ""
            if elem.table_info and elem.table_info.rows:
                elem_first_cell = elem.table_info.rows[0][0].strip()[:100].lower() if elem.table_info.rows[0] else ""

            if elem_first_cell:
                for fp_idx, (orig_idx, fp_text) in enumerate(table_fingerprints):
                    if orig_idx not in used_table_indices and fp_text and fp_text == elem_first_cell:
                        matched_tbl_idx = orig_idx
                        break

            if matched_tbl_idx < 0 and table_count_processed < len(existing_tables):
                matched_tbl_idx = table_count_processed

            if matched_tbl_idx >= 0 and matched_tbl_idx < len(existing_tables):
                curr_tbl = existing_tables[matched_tbl_idx]
                table_count_processed += 1
                used_table_indices.add(matched_tbl_idx)
                _tbl_style = getattr(rules, 'table_border_style', None)
                set_table_borders(curr_tbl, _tbl_style.value if _tbl_style else "apa")
                fit_table_to_page(curr_tbl, rules, landscape=table_needs_landscape)

                tbl_num = elem.table_info.table_number if (elem.table_info and elem.table_info.table_number > 0) else table_count_processed
                caption_text = elem.table_info.caption if elem.table_info else ""

                if in_cover:
                    in_cover = False
                    if cover_paragraph_count > 0:
                        insert_page_break_before(curr_tbl._tbl)

                # 🆕 LANDSCAPE: Wrap wide table in landscape section
                if table_needs_landscape:
                    _wrap_in_landscape_section(doc, curr_tbl._tbl)

                # Insertar un nuevo párrafo dedicado inmediatamente anterior al XML de la tabla para la etiqueta APA 7
                new_p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                curr_tbl._tbl.addprevious(new_p_xml)
                p_label = docx.text.paragraph.Paragraph(new_p_xml, doc)

                p_label.text = ""
                p_label.paragraph_format.space_before = Pt(6)
                p_label.paragraph_format.space_after = Pt(2)
                p_label.paragraph_format.line_spacing = rules.line_spacing
                p_label.paragraph_format.first_line_indent = Inches(0)
                p_label.paragraph_format.keep_with_next = True

                r_num = p_label.add_run(f"{rules.table_label_prefix} {tbl_num}\n")
                r_num.bold = True
                r_num.font.name = rules.font_family
                r_num.font.size = Pt(rules.font_size_pt)
                r_num.font.color.rgb = RGBColor(0, 0, 0)

                if caption_text:
                    r_cap = p_label.add_run(caption_text)
                    r_cap.italic = True
                    r_cap.font.name = rules.font_family
                    r_cap.font.size = Pt(rules.font_size_pt)
                    r_cap.font.color.rgb = RGBColor(0, 0, 0)
            else:
                if elem.table_info:
                    format_apa_table(doc, elem.table_info, rules)
            continue

        # ── CASO ESPECIAL: IMAGEN EXISTENTE (FORMATO ATÓMICO IN-PLACE + SIN DUPLICAR) ──
        elif elem_type == ElementType.IMAGE:
            if image_count_processed < len(existing_drawings):
                img_p = existing_drawings[image_count_processed]
                image_count_processed += 1

                fig_num = elem.image_info.figure_number if (elem.image_info and elem.image_info.figure_number > 0) else image_count_processed
                caption_text = elem.image_info.caption if elem.image_info else ""

                # 🆕 DESIGN STYLE: Apply design_style formatting
                _apply_image_design_style(img_p, elem, rules)

                if in_cover:
                    in_cover = False
                    if cover_paragraph_count > 0:
                        insert_page_break_before(img_p._element)

                # 🆕 LANDSCAPE: Wrap wide images in landscape section
                if elem.image_info and (elem.image_info.width_cm or 0) > _landscape_threshold_cm:
                    _wrap_in_landscape_section(doc, img_p._element)

                # 🆕 BUG 5 FIX: Detectar si ya existe un caption de figura en el texto
                # del parrafo de la imagen o en el parrafo inmediatamente anterior
                existing_caption = _detect_existing_figure_caption(img_p, existing_paragraphs)
                if existing_caption:
                    # Ya existe caption — no duplicar. Si el usuario escribio un
                    # caption nuevo en el UI, actualizar el texto existente en vez
                    # de crear otro parrafo.
                    if caption_text and caption_text != existing_caption:
                        _update_existing_caption_text(img_p, caption_text, rules)
                else:
                    # No hay caption previo — crear parrafo de etiqueta APA 7
                    new_p_xml = parse_xml(f'<w:p {nsdecls("w")}/>')
                    img_p._element.addprevious(new_p_xml)
                    p_label = docx.text.paragraph.Paragraph(new_p_xml, doc)

                    p_label.text = ""
                    p_label.paragraph_format.space_before = Pt(12)
                    p_label.paragraph_format.space_after = Pt(2)
                    p_label.paragraph_format.line_spacing = rules.line_spacing
                    p_label.paragraph_format.first_line_indent = Inches(0)
                    p_label.paragraph_format.keep_with_next = True

                    r_num = p_label.add_run(f"{rules.figure_label_prefix} {fig_num}\n")
                    r_num.bold = True
                    r_num.font.name = rules.font_family
                    r_num.font.size = Pt(rules.font_size_pt)
                    r_num.font.color.rgb = RGBColor(0, 0, 0)

                    if caption_text:
                        r_cap = p_label.add_run(caption_text)
                        r_cap.italic = True
                        r_cap.font.name = rules.font_family
                        r_cap.font.size = Pt(rules.font_size_pt)
                        r_cap.font.color.rgb = RGBColor(0, 0, 0)
            else:
                if elem.image_info:
                    format_apa_figure(doc, elem.image_info, rules)
            continue

        # Obtener o crear párrafo objetivo para texto (orden secuencial = orden del modelo)
        if p_idx < len(existing_paragraphs):
            p = existing_paragraphs[p_idx]
            p_idx += 1
        else:
            p = doc.add_paragraph()

        #  PAGE BREAK LOGIC (BUG 1 FIX)
        needs_break = False
        force_text = (elem.original_text or elem.text or "")

        # Cover-to-body transition: page break after cover (only if cover detected)
        if in_cover:
            in_cover = False
            if cover_paragraph_count > 0:
                needs_break = True

        # Heading 1 or force marker: every H1 starts a new page
        # Bug fix: no saltar si es un H1 inmediatamente despues de otro H1
        is_h1 = (elem_type == ElementType.HEADING and elem_level == 1)
        if _should_insert_page_break(elem_type, elem_level, elem.text or "", force_text):
            if not last_was_heading1:
                needs_break = True

        # Actualizar estado de H1 consecutivo
        if is_h1:
            last_was_heading1 = True
        elif elem.text and str(elem.text).strip():
            last_was_heading1 = False

        # Insert page break if needed (avoid double break)
        if needs_break:
            insert_page_break_before(p._element)

        # Limpiar marcadores inline de notas al pie del texto (el contenido
        # real se anexa al final del párrafo tras formatearlo).
        if elem.footnote_ids and elem.text:
            elem.text = _strip_inline_footnote_markers(elem.text)

        if elem_type == ElementType.HEADING:
            lvl: int = elem.heading_level if elem.heading_level else 1

            # Reiniciar contadores de listas numeradas al iniciar nueva sección
            if lvl <= 2:
                numbered_counters = {1: 0, 2: 0, 3: 0}
                last_numbered_level = 0

            # Auto-numeración persistente de headings
            heading_counters[lvl] = heading_counters.get(lvl, 0) + 1
            # Resetear contadores de niveles inferiores
            for lower_lvl in range(lvl + 1, 6):
                heading_counters[lower_lvl] = 0

            # Detectar estilo de numeracion especifico para este nivel
            level_style = getattr(rules, f'heading_numbering_style_lvl{lvl}', 'decimal')

            orig_text = elem.original_text or elem.text or ""
            if lvl == 1:
                marker_style = _extract_numbering_style_marker(orig_text)
                detected_style = _detect_heading_numbering_style(orig_text)
                if marker_style:
                    level_style = marker_style
                elif detected_style == 'roman':
                    level_style = 'roman'

            # Construir prefijo numerico
            number_prefix = _build_heading_prefix(heading_counters, lvl, level_style)

            # Strip existing numbering from text before adding programmatic prefix
            raw_text = elem.text or p.text
            if level_style == 'roman' and lvl == 1:
                raw_text = _strip_existing_numbering(raw_text)

            # Limpiar marcadores del texto y construir heading final
            clean_text = _strip_markers(raw_text)
            heading_text = number_prefix + clean_text
            format_heading_paragraph(p, lvl, heading_text, rules, preserve_text=(elem.has_math or elem.has_fields))

        elif elem_type == ElementType.PARAGRAPH:
            # A1: párrafos de la portada original preservada NO se formatean.
            if use_orig_cover and elem.is_cover_section:
                continue
            # Reset numbered list counters — a paragraph breaks the list sequence
            numbered_counters = {1: 0, 2: 0, 3: 0}
            last_numbered_level = 0
            format_normal_paragraph(p, elem.text or p.text, rules, preserve_text=(elem.has_math or elem.has_fields))

        elif elem_type == ElementType.PORTADA_BLOCK:
            # A1 Portada INTOCABLE: si se preserva la portada original, ni
            # formato ni reescritura de texto — representación tal cual.
            if use_orig_cover:
                pass
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if elem.alignment == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = rules.line_spacing
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                if elem.text:
                    p.text = ""
                    r = p.add_run(elem.text)
                    set_run_font(r, rules.font_family, rules.font_size_pt)
                    r.bold = elem.is_bold
                    r.italic = elem.is_italic
                    r.font.color.rgb = RGBColor(0, 0, 0)

        elif elem_type == ElementType.BLOCK_QUOTE:
            numbered_counters = {1: 0, 2: 0, 3: 0}
            last_numbered_level = 0
            format_block_quote(p, elem.text or p.text, rules)

        elif elem_type == ElementType.BULLET:
            numbered_counters = {1: 0, 2: 0, 3: 0}
            last_numbered_level = 0
            _bullet_style = rules.bullet_style_level1 if list_lvl == 1 else (rules.bullet_style_level2 if list_lvl == 2 else rules.bullet_style_level3)
            format_bullet_item(p, elem.text or p.text, level=list_lvl, rules=rules, bullet_style=_bullet_style, is_bold=elem.is_bold, is_italic=elem.is_italic)

        elif elem_type == ElementType.NUMBERED_LIST:
            num_lvl = list_lvl if list_lvl in (1, 2, 3) else 1
            if num_lvl <= last_numbered_level:
                for l in range(num_lvl + 1, 4):
                    numbered_counters[l] = 0

            numbered_counters[num_lvl] = numbered_counters.get(num_lvl, 0) + 1
            last_numbered_level = num_lvl
            item_index = numbered_counters[num_lvl]
            _num_style = rules.number_style_level1 if num_lvl == 1 else (rules.number_style_level2 if num_lvl == 2 else rules.number_style_level3)
            format_numbered_item(p, elem.text or p.text, level=num_lvl, index=item_index, rules=rules, num_style=_num_style, is_bold=elem.is_bold, is_italic=elem.is_italic)

        elif elem_type == ElementType.TOC:
            # 🆕 GENERAR ÍNDICE / TABLA DE CONTENIDOS basado en headings detectados
            _generate_toc_from_headings(doc, doc_model, rules, p)

        elif elem_type == ElementType.EQUATION:
            # Ecuaciones OMML: preservar el XML intacto (no modificar ni reformatear).
            # Las ecuaciones de Word usan Office Math Markup Language (m:oMath)
            # que es fragil y no debe ser tocado por el motor de estilos.
            eq_cfg = getattr(elem, 'equation', None)
            # Config por defecto si no se inicializo
            if eq_cfg is None:
                from models import EquationConfig
                eq_cfg = EquationConfig()

            # Alineación configurable (centro por defecto, segun APA/estilo tesis)
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            p.paragraph_format.alignment = align_map.get(
                (eq_cfg.alignment or "center"), WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = rules.line_spacing
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Numeración de ecuación: tab stop derecho → el número pegado al
            # margen derecho mientras la ecuación queda alineada segun `alignment`.
            # Si la ecuación está centrada, el tab derecho empuja el número al borde.
            if eq_cfg.show_number:
                _equation_counter += 1
                eq_number = eq_cfg.number or str(_equation_counter)
                display = _render_equation_number(eq_number, eq_cfg.number_format or "(1)")
                _apply_equation_number(p, display, eq_cfg)

            # Fuente de apoyo: aplicar a los runs de texto plano (NO al XML OMML).
            # set_run_font solo toca w:rPr de runs regulares; los m:oMath quedan intactos.
            try:
                for r in p.runs:
                    set_run_font(r, eq_cfg.font_name or rules.font_family, eq_cfg.font_size_pt or rules.font_size_pt)
            except Exception:
                pass

        elif elem_type == ElementType.PAGE_BREAK:
            if p is not None and p._element is not None:
                insert_page_break_before(p._element)

        elif elem_type == ElementType.SECTION_BREAK:
            # Salto de sección real (no solo de página): insertar un párrafo con
            # w:sectPr que hereda la configuración de la sección actual.
            try:
                if p is not None and p._element is not None:
                    insert_section_break_before(p._element, doc)
            except Exception:
                print("[WARN] No se pudo insertar salto de sección para el elemento "
                      f"'{elem.id}'")

        # ── POST-PROCESO: rehidratar hyperlinks, bookmarks y notas al pie ──────
        if (
            getattr(elem, 'hyperlinks', None)
            or getattr(elem, 'bookmarks', None)
            or getattr(elem, 'footnote_ids', None)
        ):
            try:
                if getattr(elem, 'hyperlinks', None):
                    rehydrate_hyperlinks(p, elem, doc)
                if getattr(elem, 'bookmarks', None):
                    rehydrate_bookmarks(p, elem)
                if getattr(elem, 'footnote_ids', None):
                    footnotes = doc_model.meta.footnotes if doc_model.meta else []
                    rehydrate_footnotes(p, elem, footnotes)
            except Exception:
                print(f"[WARN] Post-proceso (hyperlinks/bookmarks/footnotes) falló en '{elem.id}'")

    # 7. Sección de Referencias Bibliográficas
    if references:
        format_apa_referencias_section(doc, references, rules)

    from generation.style_engine import normalize_global_body_spacing
    normalize_global_body_spacing(doc, rules, cover_paragraph_count)

    # 7.5 Normalización dura de fuentes en TODOS los runs (cuerpo, celdas,
    # headers/footers) para eliminar mezclas heredadas de Word/internet.
    from generation.style_engine import normalize_all_fonts, normalize_cover_font_name
    # Primero normalizar solo el nombre de fuente de la portada (respetando
    # tamaños originales para mantener la jerarquía visual del diseño).
    if cover_paragraph_count > 0 and not remove_cover_paragraphs:
        normalize_cover_font_name(doc, rules.font_family, cover_paragraph_count)
    normalize_all_fonts(
        doc,
        rules,
        target_font=rules.font_family,
        target_size_pt=int(rules.font_size_pt),
        skip_body_paragraphs=cover_paragraph_count,
    )

    # 8.5 Remover párrafos de portada si el post-procesador COM se va a encargar
    if remove_cover_paragraphs and cover_paragraph_count > 0:
        for _ in range(cover_paragraph_count):
            if doc.paragraphs:
                p = doc.paragraphs[0]
                p._element.getparent().remove(p._element)

    # 9. Validar ancho de tablas antes de guardar
    width_warnings = validate_table_widths(doc, rules)
    if width_warnings:
        for warning in width_warnings:
            print(f"[WARN] {warning}")
        if getattr(doc_model.meta, 'content_warning', None):
            doc_model.meta.content_warning = f"{doc_model.meta.content_warning} | {'; '.join(width_warnings)}"
        else:
            doc_model.meta.content_warning = '; '.join(width_warnings)

    # 10. Guardar resultado final
    doc.save(out_path)

    # 11. Activar updateFields en settings.xml para que Word recalcule TOC/PAGEREF automáticamente
    try:
        from parsing.field_guard import enable_word_update_fields
        enable_word_update_fields(out_path)
    except Exception as e:
        print(f"[WARN] Error activando updateFields: {e}")

    # 12. Gate de Sanidad: verificar que no haya pérdida silenciosa de texto
    if original_file.exists():
        try:
            from parsing.sanity_check import verify_document_content_integrity
            verify_document_content_integrity(original_file, out_path, tolerance=0.05)
        except Exception as e:
            print(f"[WARN] Alerta de Sanidad en exportación: {e}")
            raise

    return out_path
