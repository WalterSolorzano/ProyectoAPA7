"""
WordAPA7 — Estructura de Documento y Encabezados (Running Head)

Maneja los elementos de nivel de documento:
- Numeración de páginas en la esquina superior derecha (mediante fldChar PAGE en XML de Word).
- Running Head (Título corto en mayúsculas a la izquierda en modo Profesional).
- Salto de página previo a secciones mayores (Referencias, Portada).
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor
from models import APAFormat, APARuleSet


def add_page_number_field(run):
    """
    Inserta un campo XML de Word `PAGE` dentro del run para numeración dinámica.
    """
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _clean_existing_page_numbers(doc: docx.Document):
    """
    Elimina cualquier campo PAGE / NUMPAGES existente en los headers y footers
    de todas las secciones del documento.  Esto evita que el numero de pagina
    APA se superponga a uno ya presente en el documento original (doble numero).
    Es seguro llamarlo aunque el documento no tenga headers (no-op).
    """
    for section in doc.sections:
        for attr in ('header', 'first_page_header', 'even_page_header',
                     'footer', 'first_page_footer', 'even_page_footer'):
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            try:
                element = hf._element
            except (AttributeError, Exception):
                continue

            for p in element.findall(f'.//{qn("w:p")}'):
                should_clean = False
                for instr in p.findall(f'.//{qn("w:instrText")}'):
                    text = (instr.text or '').upper()
                    if 'PAGE' in text or 'NUMPAGES' in text:
                        should_clean = True
                        break
                if should_clean:
                    # Elimina todos los runs del parrafo que contienen el field
                    for r in list(p.findall(qn("w:r"))):
                        p.remove(r)


_TOC_LEADER_RE = None  # lazy


def refresh_or_flag_existing_toc(doc: docx.Document) -> str | None:
    """H20: si el documento ya trae índice, NO duplicarlo.

    - TOC nativo de Word (SDT con campo TOC): se deja intacto; Word lo
      actualiza solo al imprimir/F9. Retorna 'sdt'.
    - Índice "manual" (líneas con puntos de relleno + número): se REEMPLAZA
      por un campo TOC real para que Word lo mantenga. Retorna 'legacy'.
    - Sin índice: retorna None (el generador puede añadir uno si procede).
    """
    global _TOC_LEADER_RE
    if _TOC_LEADER_RE is None:
        import re as _re
        _TOC_LEADER_RE = _re.compile(r"^(.{3,120}?)[ \t]*\.{3,}[ \t]*(\d{1,4})\s*$")

    body = doc.element.body

    # 1) SDT nativo
    for sdt in body.iter(qn("w:sdt")):
        for instr in sdt.iter(qn("w:instrText")):
            if "TOC" in (instr.text or ""):
                return "sdt"

    # 2) Índice manual: bloque >=2 párrafos con leader-dots + número.
    paras = doc.paragraphs
    hits: list[int] = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t and _TOC_LEADER_RE.match(t):
            hits.append(i)
    if len(hits) < 2:
        return None

    # Contigüidad aproximada (tolera título "Índice"/blancos intercalados).
    first, last = hits[0], hits[-1]
    if last - first > len(hits) * 3:
        return None

    start_p = paras[first]._p
    end_p = paras[last]._p

    new_p = parse_xml(
        '<w:p %s><w:fldSimple w:instr=" TOC \\\\o &quot;1-3&quot; \\\\h \\\\z \\\\u ">'
        '<w:r><w:t>Índice</w:t></w:r></w:fldSimple></w:p>' % nsdecls("w")
    )
    start_p.addprevious(new_p)

    cur = start_p
    while cur is not end_p:
        nxt = cur.getnext()
        body.remove(cur)
        cur = nxt if nxt is not None else end_p
    body.remove(end_p)
    return "legacy"


def setup_apa_header(doc: docx.Document, format_type: APAFormat, running_head_text: str, rules: APARuleSet):
    """
    Configura el encabezado de las páginas según el tipo APA (Estudiante vs Profesional).
    - Portada e Índice: No llevan número de página visible.
    - Cuerpo (Sección principal): Número de página a la derecha (correlativo, ej. 3 si va tras portada e índice).
    - Modo Profesional: Running head en mayúsculas a la izquierda + Número de página a la derecha.
    """
    # Clean existing page number fields to prevent double numbering
    _clean_existing_page_numbers(doc)

    if len(doc.sections) > 1:
        # Sección 1 (Front matter: Portada / Índice) -> Sin número de página
        sec_front = doc.sections[0]
        sec_front.different_first_page_header_footer = True
        try:
            for p in sec_front.header.paragraphs:
                p.text = ""
            for p in sec_front.first_page_header.paragraphs:
                p.text = ""
        except Exception:
            pass

        # Secciones subsiguientes (Cuerpo del documento)
        for i, sec in enumerate(doc.sections[1:], start=2):
            sec.header.is_linked_to_previous = False
            p = sec.header.paragraphs[0]
            p.text = ""
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if format_type == APAFormat.PROFESSIONAL and running_head_text:
                r_rh = p.add_run(f"{running_head_text.upper()}\t")
                r_rh.font.name = rules.font_family
                r_rh.font.size = Pt(rules.font_size_pt)
                r_rh.font.color.rgb = RGBColor(0, 0, 0)

            r_pg = p.add_run()
            r_pg.font.name = rules.font_family
            r_pg.font.size = Pt(rules.font_size_pt)
            r_pg.font.color.rgb = RGBColor(0, 0, 0)
            add_page_number_field(r_pg)

    else:
        # Documento de sección única
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        header = section.header
        header.is_linked_to_previous = False

        p = header.paragraphs[0]
        p.text = ""
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if format_type == APAFormat.PROFESSIONAL and running_head_text:
            # Running Head a la izquierda + Tabulación a la derecha + Número de página
            r_rh = p.add_run(f"{running_head_text.upper()}\t")
            r_rh.font.name = rules.font_family
            r_rh.font.size = Pt(rules.font_size_pt)
            r_rh.font.color.rgb = RGBColor(0, 0, 0)

        # Run para el número de página
        r_pg = p.add_run()
        r_pg.font.name = rules.font_family
        r_pg.font.size = Pt(rules.font_size_pt)
        r_pg.font.color.rgb = RGBColor(0, 0, 0)
        add_page_number_field(r_pg)
