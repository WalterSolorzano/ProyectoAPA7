"""
WordAPA7 — Motor de Tablas APA 7

Aplica el formato estricto de tablas APA 7:
1. Bordes horizontales unicamente en la parte superior, inferior y debajo del encabezado.
2. Cero bordes verticales (left, right, insideV eliminados).
3. Numero de tabla ("Tabla 1") en negrita arriba.
4. Titulo de la tabla en cursiva en la siguiente linea.
5. Nota de tabla al pie si existe.
"""

from typing import Optional

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from models import TableModel, APARuleSet


def _tc_element_of(cell_or_tc):
    """Normaliza un objeto python-docx Cell o un `w:tc` lxml crudo al elemento `w:tc`."""
    return cell_or_tc._tc if hasattr(cell_or_tc, '_tc') else cell_or_tc


def _safe_cell_text(cell) -> str:
    """Lee el texto visible de una celda de forma defensiva.

    Acepta una python-docx Cell o un `w:tc` crudo. Solo recopila los runs `w:t`
    y NO toca el XML interno (las tablas anidadas `<w:tbl>` dentro de la celda
    se conservan). Nunca lanza excepciones ni muta la celda.
    """
    try:
        tc = _tc_element_of(cell)
        return "".join(t.text or "" for t in tc.iter(qn('w:t')))
    except Exception:
        return ""


def _cell_has_nested_table(cell) -> bool:
    """True si la celda contiene una tabla anidada (`<w:tbl>`) en su XML."""
    try:
        tc = _tc_element_of(cell)
        return tc.find(qn('w:tbl')) is not None
    except Exception:
        return False


def _is_vmerge_continuation(cell) -> bool:
    """True si la celda es la continuación de un merge vertical (w:vMerge sin val="restart").

    En OOXML, `<w:vMerge/>` (sin val) o `val="continue"` marca la continuación;
    solo `val="restart"` es la celda superior del merge. Las celdas de
    continuación no deben recibir tratamiento de bordes que rompa el layout.
    """
    try:
        tc = _tc_element_of(cell)
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            return False
        vmerge = tcPr.find(qn('w:vMerge'))
        if vmerge is None:
            return False
        val = vmerge.get(qn('w:val'))
        if val is None or str(val).lower() in ('continue', 'cont'):
            return True
    except Exception:
        pass
    return False


def _force_run_bold(r_el) -> None:
    """Marca un `w:r` como negrita sin depender del objeto Run."""
    try:
        rPr = r_el.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r_el.insert(0, rPr)
        if rPr.find(qn('w:b')) is None:
            rPr.append(OxmlElement('w:b'))
    except Exception:
        pass


def _raw_row_tcs(row) -> list:
    """Devuelve los elementos `w:tc` crudos de una fila (incluye celdas de merge)."""
    try:
        return list(row._tr.findall(qn('w:tc')))
    except Exception:
        return []


def _available_page_width_inches(rules: APARuleSet, landscape: bool = False) -> float:
    page_width = 11.0 if landscape else 8.5
    margin_inches = rules.margins_cm / 2.54
    return max(3.0, page_width - (margin_inches * 2))


def fit_table_to_page(table, rules: APARuleSet, landscape: bool = False) -> None:
    """Force a table to stay within the available text width."""
    try:
        from docx.shared import Inches
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    available_width = _available_page_width_inches(rules, landscape=landscape)
    column_count = len(table.columns)
    if column_count <= 0:
        return

    table.autofit = False
    target_width = Inches(available_width)

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tbl_pr)
    tbl_width = tbl_pr.find(qn('w:tblW'))
    if tbl_width is None:
        tbl_width = OxmlElement('w:tblW')
        tbl_pr.append(tbl_width)
    tbl_width.set(qn('w:type'), 'dxa')
    tbl_width.set(qn('w:w'), str(int(target_width.twips)))

    widths = []
    try:
        existing_widths = [col.width.inches for col in table.columns if col.width]
        total_existing = sum(existing_widths)
        if total_existing > 0:
            scale = available_width / total_existing
            widths = [max(0.75, w * scale) for w in existing_widths]
    except Exception:
        widths = []

    if not widths:
        widths = [available_width / column_count for _ in range(column_count)]

    for col_idx, width_in in enumerate(widths):
        try:
            table.columns[col_idx].width = Inches(width_in)
        except Exception:
            pass
        for row in table.rows:
            try:
                row.cells[col_idx].width = Inches(width_in)
            except Exception:
                pass


def validate_table_widths(doc: docx.Document, rules: APARuleSet) -> list[str]:
    """Log any tables that still exceed the printable width after formatting."""
    warnings: list[str] = []
    max_width = _available_page_width_inches(rules, landscape=False)
    for idx, table in enumerate(doc.tables, start=1):
        total = 0.0
        try:
            widths = [col.width.inches for col in table.columns if col.width]
            total = sum(widths)
        except Exception:
            total = 0.0
        if total and total > (max_width + 0.05):
            warnings.append(f"Tabla {idx} supera el ancho utilizable ({total:.2f}in > {max_width:.2f}in).")
    return warnings


def set_table_borders(table, style: str = "apa") -> None:
    """
    Modifica el XML de la tabla para establecer bordes según el estilo del perfil:
    - "apa": APA 7 — borde superior e inferior simples (0.75 pt), línea bajo el
      encabezado, cero bordes verticales e internos.
    - "grid": cuadrícula completa — todos los bordes simples (0.5 pt), típico
      de revistas científicas y manuales.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Reemplazar tblBorders existente
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    if style == "grid":
        borders_xml = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders_xml)

        # Evitar division de filas entre paginas (w:cantSplit)
        for row in table.rows:
            try:
                trPr = row._tr.get_or_add_trPr()
                trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            except Exception:
                pass

        # Marcar primera fila como encabezado repetible (w:tblHeader)
        if len(table.rows) > 0:
            try:
                hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
                hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            except Exception:
                pass
        return

    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_xml)

    # Evitar division de filas entre paginas (w:cantSplit)
    for row in table.rows:
        try:
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        except Exception:
            pass

    # Marcar primera fila como encabezado repetible (w:tblHeader)
    if len(table.rows) > 0:
        try:
            hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
            hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        except Exception:
            pass

    # Borde inferior SOLO en la primera fila (encabezado) y forzar negrita en los runs del encabezado
    # Defensivo: se iteran los `w:tc` crudos de la fila (no `row.cells`, que mapea
    # posiciones de merge a la celda restart) para poder saltar las celdas de
    # continuación de merge vertical (w:vMerge) y las celdas con tablas anidadas.
    if len(table.rows) > 0:
        for tc_el in _raw_row_tcs(table.rows[0]):
            try:
                # Celdas de continuacion de merge vertical: no tocar bordes
                if _is_vmerge_continuation(tc_el):
                    continue
                # Celdas que contienen una tabla anidada (sin texto propio): preservar XML
                if _cell_has_nested_table(tc_el) and not _safe_cell_text(tc_el).strip():
                    continue

                for p_el in tc_el.findall(qn('w:p')):
                    for r_el in p_el.findall(qn('w:r')):
                        _force_run_bold(r_el)

                tcPr = tc_el.get_or_add_tcPr()

                # Eliminar bordes existentes de la celda
                existing_cell_borders = tcPr.find(qn("w:tcBorders"))
                if existing_cell_borders is not None:
                    tcPr.remove(existing_cell_borders)

                # Agregar solo borde inferior al encabezado
                cell_borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                    f'  <w:top w:val="none"/>'
                    f'  <w:left w:val="none"/>'
                    f'  <w:right w:val="none"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(cell_borders)
            except Exception:
                pass

    # Limpiar bordes de celdas en filas de datos (sin tocar merges ni tablas anidadas)
    for row in table.rows[1:]:
        for tc_el in _raw_row_tcs(row):
            try:
                # Celdas de continuacion de merge vertical: no tocar bordes
                if _is_vmerge_continuation(tc_el):
                    continue
                # Celdas con tabla anidada: preservar el XML interno intacto
                if _cell_has_nested_table(tc_el):
                    continue

                tcPr = tc_el.get_or_add_tcPr()
                existing_cell_borders = tcPr.find(qn("w:tcBorders"))
                if existing_cell_borders is not None:
                    tcPr.remove(existing_cell_borders)
                # Agregar sin bordes
                cell_borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="none"/>'
                    f'  <w:bottom w:val="none"/>'
                    f'  <w:left w:val="none"/>'
                    f'  <w:right w:val="none"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(cell_borders)
            except Exception:
                pass


def set_table_apa7_borders(table) -> None:
    """Alias de compatibilidad: bordes APA 7 (solo horizontales)."""
    set_table_borders(table, "apa")


def format_apa_table(
    doc: docx.Document,
    table_data: TableModel,
    rules: APARuleSet,
) -> None:
    """
    Inserta y formatea una tabla completa siguiendo las normas del perfil activo.
    """
    font_name: str = rules.font_family
    font_size: Pt = Pt(rules.font_size_pt)

    # 1. Numero de Tabla ("Tabla 1") en negrita
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = Pt(6)
    p_num.paragraph_format.space_after = Pt(0)
    p_num.paragraph_format.line_spacing = rules.line_spacing
    p_num.paragraph_format.first_line_indent = Inches(0)
    r_num = p_num.add_run(f"{rules.table_label_prefix} {table_data.table_number}")
    r_num.bold = True
    r_num.font.name = font_name
    r_num.font.size = font_size
    r_num.font.color.rgb = RGBColor(0, 0, 0)

    # 2. Titulo de Tabla en cursiva
    if table_data.caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(6)
        p_cap.paragraph_format.line_spacing = rules.line_spacing
        p_cap.paragraph_format.first_line_indent = Inches(0)
        r_cap = p_cap.add_run(table_data.caption)
        r_cap.italic = True
        r_cap.font.name = font_name
        r_cap.font.size = font_size
        r_cap.font.color.rgb = RGBColor(0, 0, 0)

    # 3. Crear Estructura de Tabla
    num_rows: int = len(table_data.rows) + (1 if table_data.headers else 0)
    num_cols: int = (
        len(table_data.headers)
        if table_data.headers
        else (len(table_data.rows[0]) if table_data.rows else 1)
    )

    table = doc.add_table(rows=num_rows, cols=num_cols)
    set_table_borders(table, rules.table_border_style.value)
    fit_table_to_page(table, rules, landscape=False)

    current_row_idx: int = 0

    # Llenar Encabezados (primera fila en negrita)
    if table_data.headers:
        hdr_cells = table.rows[0].cells
        for col_idx, cell_text in enumerate(table_data.headers):
            if col_idx >= len(hdr_cells):
                continue
            try:
                cell = hdr_cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(cell_text)
                r.bold = True
                r.font.name = font_name
                r.font.size = font_size
                r.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
        current_row_idx = 1

    # Llenar Filas de Datos
    for row_data in table_data.rows:
        if current_row_idx >= len(table.rows):
            break
        row_cells = table.rows[current_row_idx].cells
        for col_idx, val in enumerate(row_data):
            if col_idx >= len(row_cells):
                continue
            try:
                cell = row_cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(val))
                r.font.name = font_name
                r.font.size = font_size
                r.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
        current_row_idx += 1

    # 4. Nota al pie de tabla si existe
    if table_data.note:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after = Pt(12)
        p_note.paragraph_format.first_line_indent = Inches(0)
        p_note.paragraph_format.line_spacing = rules.line_spacing

        r_label = p_note.add_run("Nota. ")
        r_label.italic = True
        r_label.font.name = font_name
        r_label.font.size = Pt(10)
        r_label.font.color.rgb = RGBColor(0, 0, 0)

        r_text = p_note.add_run(table_data.note)
        r_text.font.name = font_name
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = RGBColor(0, 0, 0)
