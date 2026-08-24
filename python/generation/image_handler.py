"""
WordAPA7 — Manejador de Imagenes y Figuras APA 7 (Version Configurable)

Inserta y formatea figuras de acuerdo con el estandar APA 7:
1. "Figura X" en negrita (arriba o abajo segun caption_position).
2. Titulo de la figura en cursiva.
3. Imagen con alineacion configurable (centrada por defecto).
4. Nota de la figura opcional abajo.

SOPORTA ESTILOS DE DISENO:
- standard:   Figura centrada, caption debajo (APA estandar)
- sidebar:    Cuadro lateral derecho con borde, texto alrededor
- scientific: Borde negro fino estilo revista, "Figura X." arriba
- corner:     Esquina superior derecha, texto fluye alrededor
- full_width: Ancho completo de pagina, caption centrado
"""

import os
from copy import deepcopy

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from models import APARuleSet, ImageModel


def _apply_image_rotation(run_element, rotation: int) -> None:
    """
    Aplica rotacion (en grados, multiplos de 90) a la imagen inline del run.
    Escribe a:xfrm/a:rot (en 60000ths de grado) dentro del drawing.
    """
    try:
        drawing = run_element.find(qn('w:drawing'))
        if drawing is None:
            return
        inline = drawing.find(qn('wp:inline'))
        if inline is None:
            anchor = drawing.find(qn('wp:anchor'))
            if anchor is not None:
                inline = anchor
        if inline is None:
            return
        # El a:xfrm con rot esta en wp:extent -> a:graphic -> a:graphicData -> pic:pic -> pic:spPr -> a:xfrm
        extent = inline.find(qn('wp:extent'))
        if extent is None:
            return
        graphic = inline.find(qn('a:graphic'))
        if graphic is None:
            return
        graphicData = graphic.find(qn('a:graphicData'))
        if graphicData is None:
            return
        pic = graphicData.find(qn('pic:pic'))
        if pic is None:
            return
        spPr = pic.find(qn('pic:spPr'))
        if spPr is None:
            return
        xfrm = spPr.find(qn('a:xfrm'))
        if xfrm is None:
            xfrm = OxmlElement('a:xfrm')
            spPr.insert(0, xfrm)
        rot = xfrm.find(qn('a:rot'))
        if rot is None:
            rot = OxmlElement('a:rot')
            xfrm.append(rot)
        rot.set(qn('val'), str(int(rotation * 60000)))
        # Si rotamos 90/270, intercambiamos cx/cy para mantener proporciones
        if rotation % 180 != 0:
            cx = extent.get('cx')
            cy = extent.get('cy')
            if cx and cy:
                extent.set('cx', cy)
                extent.set('cy', cx)
    except Exception:
        pass


def _get_alignment(align_str: str):
    """Convierte string de alineacion a constante WD_ALIGN_PARAGRAPH."""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align_str, WD_ALIGN_PARAGRAPH.CENTER)


def _add_paragraph_border(p_element, size_pt: int = 1, color: str = "000000"):
    """Agrega borde negro alrededor de un elemento de parrafo."""
    pPr = p_element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size_pt * 8))  # size_pt * 8 = eighths of a point
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)


def _convert_inline_to_anchor(drawing_elem, width_emu: int, height_emu: int,
                              align_h: str = "right", wrap: str = "square"):
    """
    Convierte un elemento drawing con wp:inline a wp:anchor para imagen flotante.

    Args:
        drawing_elem: Elemento w:drawing
        width_emu: Ancho en EMU
        height_emu: Alto en EMU
        align_h: "left" | "right" | "center"
        wrap: "square" | "tight" | "topAndBottom"
    """
    from docx.oxml import OxmlElement

    # Extraer el wp:inline
    inline = drawing_elem.find(qn('wp:inline'))
    if inline is None:
        return  # Ya es anchor o no es manipulable

    # Posicion horizontal
    pos_h_map = {
        "left": "left",
        "center": "center",
        "right": "right",
    }
    align = pos_h_map.get(align_h, "right")

    # Crear wp:anchor
    anchor_xml = (
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="2" '
        f'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" '
        f'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column">'
        f'<wp:align>{align}</wp:align>'
        f'</wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph">'
        f'<wp:posOffset>0</wp:posOffset>'
        f'</wp:positionV>'
        f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="1" name="Floating"/>'
        f'</wp:anchor>'
    )

    anchor = parse_xml(anchor_xml)

    # Elegir wrap segun estilo
    wrap_elem = anchor.find(qn('wp:wrapNone'))
    if wrap == "square":
        sq = OxmlElement('wp:wrapSquare')
        sq.set('wrapText', 'both')
        sq.set('distL', '114300')  # ~0.1 inch
        sq.set('distR', '114300')
        sq.set('distT', '0')
        sq.set('distB', '0')
        anchor.replace(wrap_elem, sq)
    elif wrap == "tight":
        ti = OxmlElement('wp:wrapTight')
        ti.set('wrapText', 'both')
        ti.set('distL', '114300')
        ti.set('distR', '114300')
        anchor.replace(wrap_elem, ti)
    elif wrap == "topAndBottom":
        tb = OxmlElement('wp:wrapTopAndBottom')
        tb.set('distT', '0')
        tb.set('distB', '0')
        anchor.replace(wrap_elem, tb)

    # Mover contenido de inline al anchor
    for child in list(inline):
        anchor.append(deepcopy(child))

    # Reemplazar inline con anchor en el drawing
    drawing_elem.replace(inline, anchor)


def format_apa_figure(doc: docx.Document, img_data: ImageModel, rules: APARuleSet):
    """
    Inserta la imagen en el documento con estructura APA 7 configurable.

    Respeta:
    - img_data.design_style (standard|sidebar|scientific|corner|full_width)
    - img_data.alignment (left|center|right)
    - img_data.width_inches / width_cm
    - img_data.caption_position (above|below)
    - img_data.wrap_style
    """
    design = img_data.design_style or "standard"
    alignment = _get_alignment(img_data.alignment)

    # Ajustes segun estilo de diseno
    effective_width_cm = img_data.width_cm
    effective_alignment = img_data.alignment
    effective_wrap = img_data.wrap_style
    add_border = False

    if design == "full_width":
        usable_cm = None
        try:
            sec = doc.sections[0] if doc.sections else None
            if (sec is not None and sec.page_width is not None
                    and sec.left_margin is not None and sec.right_margin is not None):
                usable_cm = float((sec.page_width - sec.left_margin - sec.right_margin) / 360000)
        except Exception:
            usable_cm = None
        effective_width_cm = usable_cm if usable_cm else 16.0  # Ancho completo del area de texto APA
        effective_alignment = "center"
    elif design == "scientific":
        add_border = True
        if img_data.caption_position == "above" or img_data.caption_position != "below":
            pass  # Figura label arriba ya se maneja
    elif design == "sidebar":
        effective_alignment = "right"
        effective_wrap = "square"
    elif design == "corner":
        effective_alignment = "right"
        effective_wrap = "tight"

    # 1. Numero de Figura + Titulo
    if img_data.caption_position == "above" or img_data.caption_position != "below":
        _add_figure_label_and_caption(doc, img_data, rules)

    # 2. Imagen
    p_img = doc.add_paragraph()
    p_img.alignment = _get_alignment(effective_alignment)
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(6)
    p_img.paragraph_format.first_line_indent = Inches(0)
    # Evitar que la imagen quede partida entre dos páginas
    p_img.paragraph_format.keep_together = True
    p_img.paragraph_format.widow_control = True

    if img_data.file_path and os.path.exists(img_data.file_path):
        run = p_img.add_run()
        run_element = run._element

        # Calcular ancho
        width = None
        if effective_width_cm and effective_width_cm > 0:
            width = Inches(effective_width_cm / 2.54)

        if width is not None:
            run.add_picture(img_data.file_path, width=width)
        else:
            run.add_picture(img_data.file_path)

        # Aplicar rotacion (0, 90, 180, 270) via a:xfrm/a:rot dentro del drawing
        rotation = int(getattr(img_data, "rotation", 0) or 0)
        if rotation:
            _apply_image_rotation(run_element, rotation)

        # Texto alternativo (accesibilidad) — se escribe en el drawing y el docProps
        alt_text = (getattr(img_data, "alt_text", "") or "").strip()
        if alt_text:
            try:
                drawing = run_element.find(qn('w:drawing'))
                if drawing is not None:
                    wp_inline = drawing.find(qn('wp:inline'))
                    target = wp_inline if wp_inline is not None else drawing.find(qn('wp:anchor'))
                    if target is not None:
                        descr = OxmlElement('wp:docPr')
                        descr.set(qn('id'), '0')
                        descr.set(qn('name'), 'Imagen')
                        descr.set(qn('descr'), alt_text[:255])
                        target.insert(0, descr)
            except Exception:
                pass

        # Aplicar estilos de diseno avanzados (floating, borders)
        if design in ("sidebar", "corner"):
            # Convertir imagen inline a floating anchor
            drawing = run_element.find(qn('w:drawing'))
            if drawing is not None:
                # Extraer dimensiones EMU del inline
                inline = drawing.find(qn('wp:inline'))
                if inline is not None:
                    extent = inline.find(qn('wp:extent'))
                    if extent is not None:
                        cx = extent.get('cx', '0')
                        cy = extent.get('cy', '0')
                        wrap_map = {
                            "square": "square",
                            "tight": "tight",
                            "top_and_bottom": "topAndBottom",
                        }
                        _convert_inline_to_anchor(
                            drawing,
                            int(cx), int(cy),
                            align_h=effective_alignment,
                            wrap=wrap_map.get(effective_wrap, "square"),
                        )
        elif design == "scientific" and add_border:
            _add_paragraph_border(p_img._element, size_pt=1, color="000000")
    else:
        r_ph = p_img.add_run(f"[Imagen: {img_data.filename}]")
        r_ph.italic = True

    # 3. Titulo / caption ABAJO si se configuro asi
    if img_data.caption_position == "below":
        _add_figure_label_and_caption(doc, img_data, rules)

    # 4. Nota al pie de figura si existe
    internal_tokens = {"author_card", "vertical_line", "shape_group", "shape_textbox"}
    if img_data.note and img_data.note.strip() not in internal_tokens:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after = Pt(12)
        p_note.paragraph_format.first_line_indent = Inches(0)

        # Para estilo scientific, la nota tambien va con borde
        if design == "scientific":
            _add_paragraph_border(p_note._element, size_pt=1, color="000000")

        r_label = p_note.add_run("Nota. ")
        r_label.italic = True
        r_label.font.name = rules.font_family
        r_label.font.size = Pt(10)

        r_text = p_note.add_run(img_data.note)
        r_text.font.name = rules.font_family
        r_text.font.size = Pt(10)


def _add_figure_label_and_caption(doc: docx.Document, img_data: ImageModel, rules: APARuleSet):
    """Anade el numero de figura y caption como parrafo independiente."""
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = Pt(12)
    p_num.paragraph_format.space_after = Pt(0)
    p_num.paragraph_format.line_spacing = rules.line_spacing
    p_num.paragraph_format.first_line_indent = Inches(0)

    r_num = p_num.add_run(f"Figura {img_data.figure_number}")
    r_num.bold = True
    r_num.font.name = rules.font_family
    r_num.font.size = Pt(rules.font_size_pt)
    p_num.paragraph_format.keep_with_next = True

    # Titulo de la Figura en Cursiva
    if img_data.caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(6)
        p_cap.paragraph_format.line_spacing = rules.line_spacing
        p_cap.paragraph_format.first_line_indent = Inches(0)
        p_cap.paragraph_format.keep_with_next = True

        r_cap = p_cap.add_run(img_data.caption)
        r_cap.italic = True
        r_cap.font.name = rules.font_family
        r_cap.font.size = Pt(rules.font_size_pt)
