"""
Script de Auditoría y Verificación Pipeline E2E para Estudio_Trabajo (1).docx
"""

import sys
from pathlib import Path
import json
import requests
from docx import Document

# Asegurar path de python
sys.path.insert(0, str(Path(__file__).resolve().parent))

from parsing.docx_parser import parse_docx_bytes

TARGET_FILE = Path(r"C:\Users\--X\Desktop\Estudio_Trabajo (1).docx")
SERVER_URL = "http://localhost:8742"

def main():
    print(f"=== AUDITORÍA PIPELINE E2E: {TARGET_FILE.name} ===")
    if not TARGET_FILE.exists():
        print(f"[ERROR] No se encontró el archivo {TARGET_FILE}")
        return

    # 1. Parseo interno de diagnóstico
    with open(TARGET_FILE, "rb") as f:
        file_bytes = f.read()

    doc_model = parse_docx_bytes(file_bytes, TARGET_FILE.name, "session_test_estudio", Path("storage"))
    cover_elems = [e for e in doc_model.elements if e.is_cover_section]
    headings = [e for e in doc_model.elements if e.type == "heading"]
    tables = [e for e in doc_model.elements if e.type == "table"]
    images = [e for e in doc_model.elements if e.type == "image"]
    bullets = [e for e in doc_model.elements if e.type in ("bullet", "numbered_list")]

    print(f"[1] Diagnóstico de Clasificación:")
    print(f"    - Elementos de Portada: {len(cover_elems)}")
    print(f"    - Títulos Validados: {len(headings)}")
    print(f"    - Tablas: {len(tables)}")
    print(f"    - Imágenes: {len(images)}")

    # 2. Upload al servidor FastAPI
    print(f"\n[2] Enviando archivo a {SERVER_URL}/api/upload ...")
    with open(TARGET_FILE, "rb") as f:
        resp = requests.post(f"{SERVER_URL}/api/upload", files={"file": (TARGET_FILE.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})

    if resp.status_code != 200:
        print(f"[ERROR] Fail en upload: {resp.status_code} - {resp.text}")
        return

    upload_data = resp.json()
    session_id = upload_data.get("session_id")
    print(f"    - Upload Exitoso! Session ID: {session_id}")

    # 3. Generar DOCX APA 7
    print(f"\n[3] Generando documento APA 7 vía {SERVER_URL}/api/generate ...")
    gen_payload = {
        "session_id": session_id,
        "apa_format": "student",
        "work_mode": "review"
    }
    gen_resp = requests.post(f"{SERVER_URL}/api/generate", json=gen_payload)
    if gen_resp.status_code != 200:
        print(f"[ERROR] Fail en generación: {gen_resp.status_code} - {gen_resp.text}")
        return

    download_url = gen_resp.json().get("download_url")
    print(f"    - Generación Exitosa! Download URL: {download_url}")

    # 4. Descargar y verificar DOCX
    out_file = Path("output_estudio_trabajo_apa7.docx")
    file_resp = requests.get(f"{SERVER_URL}{download_url}")
    with open(out_file, "wb") as f:
        f.write(file_resp.content)

    print(f"\n[4] DOCX generado guardado en {out_file.absolute()}")
    print(f"    - Tamaño de salida: {len(file_resp.content)} bytes")

    out_doc = Document(out_file)
    print(f"    - Párrafos en DOCX salida: {len(out_doc.paragraphs)}")
    print(f"    - Tablas en DOCX salida: {len(out_doc.tables)}")

    print("\n--- Primeros 10 Párrafos del Documento Generado ---")
    for idx, p in enumerate(out_doc.paragraphs[:10]):
        text_trunc = (p.text or "")[:70].replace("\n", " ")
        style_name = p.style.name if p.style else "Normal"
        xml_snippet = p._element.xml[:120].replace("\n", "")
        print(f"  P#{idx:02d} | style={style_name:<20} | text='{text_trunc}' | xml='{xml_snippet}'")

if __name__ == "__main__":
    main()
