"""
WordAPA7 — Generador de Corpus de Prueba

Genera documentos .docx de prueba para la suite de tests automatizados.
Cada documento simula un escenario real del mundo académico hispano.

Ejecutar: cd python && python tests/generate_test_docs.py

Los documentos se generan en: corpus/
"""

import sys
from pathlib import Path

# Asegurar path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

CORPUS_DIR = Path(__file__).parent.parent.parent / "corpus"


def ensure_corpus_dir():
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Directorio de corpus: {CORPUS_DIR}")


# ── doc_01_limpio.docx ─────────────────────────────────────────────────────────

def build_doc_01_limpio():
    """Documento bien escrito con estilos Word correctos."""
    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Heading 1
    h1 = doc.add_heading("Impacto de la Tecnología en la Educación Superior", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # Párrafo
    p = doc.add_paragraph(
        "La integración de tecnologías digitales en el ámbito educativo ha transformado "
        "significativamente los procesos de enseñanza y aprendizaje en las instituciones "
        "de educación superior. Investigaciones recientes (García, 2023) han demostrado "
        "que el uso de plataformas virtuales mejora la retención de conocimientos en "
        "estudiantes universitarios de primer año."
    )
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading 2
    h2 = doc.add_heading("Marco Teórico de la Investigación", level=2)
    for run in h2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)

    # Párrafo
    p2 = doc.add_paragraph(
        "El constructivismo social de Vygotsky (1978) proporciona un marco teórico "
        "sólido para entender cómo los estudiantes construyen conocimiento a través "
        "de la interacción con sus pares y con herramientas tecnológicas. Según "
        "López et al. (2021), los entornos virtuales de aprendizaje facilitan la "
        "colaboración entre estudiantes geográficamente dispersos."
    )
    p2.style = doc.styles['Normal']
    for run in p2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading 3
    h3 = doc.add_heading("Diseño Metodológico del Estudio", level=3)
    for run in h3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Párrafo
    p3 = doc.add_paragraph(
        "Se utilizó un diseño de investigación mixto que combinó métodos cuantitativos "
        "y cualitativos para obtener una comprensión integral del fenómeno estudiado. "
        "La muestra estuvo compuesta por 250 estudiantes de tres universidades públicas "
        "mexicanas durante el ciclo escolar 2024-2025."
    )
    p3.style = doc.styles['Normal']
    for run in p3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    path = CORPUS_DIR / "doc_01_limpio.docx"
    doc.save(str(path))
    print(f"  [OK] doc_01_limpio.docx — {len(doc.paragraphs)} párrafos")
    return path


# ── doc_02_fuentes_mixtas.docx ─────────────────────────────────────────────────

def build_doc_02_fuentes_mixtas():
    """Documento con Calibri + Arial + Times mezclados (pegado de internet)."""
    doc = Document()

    # Párrafo con Calibri
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Este párrafo fue copiado de un sitio web y está en Calibri 11pt. ")
    r1.font.name = "Calibri"
    r1.font.size = Pt(11)
    r2 = p1.add_run("La información proviene de múltiples fuentes en línea.")
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)

    # Título en Arial
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Análisis de Datos Cuantitativos")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(16)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Párrafo con Arial
    p2 = doc.add_paragraph()
    r3 = p2.add_run("El análisis estadístico se realizó utilizando SPSS versión 28. ")
    r3.font.name = "Arial"
    r3.font.size = Pt(11)
    r4 = p2.add_run("Los resultados muestran una correlación significativa (p < .001).")
    r4.font.name = "Arial"
    r4.font.size = Pt(11)

    # Subtítulo en Times New Roman
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Resultados Principales")
    r_sub.bold = True
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(13)

    # Párrafo con Times
    p3 = doc.add_paragraph()
    r5 = p3.add_run("Los hallazgos principales indican que el 78% de los participantes ")
    r5.font.name = "Times New Roman"
    r5.font.size = Pt(12)
    r6 = p3.add_run("reportó una mejora significativa en su rendimiento académico tras "
                     "la implementación del programa de tutoría virtual personalizada.")
    r6.font.name = "Times New Roman"
    r6.font.size = Pt(12)

    # Mezcla en el mismo párrafo
    p4 = doc.add_paragraph()
    r7 = p4.add_run("Nota: ")
    r7.bold = True
    r7.font.name = "Arial"
    r7.font.size = Pt(10)

    r8 = p4.add_run("Este texto tiene tres fuentes diferentes. ")
    r8.font.name = "Calibri"
    r8.font.size = Pt(11)

    r9 = p4.add_run("La normalización debe unificar todo a una sola fuente APA 7.")
    r9.font.name = "Times New Roman"
    r9.font.size = Pt(12)

    path = CORPUS_DIR / "doc_02_fuentes_mixtas.docx"
    doc.save(str(path))
    print("  [OK] doc_02_fuentes_mixtas.docx — Calibri + Arial + Times mezclados")
    return path


# ── doc_03_bullets_manuales.docx ───────────────────────────────────────────────

def build_doc_03_bullets_manuales():
    """Documento con viñetas hechas con caracteres manuales (•, -, –)."""
    doc = Document()

    h = doc.add_paragraph()
    r = h.add_run("Lista de Verificación para Investigación")
    r.bold = True
    r.font.size = Pt(14)

    bullets = [
        "• Revisar la literatura existente sobre el tema de investigación.",
        "- Definir las variables dependientes e independientes del estudio.",
        "– Seleccionar la muestra representativa de la población objetivo.",
        "• Diseñar los instrumentos de recolección de datos apropiados.",
        "- Obtener la aprobación del comité de ética institucional.",
        "– Realizar una prueba piloto con al menos 30 participantes.",
        "1. Analizar los datos utilizando software estadístico especializado.",
        "2. Redactar el informe final siguiendo el formato APA 7.",
        "3. Someter el artículo a revisión por pares en una revista indexada.",
    ]

    for text in bullets:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    path = CORPUS_DIR / "doc_03_bullets_manuales.docx"
    doc.save(str(path))
    print("  [OK] doc_03_bullets_manuales.docx — •, -, – manuales + lista numerada")
    return path


# ── doc_04_headings_sin_estilo.docx ────────────────────────────────────────────

def build_doc_04_headings_sin_estilo():
    """Documento con headings en negrita manual, no estilo Word."""
    doc = Document()

    # Heading simulado: centrado + negrita (nivel 1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Evaluación del Aprendizaje en Entornos Virtuales")
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.name = "Times New Roman"
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_text1 = doc.add_paragraph(
        "La evaluación del aprendizaje en entornos virtuales representa un desafío "
        "significativo para los docentes universitarios, quienes deben adaptar sus "
        "estrategias de evaluación a las características específicas de la modalidad "
        "en línea sin comprometer la validez y confiabilidad de los instrumentos."
    )
    for run in p_text1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading simulado: izquierda + negrita (nivel 2)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Tipos de Evaluación en Línea")
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.name = "Times New Roman"

    p_text2 = doc.add_paragraph(
        "Existen diversos tipos de evaluación que pueden implementarse en entornos "
        "virtuales, incluyendo la evaluación diagnóstica, formativa y sumativa. "
        "Cada una cumple funciones específicas dentro del proceso educativo."
    )
    for run in p_text2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading simulado: izquierda + negrita + cursiva (nivel 3)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Evaluación formativa automatizada.")
    r3.bold = True
    r3.italic = True
    r3.font.size = Pt(12)
    r3.font.name = "Times New Roman"

    p_text3 = doc.add_paragraph(
        "La evaluación formativa automatizada utiliza algoritmos de inteligencia "
        "artificial para proporcionar retroalimentación inmediata a los estudiantes "
        "sobre su desempeño en tiempo real durante el proceso de aprendizaje."
    )
    for run in p_text3.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    path = CORPUS_DIR / "doc_04_headings_sin_estilo.docx"
    doc.save(str(path))
    print("  [OK] doc_04_headings_sin_estilo.docx — Headings con formato manual")
    return path


# ── doc_11_citas_mal.docx ──────────────────────────────────────────────────────

def build_doc_11_citas_mal():
    """Documento con citas que contienen errores comunes de formato."""
    doc = Document()

    h = doc.add_paragraph()
    r = h.add_run("Revisión de Literatura con Citas Problemáticas")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    paragraphs_with_errors = [
        # Sin coma entre autor y año
        "Diversos estudios han señalado la importancia de la práctica deliberada "
        "en el desarrollo de habilidades cognitivas complejas (Ericsson 2016).",

        # et al sin punto
        "Las investigaciones más recientes en el campo de la neurociencia educativa "
        "(García et al, 2023) sugieren que el cerebro adolescente es particularmente "
        "sensible a los estímulos ambientales durante el aprendizaje.",

        # Falta & en múltiples autores
        "El metaanálisis realizado por Johnson y Smith (2022) reveló que las "
        "intervenciones basadas en mindfulness tienen un efecto moderado pero "
        "significativo sobre la reducción del estrés académico en universitarios.",

        # Año sin paréntesis en cita narrativa
        "Martínez 2021 encontró que los estudiantes que utilizan técnicas de "
        "estudio espaciado obtienen calificaciones significativamente más altas "
        "que aquellos que utilizan técnicas de estudio masivo.",

        # Falta "p." o "pp." en cita con página
        "Como señalan diversos autores (Rodríguez, 2020, 145), la motivación "
        "intrínseca es un predictor más robusto del rendimiento académico que "
        "la motivación extrínseca en estudiantes de posgrado.",

        # Cita correcta (control)
        "La teoría del aprendizaje significativo propuesta por Ausubel (1963) "
        "continúa siendo uno de los pilares fundamentales de la psicología "
        "educativa contemporánea (García, 2023, p. 45).",
    ]

    for text in paragraphs_with_errors:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    path = CORPUS_DIR / "doc_11_citas_mal.docx"
    doc.save(str(path))
    print("  [OK] doc_11_citas_mal.docx — Citas con errores comunes")
    return path


# ── MAIN ───────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print(" Generando Corpus de Prueba WordAPA7")
    print("=" * 60)

    ensure_corpus_dir()

    build_doc_01_limpio()
    build_doc_02_fuentes_mixtas()
    build_doc_03_bullets_manuales()
    build_doc_04_headings_sin_estilo()
    build_doc_11_citas_mal()

    print("=" * 60)
    print(" Corpus generado exitosamente.")
    print(f" Ubicación: {CORPUS_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
