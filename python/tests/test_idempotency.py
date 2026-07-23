"""
WordAPA7 — Tests de Idempotencia

Verifica que reprocesar un documento que ya pasó por WordAPA7:
1. No duplique portadas
2. No re-numere figuras incorrectamente
3. SHA-256 del mismo archivo retorna el mismo hash
4. Documento ya procesado sea correctamente detectado
"""

import sys
import io
import hashlib
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
import docx
from docx.shared import Inches

from models import (
    DocumentModel, ElementModel, ElementType, APARuleSet,
    PortadaData, APAFormat
)
from generation.generator import generate_apa7_docx


# ── HELPERS ─────────────────────────────────────────────────────────────────────

def compute_sha256(filepath: Path) -> str:
    """Calcula el hash SHA-256 de un archivo."""
    sha = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha.update(chunk)
    return sha.hexdigest()


def compute_bytes_sha256(data: bytes) -> str:
    """Calcula el hash SHA-256 de bytes."""
    return hashlib.sha256(data).hexdigest()


# ── TESTS DE HASH ──────────────────────────────────────────────────────────────

class TestSHA256Hashing:
    """Verifica el comportamiento del hashing SHA-256 para documentos."""

    def test_same_bytes_same_hash(self):
        """Los mismos bytes deben producir el mismo hash."""
        content = b"Este es un contenido de prueba para hashing."
        hash1 = compute_bytes_sha256(content)
        hash2 = compute_bytes_sha256(content)
        assert hash1 == hash2

    def test_different_bytes_different_hash(self):
        """Bytes diferentes deben producir hash diferente."""
        content1 = b"Contenido A"
        content2 = b"Contenido B"
        hash1 = compute_bytes_sha256(content1)
        hash2 = compute_bytes_sha256(content2)
        assert hash1 != hash2

    def test_same_file_same_hash(self, test_output_dir):
        """El mismo archivo debe producir el mismo hash en lecturas sucesivas."""
        test_file = test_output_dir / "hash_test.bin"
        test_file.write_bytes(b"Contenido de prueba para hash de archivo (x 10)\n" * 10)

        hash1 = compute_sha256(test_file)
        hash2 = compute_sha256(test_file)
        assert hash1 == hash2

    def test_file_hash_stable_across_generations(self, rules, test_output_dir):
        """Generar el mismo documento dos veces debe producir hashes consistentes
        (el contenido interno debe ser determinístico)."""
        doc_model = DocumentModel(
            session_id="test_hash_stable",
            file_name="hash_stable.docx",
            elements=[
                ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                            text="Documento de Prueba de Hash"),
                ElementModel(id="e2", type=ElementType.PARAGRAPH,
                            text="Este es un párrafo de prueba."),
            ],
        )

        out1 = test_output_dir / "hash_gen_1.docx"
        out2 = test_output_dir / "hash_gen_2.docx"

        generate_apa7_docx(doc_model, out1, rules)
        generate_apa7_docx(doc_model, out2, rules)

        hash1 = compute_sha256(out1)
        hash2 = compute_sha256(out2)

        # Los hashes deben ser iguales (el generador es determinístico)
        assert hash1 == hash2, (
            f"Generación no determinística: hash1={hash1[:16]}... != hash2={hash2[:16]}..."
        )


# ── TESTS DE REPROCESAMIENTO ────────────────────────────────────────────────────

class TestReprocessingDetection:
    """Verifica detección de documentos ya procesados."""

    def test_previously_processed_flag(self):
        """DocumentModel.meta.previously_processed debe ser False en docs nuevos."""
        doc_model = DocumentModel(
            session_id="test_new",
            file_name="nuevo.docx",
        )
        assert doc_model.meta.previously_processed is False

    def test_previously_processed_can_be_true(self):
        """Cuando se marca como procesado, debe reflejarse."""
        doc_model = DocumentModel(
            session_id="test_processed",
            file_name="procesado.docx",
        )
        doc_model.meta.previously_processed = True
        assert doc_model.meta.previously_processed is True

    def test_processed_document_hash_detected(self):
        """Simular detección de documento ya procesado usando source_hash."""
        original_hash = "abc123def456"

        doc_model = DocumentModel(
            session_id="test_hash_detect",
            file_name="doc.docx",
        )
        doc_model.meta.source_hash = original_hash

        # Verificar que podemos leer el hash
        assert doc_model.meta.source_hash == original_hash

    def test_different_documents_have_different_hashes(self, rules, test_output_dir):
        """Documentos con diferente contenido deben tener diferente hash."""
        # Documento A
        doc_a = DocumentModel(
            session_id="test_a",
            file_name="doc_a.docx",
            elements=[
                ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Contenido A"),
            ],
        )
        out_a = test_output_dir / "idem_a.docx"
        generate_apa7_docx(doc_a, out_a, rules)

        # Documento B (diferente contenido)
        doc_b = DocumentModel(
            session_id="test_b",
            file_name="doc_b.docx",
            elements=[
                ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Contenido B"),
            ],
        )
        out_b = test_output_dir / "idem_b.docx"
        generate_apa7_docx(doc_b, out_b, rules)

        hash_a = compute_sha256(out_a)
        hash_b = compute_sha256(out_b)

        assert hash_a != hash_b, "Documentos diferentes no deberían tener el mismo hash"


# ── TESTS DE NO DUPLICACIÓN ────────────────────────────────────────────────────

class TestNoDuplication:
    """Verifica que al reprocesar no se dupliquen elementos."""

    def test_heading_not_duplicated_on_regeneration(self, rules, test_output_dir):
        """Regenerar el mismo modelo no debe duplicar headings."""
        elements = [
            ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                        text="Título Único"),
            ElementModel(id="e2", type=ElementType.PARAGRAPH,
                        text="Párrafo de prueba."),
        ]

        doc_model = DocumentModel(
            session_id="test_nodup",
            file_name="nodup.docx",
            elements=elements,
        )

        out = test_output_dir / "nodup_output.docx"
        generate_apa7_docx(doc_model, out, rules)

        doc = docx.Document(str(out))
        # Contar ocurrencias de "Título Único"
        title_count = sum(1 for p in doc.paragraphs if "Título Único" in p.text)
        assert title_count == 1, (
            f"El título aparece {title_count} veces, debería aparecer 1"
        )

    def test_portada_not_duplicated(self, rules, test_output_dir):
        """La portada no debe duplicarse al regenerar."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH,
                        text="Contenido después de la portada."),
        ]

        portada = PortadaData(
            apa_format=APAFormat.STUDENT,
            title="Trabajo de Investigación",
            author="Autor Principal",
            institution="Universidad de Prueba",
            course="Curso de Metodología",
            instructor="Dr. Profesor Titular",
            date="2025",
        )

        doc_model = DocumentModel(
            session_id="test_portada_dup",
            file_name="portada_test.docx",
            elements=elements,
        )

        out = test_output_dir / "portada_output.docx"
        generate_apa7_docx(doc_model, out, rules, portada=portada)

        doc = docx.Document(str(out))
        author_count = sum(1 for p in doc.paragraphs if "Autor Principal" in p.text)
        assert author_count == 1, (
            f"El autor aparece {author_count} veces, debería aparecer 1 vez"
        )

    def test_references_not_duplicated(self, rules, test_output_dir, references_sample):
        """La sección de referencias no debe duplicarse al regenerar."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH,
                        text="Texto con cita a García (2023)."),
        ]

        doc_model = DocumentModel(
            session_id="test_ref_dup",
            file_name="ref_dup_test.docx",
            elements=elements,
        )

        out = test_output_dir / "ref_dup_output.docx"
        generate_apa7_docx(doc_model, out, rules, references=references_sample)

        doc = docx.Document(str(out))
        ref_headers = sum(1 for p in doc.paragraphs if p.text.strip() == "Referencias")
        assert ref_headers <= 1, (
            f"'Referencias' aparece {ref_headers} veces, debería aparecer <= 1"
        )
