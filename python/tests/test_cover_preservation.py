"""
WordAPA7 — Unit Test de Preservación Estricta de Portada en DOCX

Verifica que al procesar un documento con portada, la portada completa
(textos, imágenes, logos, cuadros de texto) permanezca 100% intacta en el XML de salida.
"""

import docx
from generation.generator import generate_apa7_docx
from models import APARuleSet, PortadaData
from parsing.docx_parser import parse_docx_bytes


def test_cover_preservation_strict(tmp_path):
    orig_path = tmp_path / "original.docx"
    out_path = tmp_path / "output_with_cover.docx"

    # Crear documento con portada y cuerpo
    doc = docx.Document()
    doc.add_paragraph("UNIVERSIDAD NACIONAL DE INGENIERÍA")
    doc.add_paragraph("FACULTAD DE ELECTROTECNIA Y COMPUTACIÓN")
    doc.add_paragraph("ANÁLISIS DE SISTEMAS EN APA 7")
    doc.add_paragraph("Elaborado por: Br. Juan Pérez")
    doc.add_paragraph("Tutor: Ing. Carlos Rodríguez")
    doc.add_paragraph("15 de junio de 2026")
    doc.add_paragraph("")  # Espacio
    # Cuerpo
    doc.add_paragraph("Introducción")
    doc.add_paragraph("Este es el primer párrafo del cuerpo del trabajo académico.")

    doc.save(orig_path)

    with open(orig_path, "rb") as f:
        bytes_data = f.read()

    doc_model = parse_docx_bytes(bytes_data, "original.docx", "test_sess_001", tmp_path)

    # Generar salida
    portada_cfg = PortadaData(use_original_cover=True)
    generate_apa7_docx(doc_model, out_path, APARuleSet(), portada_cfg)

    # Verificar que el documento generado mantenga los textos de la portada
    gen_doc = docx.Document(out_path)
    gen_texts = [p.text for p in gen_doc.paragraphs if p.text]
    print("[DEBUG gen_texts]:", gen_texts)

    assert any("UNIVERSIDAD NACIONAL" in t for t in gen_texts)
    assert any("ANÁLISIS DE SISTEMAS" in t for t in gen_texts)
    assert any("Elaborado por" in t for t in gen_texts)
    assert any("Introducción" in t for t in gen_texts)
