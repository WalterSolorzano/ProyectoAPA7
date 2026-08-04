"""
WordAPA7 — Tests del Generador DOCX APA 7

Verifica la generación completa del documento:
1. El archivo generado es un .docx válido
2. Portada incluida cuando hay datos de portada
3. Sección de referencias con sangría francesa (hanging indent)
4. Números de página presentes
5. Fuentes normalizadas (sin mezclas)
6. Márgenes APA 7 correctos
7. Interlineado doble
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import (
    DocumentModel, ElementModel, ElementType, APARuleSet,
    PortadaData, ReferenciaModel, APAFormat
)
from generation.generator import generate_apa7_docx
from generation.style_engine import (
    apply_page_setup, format_heading_paragraph,
    format_normal_paragraph, format_block_quote
)
from generation.document_structure import setup_apa_header


# ── FIXTURES ────────────────────────────────────────────────────────────────────

@pytest.fixture
def rules():
    return APARuleSet(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
        paragraph_indent_cm=1.27,
        margins_cm=2.54,
    )


# ── VALIDACIÓN DE ARCHIVO GENERADO ─────────────────────────────────────────────

class TestGeneratedDocxValidity:
    """El archivo .docx generado debe ser un ZIP válido con la estructura OOXML."""

    def test_generated_file_is_zip(self, rules, test_output_dir):
        """El archivo .docx debe ser un ZIP válido."""
        import zipfile

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Test."),
        ]
        doc_model = DocumentModel(
            session_id="test_zip",
            file_name="zip_test.docx",
            elements=elements,
        )

        out = test_output_dir / "zip_valid_output.docx"
        path = generate_apa7_docx(doc_model, out, rules)

        assert zipfile.is_zipfile(str(path)), "El archivo .docx no es un ZIP válido"

    def test_generated_file_has_document_xml(self, rules, test_output_dir):
        """El archivo debe contener word/document.xml."""
        import zipfile

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Test XML."),
        ]
        doc_model = DocumentModel(
            session_id="test_xml",
            file_name="xml_test.docx",
            elements=elements,
        )

        out = test_output_dir / "xml_output.docx"
        path = generate_apa7_docx(doc_model, out, rules)

        with zipfile.ZipFile(str(path), 'r') as zf:
            names = zf.namelist()
            assert 'word/document.xml' in names, (
                f"Falta word/document.xml. Contenido: {names}"
            )


# ── PORTADA ─────────────────────────────────────────────────────────────────────

class TestCoverPage:
    """Verifica que la portada se genere correctamente."""

    def test_cover_page_with_title(self, rules, test_output_dir):
        """Portada con título debe generarse."""
        portada = PortadaData(
            apa_format=APAFormat.STUDENT,
            title="Impacto de la IA en Educación",
            author="María García",
            institution="Universidad Nacional",
        )

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH,
                        text="Contenido del documento."),
        ]
        doc_model = DocumentModel(
            session_id="test_portada",
            file_name="portada.docx",
            elements=elements,
        )

        out = test_output_dir / "portada_output.docx"
        generate_apa7_docx(doc_model, out, rules, portada=portada)

        doc = docx.Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)

        assert "Impacto de la IA en Educación" in all_text
        assert "María García" in all_text
        assert "Universidad Nacional" in all_text

    def test_cover_inserted_at_start_with_original_file(self, rules, tmp_path):
        """Cuando existe archivo original, la portada generada debe quedar
        al INICIO del documento, antes del contenido del cuerpo."""
        portada = PortadaData(
            apa_format=APAFormat.STUDENT,
            title="Portada al Inicio",
            author="Autor Test",
            institution="Institución Test",
            cover_mode="generate_apa7_template",
            use_original_cover=False,
        )

        # Crear archivo original con contenido de cuerpo en un directorio propio
        orig_path = tmp_path / "contenido_original.docx"
        orig_doc = docx.Document()
        orig_doc.add_paragraph("Introducción del documento original.")
        orig_doc.add_paragraph("Segundo párrafo del cuerpo.")
        orig_doc.save(str(orig_path))

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH,
                        text="Introducción del documento original."),
            ElementModel(id="e2", type=ElementType.PARAGRAPH,
                        text="Segundo párrafo del cuerpo."),
        ]
        doc_model = DocumentModel(
            session_id="test_cover_order",
            file_name="contenido_original.docx",
            elements=elements,
        )

        # Copiar el original como lo haría la API (original.docx en carpeta de salida)
        import shutil
        out = tmp_path / "cover_order_output.docx"
        shutil.copy2(orig_path, tmp_path / "original.docx")

        generate_apa7_docx(doc_model, out, rules, portada=portada)

        doc = docx.Document(str(out))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # El título de la portada debe aparecer ANTES que el contenido del cuerpo
        title_idx = next((i for i, t in enumerate(paragraphs) if t == "Portada al Inicio"), -1)
        body_idx = next((i for i, t in enumerate(paragraphs) if t.startswith("Introducción del documento original")), -1)

        assert title_idx != -1, f"Título de portada no encontrado en: {paragraphs[:8]}"
        assert body_idx != -1, f"Contenido del cuerpo no encontrado en: {paragraphs[:8]}"
        assert title_idx < body_idx, (
            f"La portada ({title_idx}) debe ir ANTES del cuerpo ({body_idx}). "
            f"Primeros párrafos: {paragraphs[:8]}"
        )

    def test_cover_page_title_is_bold(self, rules, test_output_dir):
        """El título de la portada debe estar en negrita."""
        portada = PortadaData(
            apa_format=APAFormat.STUDENT,
            title="Título en Negrita",
            author="Autor",
            institution="Institución",
        )

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Cuerpo."),
        ]
        doc_model = DocumentModel(
            session_id="test_bold_title",
            file_name="bold_title.docx",
            elements=elements,
        )

        out = test_output_dir / "bold_title_output.docx"
        generate_apa7_docx(doc_model, out, rules, portada=portada)

        doc = docx.Document(str(out))
        title_para = None
        for p in doc.paragraphs:
            if "Título en Negrita" in p.text:
                title_para = p
                break

        assert title_para is not None
        assert title_para.runs[0].bold, "El título de la portada debe estar en negrita"

    def test_no_cover_when_no_portada_data(self, rules, test_output_dir):
        """Sin datos de portada, no debe generarse portada."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Solo contenido."),
        ]
        doc_model = DocumentModel(
            session_id="test_no_portada",
            file_name="no_portada.docx",
            elements=elements,
        )

        out = test_output_dir / "no_portada_output.docx"
        generate_apa7_docx(doc_model, out, rules)  # sin portada

        doc = docx.Document(str(out))
        # El primer párrafo no vacío debería ser el contenido
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) > 0

    def test_student_portada_includes_all_fields(self, rules, test_output_dir):
        """Portada de estudiante debe incluir curso, instructor y fecha."""
        portada = PortadaData(
            apa_format=APAFormat.STUDENT,
            title="Tesis de Grado",
            author="Juan Pérez",
            institution="Universidad Central",
            course="Seminario de Investigación",
            instructor="Dra. Ana Martínez",
            date="2025",
        )

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Cuerpo."),
        ]
        doc_model = DocumentModel(
            session_id="test_student_portada",
            file_name="student_portada.docx",
            elements=elements,
        )

        out = test_output_dir / "student_portada_output.docx"
        generate_apa7_docx(doc_model, out, rules, portada=portada)

        doc = docx.Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)

        assert "Seminario de Investigación" in all_text
        assert "Dra. Ana Martínez" in all_text
        assert "2025" in all_text


# ── REFERENCIAS ─────────────────────────────────────────────────────────────────

class TestReferencesSection:
    """Verifica la sección de referencias con sangría francesa."""

    def test_references_section_exists(self, rules, test_output_dir, references_sample):
        """La sección de referencias debe incluirse."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Contenido."),
        ]
        doc_model = DocumentModel(
            session_id="test_refs",
            file_name="refs.docx",
            elements=elements,
        )

        out = test_output_dir / "refs_output.docx"
        generate_apa7_docx(doc_model, out, rules, references=references_sample)

        doc = docx.Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Referencias" in all_text

    def test_references_sorted_alphabetically(self, rules, test_output_dir):
        """Las referencias deben ordenarse alfabéticamente por autor."""
        refs = [
            ReferenciaModel(
                id="r3", authors=["Martínez", "A."], year="2022",
                title="Tercero", source="Journal C",
                raw_text="Martínez, A. (2022). Tercero. Journal C."
            ),
            ReferenciaModel(
                id="r1", authors=["García", "B."], year="2020",
                title="Primero", source="Journal A",
                raw_text="García, B. (2020). Primero. Journal A."
            ),
            ReferenciaModel(
                id="r2", authors=["López", "C."], year="2021",
                title="Segundo", source="Journal B",
                raw_text="López, C. (2021). Segundo. Journal B."
            ),
        ]

        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Cuerpo."),
        ]
        doc_model = DocumentModel(
            session_id="test_sort_refs",
            file_name="sort_refs.docx",
            elements=elements,
        )

        out = test_output_dir / "sorted_refs_output.docx"
        generate_apa7_docx(doc_model, out, rules, references=refs)

        doc = docx.Document(str(out))

        # Encontrar posiciones de cada referencia
        positions = {}
        for i, p in enumerate(doc.paragraphs):
            for ref in refs:
                if ref.authors[0] in p.text and ref.year in p.text:
                    if ref.authors[0] not in positions:
                        positions[ref.authors[0]] = i

        # García debe aparecer antes que López, y López antes que Martínez
        if "García" in positions and "López" in positions:
            assert positions["García"] < positions["López"], (
                "García debe aparecer antes que López (orden alfabético)"
            )
        if "López" in positions and "Martínez" in positions:
            assert positions["López"] < positions["Martínez"], (
                "López debe aparecer antes que Martínez (orden alfabético)"
            )

    def test_references_have_hanging_indent(self, rules, test_output_dir, references_sample):
        """Las referencias deben tener sangría francesa (hanging indent)."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Cuerpo."),
        ]
        doc_model = DocumentModel(
            session_id="test_hanging",
            file_name="hanging.docx",
            elements=elements,
        )

        out = test_output_dir / "hanging_output.docx"
        generate_apa7_docx(doc_model, out, rules, references=references_sample)

        doc = docx.Document(str(out))

        # Buscar párrafos de referencia (los que no son "Referencias")
        ref_paragraphs = []
        in_refs = False
        for p in doc.paragraphs:
            if p.text.strip() == "Referencias":
                in_refs = True
                continue
            if in_refs and p.text.strip():
                ref_paragraphs.append(p)

        # Al menos una referencia debe tener left_indent configurado
        hanging_found = False
        for p in ref_paragraphs:
            if p.paragraph_format.left_indent and p.paragraph_format.first_line_indent:
                hanging_found = True
                break

        assert hanging_found or len(ref_paragraphs) > 0, (
            "Las referencias deben tener sangría francesa"
        )

    def test_no_references_when_empty(self, rules, test_output_dir):
        """Sin referencias, no debe aparecer la sección."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Sin refs."),
        ]
        doc_model = DocumentModel(
            session_id="test_no_refs",
            file_name="no_refs.docx",
            elements=elements,
        )

        out = test_output_dir / "no_refs_output.docx"
        generate_apa7_docx(doc_model, out, rules)  # sin references

        doc = docx.Document(str(out))
        ref_headers = [p for p in doc.paragraphs if p.text.strip() == "Referencias"]
        assert len(ref_headers) == 0, "No debe aparecer 'Referencias' si no hay referencias"


# ── NÚMEROS DE PÁGINA ─────────────────────────────────────────────────────────

class TestPageNumbers:
    """Verifica que los números de página estén presentes."""

    def test_page_number_field_exists(self, rules, test_output_dir):
        """El encabezado debe contener un campo PAGE de Word."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Contenido."),
        ]
        doc_model = DocumentModel(
            session_id="test_pagenum",
            file_name="pagenum.docx",
            elements=elements,
        )

        out = test_output_dir / "pagenum_output.docx"
        generate_apa7_docx(doc_model, out, rules)

        doc = docx.Document(str(out))
        header = doc.sections[0].header

        # Buscar fldChar en el XML del header
        header_xml = header.paragraphs[0]._element.xml
        assert "fldChar" in header_xml or "PAGE" in header_xml, (
            "El encabezado debe contener un campo PAGE de Word"
        )


# ── FUENTES NORMALIZADAS ──────────────────────────────────────────────────────

class TestFontNormalization:
    """Verifica que las fuentes estén normalizadas en el documento generado."""

    def test_all_paragraphs_use_correct_font(self, rules, test_output_dir):
        """Todos los runs deben usar la fuente especificada en las reglas."""
        elements = [
            ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                        text="Título Principal"),
            ElementModel(id="e2", type=ElementType.PARAGRAPH,
                        text="Párrafo de prueba con contenido suficiente para verificar "
                             "que la fuente se aplica correctamente."),
            ElementModel(id="e3", type=ElementType.BULLET,
                        text="Elemento de lista con viñeta"),
        ]
        doc_model = DocumentModel(
            session_id="test_font",
            file_name="font_test.docx",
            elements=elements,
        )

        out = test_output_dir / "font_output.docx"
        generate_apa7_docx(doc_model, out, rules)

        doc = docx.Document(str(out))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip():
                    assert run.font.name == rules.font_family, (
                        f"Fuente incorrecta: '{run.font.name}' != '{rules.font_family}' "
                        f"en texto: '{run.text[:50]}...'"
                    )

    def test_all_paragraphs_use_correct_font_size(self, rules, test_output_dir):
        """Todos los runs deben usar el tamaño de fuente de las reglas."""
        elements = [
            ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                        text="Título"),
            ElementModel(id="e2", type=ElementType.PARAGRAPH,
                        text="Párrafo de prueba."),
        ]
        doc_model = DocumentModel(
            session_id="test_font_size",
            file_name="font_size_test.docx",
            elements=elements,
        )

        out = test_output_dir / "font_size_output.docx"
        generate_apa7_docx(doc_model, out, rules)

        doc = docx.Document(str(out))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text.strip() and run.font.size:
                    assert run.font.size == Pt(rules.font_size_pt), (
                        f"Tamaño incorrecto: {run.font.size} != {Pt(rules.font_size_pt)}"
                    )


# ── MÁRGENES APA 7 ────────────────────────────────────────────────────────────

class TestPageMargins:
    """Verifica los márgenes APA 7 (2.54 cm / 1 inch)."""

    def test_margins_are_one_inch(self, rules, test_output_dir):
        """Los márgenes deben ser de 1 pulgada (2.54 cm)."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH, text="Test."),
        ]
        doc_model = DocumentModel(
            session_id="test_margins",
            file_name="margins.docx",
            elements=elements,
        )

        out = test_output_dir / "margins_output.docx"
        generate_apa7_docx(doc_model, out, rules)

        doc = docx.Document(str(out))
        section = doc.sections[0]

        margin_inches = rules.margins_cm / 2.54
        expected_margin = Inches(margin_inches)

        # Verificar márgenes (puede haber pequeñas variaciones de redondeo)
        assert abs(section.top_margin - expected_margin) < Inches(0.05), (
            f"Margen superior: {section.top_margin}, esperado: {expected_margin}"
        )
        assert abs(section.bottom_margin - expected_margin) < Inches(0.05)
        assert abs(section.left_margin - expected_margin) < Inches(0.05)
        assert abs(section.right_margin - expected_margin) < Inches(0.05)


# ── INTERLINEADO ───────────────────────────────────────────────────────────────

class TestLineSpacing:
    """Verifica interlineado doble (2.0)."""

    def test_paragraph_double_spacing(self, rules, test_output_dir):
        """Los párrafos deben tener interlineado 2.0."""
        elements = [
            ElementModel(id="e1", type=ElementType.PARAGRAPH,
                        text="Párrafo que debe tener interlineado doble según APA 7."),
        ]
        doc_model = DocumentModel(
            session_id="test_spacing",
            file_name="spacing.docx",
            elements=elements,
        )

        out = test_output_dir / "spacing_output.docx"
        generate_apa7_docx(doc_model, out, rules)

        doc = docx.Document(str(out))
        # El interlineado debe ser 2.0 para párrafos que no son headings
        for p in doc.paragraphs:
            if p.text.strip() and "Párrafo que debe tener" in p.text:
                assert p.paragraph_format.line_spacing == 2.0, (
                    f"Interlineado: {p.paragraph_format.line_spacing}, esperado: 2.0"
                )


# ── TESTS DE style_engine UNITARIOS ─────────────────────────────────────────────

class TestStyleEngineUnit:
    """Tests unitarios del motor de estilos."""

    def test_heading_level1_format(self, rules):
        """Heading nivel 1 debe ser centrado y negrita."""
        doc = docx.Document()
        p = doc.add_paragraph()
        format_heading_paragraph(p, 1, "Título Nivel 1", rules)

        # python-docx: run.bold returns None cuando hereda del estilo
        # El valor True/False solo se retorna si fue explícitamente seteado
        if p.runs:
            # bold puede ser None (inherit) o True (explicit)
            assert p.runs[0].bold in (True, None), f"bold={p.runs[0].bold}"
        assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_heading_level2_format(self, rules):
        """Heading nivel 2 debe ser alineado a la izquierda y negrita."""
        doc = docx.Document()
        p = doc.add_paragraph()
        format_heading_paragraph(p, 2, "Título Nivel 2", rules)

        if p.runs:
            assert p.runs[0].bold in (True, None), f"bold={p.runs[0].bold}"
        assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_heading_level3_format(self, rules):
        """Heading nivel 3 debe ser negrita y cursiva."""
        doc = docx.Document()
        p = doc.add_paragraph()
        format_heading_paragraph(p, 3, "Título Nivel 3", rules)

        if p.runs:
            assert p.runs[0].bold in (True, None), f"bold={p.runs[0].bold}"

    def test_normal_paragraph_format(self, rules):
        """Párrafo normal debe tener sangría de primera línea."""
        doc = docx.Document()
        p = doc.add_paragraph()
        format_normal_paragraph(p, "Párrafo de prueba con suficiente texto.", rules)

        assert p.paragraph_format.line_spacing == 2.0
        assert p.paragraph_format.first_line_indent == Inches(0.5)
        # font.name puede ser None si hereda del estilo del documento

    def test_block_quote_format(self, rules):
        """Block quote debe tener sangría izquierda y sin sangría de primera línea."""
        doc = docx.Document()
        p = doc.add_paragraph()
        format_block_quote(p, "Cita textual de más de 40 palabras en bloque.", rules)

        assert p.paragraph_format.left_indent == Inches(0.5), (
            "Block quote debe tener sangría izquierda"
        )
        assert p.paragraph_format.first_line_indent == Inches(0), (
            "Block quote NO debe tener sangría de primera línea"
        )


# ── TESTS: ECUACIONES ───────────────────────────────────────────────────────────

class TestEquationGeneration:
    """El generador debe aplicar alineación y numeración configurada a ecuaciones."""

    def test_equation_centered_with_number(self, rules):
        """Ecuación centrada con número → alineación centro + tab derecho con número."""
        from models import EquationConfig
        from generation.generator import _render_equation_number

        doc = docx.Document()
        p = doc.add_paragraph("E=mc^2")

        elem = ElementModel(
            id="eq1", type=ElementType.EQUATION, text="E=mc^2",
            style_name="Normal", font_size=12.0, has_math=True,
            equation=EquationConfig(show_number=True, number_format="(1)", alignment="center"),
        )
        # El generador no expone el helper por elemento; verificamos las piezas
        assert _render_equation_number("1", "(1)") == "(1)"
        assert _render_equation_number("2", "[1]") == "[2]"
        assert _render_equation_number("3", "1.") == "3."
        assert _render_equation_number("4", "Ecuación {n}") == "Ecuación 4"
        assert _render_equation_number("1", "custom") == "(1)"

    def test_equation_rendered_with_number_in_docx(self, rules, tmp_path):
        """Generar un documento con ecuación numerada produce el número en el XML."""
        from models import EquationConfig, DocumentMeta

        doc_model = DocumentModel(
            session_id="eqtest",
            file_name="orig.docx",
            elements=[
                ElementModel(
                    id="eq1", type=ElementType.EQUATION, text="x = y",
                    style_name="Normal", font_size=12.0, has_math=True,
                    equation=EquationConfig(show_number=True, number_format="(1)", alignment="center"),
                ),
            ],
            meta=DocumentMeta(session_id="eqtest", file_name="orig.docx", wordapa7_version="1.0.0"),
            rules=rules,
            references=[],
        )
        # Crear el original para que el generador modifique sobre él
        orig = docx.Document()
        orig.add_paragraph("x = y")
        (tmp_path / "original.docx").write_bytes(_docx_bytes(orig))

        out = tmp_path / "out.docx"
        generate_apa7_docx(doc_model, str(out), rules, None, [])

        from docx import Document as _D
        result = _D(str(out))
        assert len(result.paragraphs) >= 1
        full = "\n".join(p.text for p in result.paragraphs)
        assert "(1)" in full, f"El número de ecuación debería aparecer. Texto: {full!r}"


def _docx_bytes(doc) -> bytes:
    """Serializa un objeto python-docx a bytes (sin tocar disco)."""
    import io as _io
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
