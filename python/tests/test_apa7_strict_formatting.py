"""
WordAPA7 — Test de Verificación de Reglas Estrictas APA 7 (Parte 5)

Verifica:
1. Interlineado Doble (w:line=480) global en el 100% de párrafos del cuerpo.
2. Títulos Nivel 1 centrados (alignment = CENTER) y a 12pt.
3. Celdas de la primera fila de tabla (encabezados) con negrita (bold = True).
"""

from pathlib import Path
import docx

from models import DocumentModel, ElementModel, ElementType, APARuleSet, TableModel
from generation.generator import generate_apa7_docx


def test_apa7_strict_body_double_spacing_and_heading1(tmp_path):
    orig_path = tmp_path / "original.docx"
    out_path = tmp_path / "output_strict.docx"

    doc = docx.Document()
    p1 = doc.add_paragraph("Introducción")
    p2 = doc.add_paragraph("Este es un párrafo del cuerpo del trabajo académico en APA 7.")
    doc.save(orig_path)

    elements = [
        ElementModel(id="h1", type=ElementType.HEADING, heading_level=1, text="Introducción", is_cover_section=False),
        ElementModel(id="p1", type=ElementType.PARAGRAPH, text="Este es un párrafo del cuerpo del trabajo académico en APA 7.", is_cover_section=False),
    ]

    doc_model = DocumentModel(session_id="strict_test", file_name="original.docx", elements=elements)
    rules = APARuleSet(font_family="Times New Roman", font_size_pt=12, line_spacing=2.0)

    generate_apa7_docx(doc_model, out_path, rules)

    gen_doc = docx.Document(out_path)

    print("[DEBUG paragraphs]:", [(i, p.text, [r.bold for r in p.runs]) for i, p in enumerate(gen_doc.paragraphs)])

    # 1. Heading 1 centrado
    h1_para = gen_doc.paragraphs[0]
    assert h1_para.text == "Introducción"
    assert h1_para.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    assert any(r.bold for r in h1_para.runs) is True

    # 2. Párrafo con interlineado doble (2.0)
    body_para = gen_doc.paragraphs[1]
    assert body_para.paragraph_format.line_spacing == 2.0


def test_apa7_strict_table_header_bold(tmp_path):
    orig_path = tmp_path / "original.docx"
    out_path = tmp_path / "output_table.docx"

    doc = docx.Document()
    doc.add_paragraph("Introducción")
    doc.save(orig_path)

    tbl_data = TableModel(
        element_id="t1",
        table_number=1,
        caption="Estadísticas de Rendimiento",
        headers=["Módulo", "Estudiantes", "Aprobados"],
        rows=[["Matemáticas", "50", "45"], ["Física", "40", "38"]],
    )

    elements = [
        ElementModel(id="h1", type=ElementType.HEADING, heading_level=1, text="Introducción", is_cover_section=False),
        ElementModel(id="t1", type=ElementType.TABLE, table_info=tbl_data, is_cover_section=False),
    ]

    doc_model = DocumentModel(session_id="table_test", file_name="original.docx", elements=elements)
    rules = APARuleSet()

    generate_apa7_docx(doc_model, out_path, rules)

    gen_doc = docx.Document(out_path)
    assert len(gen_doc.tables) > 0

    table = gen_doc.tables[0]
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        assert any(r.bold for r in cell.paragraphs[0].runs) is True
