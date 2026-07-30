"""
WordAPA7 — Test de Integración Real de Exportación DOCX
Verifica directamente sobre la salida real generada por la API:
1. No duplicación de <w:tbl> ni <w:drawing>
2. Presencia de w:line="480" y color="000000" en styles.xml
3. Etiquetas 'Tabla 1', 'Tabla 2' e 'Figura 1' pegadas inmediatamente a sus nodos XML
4. Ausencia de fugas de tokens internos (author_card, etc.)
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import docx
import zipfile
import re
import io
import pytest

from models import DocumentModel, ElementModel, ElementType, APARuleSet, PortadaData, TableModel, ImageModel
from generation.generator import generate_apa7_docx


def create_real_sample_docx(orig_path: Path):
    """Crea un documento .docx con 2 tablas, 1 imagen y 3 párrafos para pruebas reales."""
    doc = docx.Document()
    doc.add_heading("Título Principal del Documento", level=1)
    doc.add_paragraph("Este es un párrafo de introducción sobre el estudio academico.")
    
    # Tabla 1
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "Material"
    t1.rows[0].cells[1].text = "Cantidad"
    t1.rows[1].cells[0].text = "Acero"
    t1.rows[1].cells[1].text = "10 kg"

    doc.add_heading("Sección de Métodos", level=2)
    doc.add_paragraph("Descripción de la metodología aplicada en el proceso.")

    # Tabla 2
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "Etapa"
    t2.rows[0].cells[1].text = "Tiempo"
    t2.rows[1].cells[0].text = "Mezclado"
    t2.rows[1].cells[1].text = "5 min"

    doc.save(orig_path)


def test_real_pipeline_export_integrity(tmp_path):
    orig_path = tmp_path / "original.docx"
    out_path = tmp_path / "APA7_output.docx"

    create_real_sample_docx(orig_path)

    # Contar nodos XML del original
    with zipfile.ZipFile(orig_path, "r") as z:
        xml_orig = z.read("word/document.xml").decode("utf-8")
        tbl_orig_count = len(re.findall(r'<w:tbl\b[^>]*>', xml_orig))

    # Parsear y generar
    from parsing.docx_parser import parse_docx_bytes
    orig_bytes = orig_path.read_bytes()
    doc_model = parse_docx_bytes(orig_bytes, "original.docx", "test_session_123", tmp_path)

    generate_apa7_docx(
        doc_model,
        out_path,
        rules=APARuleSet(),
        portada=PortadaData(use_original_cover=True),
    )

    # 1. Inspección del ZIP generado real
    with zipfile.ZipFile(out_path, "r") as z:
        xml_out = z.read("word/document.xml").decode("utf-8")
        styles_xml_out = z.read("word/styles.xml").decode("utf-8")
        tbl_gen_count = len(re.findall(r'<w:tbl\b[^>]*>', xml_out))

    # Aserción 1: No duplicación de tablas
    assert tbl_gen_count == tbl_orig_count, f"Duplicación detectada: {tbl_orig_count} orig vs {tbl_gen_count} gen"

    # Aserción 2: Reescritura de styles.xml con interlineado doble (480) y color negro (000000)
    assert 'w:line="480"' in styles_xml_out, "styles.xml no contiene w:line='480'"
    assert 'w:color w:val="000000"' in styles_xml_out or 'w:color w:val="000000"' in styles_xml_out, "styles.xml no forzó color negro"

    # Aserción 3: Orden de etiquetas "Tabla 1" y "Tabla 2" pegadas inmediatamente a sus tablas
    doc_gen = docx.Document(out_path)
    for idx, table in enumerate(doc_gen.tables):
        prev_elem = table._tbl.getprevious()
        assert prev_elem is not None and prev_elem.tag.endswith("p"), f"La tabla {idx+1} no tiene párrafo inmediatamente anterior"
        prev_p = docx.text.paragraph.Paragraph(prev_elem, doc_gen)
        expected_label = f"Tabla {idx+1}"
        assert expected_label in prev_p.text, f"La tabla {idx+1} no tiene la etiqueta esperada '{expected_label}'. Encontrado: '{prev_p.text}'"

    # Aserción 4: No fugas de tokens internos
    internal_leaks = ["author_card", "vertical_line", "shape_group", "shape_textbox"]
    for leak in internal_leaks:
        assert leak not in xml_out, f"Fuga de token interno '{leak}' detectada en document.xml"

    print("\n[OK] TEST DE INTEGRACIÓN REAL PASADO EXITOSAMENTE:")
    print(f" - Tablas original: {tbl_orig_count} | Tablas generado: {tbl_gen_count} (Cero duplicación)")
    print(" - styles.xml reescrito con interlineado 480 y color 000000")
    print(" - Etiquetas 'Tabla 1' y 'Tabla 2' pegadas en orden consecutivo exacto a sus tablas")
    print(" - Cero fugas de tokens de código en el texto visible")


if __name__ == "__main__":
    import tempfile
    with tempfile.TemporaryDirectory() as tmpdir:
        test_real_pipeline_export_integrity(Path(tmpdir))
