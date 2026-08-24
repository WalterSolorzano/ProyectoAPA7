"""Garantía contractual del motor por alcances: solo se toca lo pedido."""

import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import docx  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from modules.scoped_apply import apply_scopes  # noqa: E402


def _build_doc() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Título del Trabajo")
    doc.add_paragraph("Br. Autor Uno | Carnet: 2021-0251")
    doc.add_heading("Introducción", level=1)
    doc.add_paragraph("Párrafo del cuerpo que no debe cambiar.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(1, 1).text = "B"
    doc.add_paragraph("Autor, A. (2019). Libro de prueba. Editorial X.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _body_texts(data: bytes):
    d = docx.Document(io.BytesIO(data))
    return [p.text for p in d.paragraphs]


RULES = {"font_family": "Times New Roman", "font_size_pt": 12, "line_spacing": 2.0}


def test_tablas_scope_no_toca_texto_ni_portada():
    original = _build_doc()
    out, summary = apply_scopes(original, ["tablas_imagenes"], RULES)
    before = _body_texts(original)
    after = _body_texts(out)
    # Los párrafos ORIGINALES conservan su texto exacto (solo se insertan captions).
    for b in before:
        assert b in after, f"texto original alterado: {b!r}"
    assert any(t.startswith("Tabla 1") for t in after)
    assert summary["tablas"] == 1


def test_bibliografia_scope_no_toca_cuerpo():
    original = _build_doc()
    out, _ = apply_scopes(original, ["bibliografia"], RULES)
    before = _body_texts(original)
    after = _body_texts(out)
    # Cuerpo y portada idénticos
    assert before[:4] == after[:4]
    # La entrada de referencia existe igual
    assert "Libro de prueba" in " ".join(after)


def test_texto_scope_no_borra_contenido():
    original = _build_doc()
    out, _ = apply_scopes(original, ["texto"], RULES)
    before = _body_texts(original)
    after = _body_texts(out)
    assert before == after
