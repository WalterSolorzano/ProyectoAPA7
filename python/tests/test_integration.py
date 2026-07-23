"""
WordAPA7 — Tests de Integración (API End-to-End)

Verifica el flujo completo desde upload hasta download:
1. Health check del servidor
2. Upload de archivo .docx
3. Clasificación de elementos
4. Validación APA
5. Generación de documento APA 7
6. Descarga y validación del archivo generado

Requiere que el servidor FastAPI esté corriendo en http://localhost:8742.
Ejecutar con: pytest python/tests/test_integration.py -v -m integration
"""

import sys
import io
import json
import urllib.request
import urllib.error
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import pytest
import docx
from docx.shared import Pt


# Marcador para tests de integración (requieren servidor corriendo)
pytestmark = pytest.mark.integration


# ── HELPERS ─────────────────────────────────────────────────────────────────────

BASE_URL = "http://127.0.0.1:8742"


def server_is_running():
    """Verifica si el servidor FastAPI está corriendo."""
    try:
        req = urllib.request.Request(f"{BASE_URL}/api/health")
        with urllib.request.urlopen(req, timeout=3) as resp:
            return resp.status == 200
    except Exception:
        return False


def create_test_docx_bytes():
    """Crea un documento .docx de prueba en memoria."""
    doc = docx.Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = docx.shared.Cm(2.54)

    # Heading 1
    h1 = doc.add_heading("Estudio sobre Inteligencia Artificial en Educación", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    # Párrafo con cita
    p1 = doc.add_paragraph(
        "La inteligencia artificial ha transformado radicalmente el panorama "
        "educativo en la última década (García, 2023, p. 45). Los sistemas "
        "de tutoría inteligente han demostrado ser particularmente efectivos "
        "en la personalización del aprendizaje."
    )
    for run in p1.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Heading 2
    h2 = doc.add_heading("Metodología de Investigación", level=2)
    for run in h2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)

    # Párrafo con cita narrativa
    p2 = doc.add_paragraph(
        "López et al. (2021) encontraron que los estudiantes que utilizan "
        "plataformas de aprendizaje adaptativo mejoran su rendimiento en un "
        "promedio de 23% comparado con métodos tradicionales."
    )
    for run in p2.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # Guardar a bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def multipart_post(url, file_bytes, filename="test.docx"):
    """Hace un POST multipart/form-data a la URL especificada."""
    boundary = '----WebKitFormBoundary7MA4YWxkTrZu0gW'
    body = (
        f'--{boundary}\r\n'
        f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
        f'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n'
    ).encode('utf-8') + file_bytes + f'\r\n--{boundary}--\r\n'.encode('utf-8')

    req = urllib.request.Request(
        url,
        data=body,
        headers={'Content-Type': f'multipart/form-data; boundary={boundary}'}
    )
    return req


# ── FIXTURE: SKIP SI EL SERVIDOR NO ESTÁ CORRIENDO ─────────────────────────────

@pytest.fixture(scope="module")
def server_available():
    """Skip todos los tests de integración si el servidor no está corriendo."""
    if not server_is_running():
        pytest.skip(
            "Servidor FastAPI no está corriendo en http://localhost:8742. "
            "Inícialo con: cd python && python main.py"
        )
    return True


# ── TESTS ──────────────────────────────────────────────────────────────────────

class TestHealthCheck:
    """Verifica que el servidor responda correctamente."""

    def test_health_endpoint(self, server_available):
        """GET /api/health debe retornar status ok."""
        req = urllib.request.Request(f"{BASE_URL}/api/health")
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            assert data["status"] == "ok"
            assert data["app"] == "WordAPA7"


class TestUploadFlow:
    """Verifica el flujo de upload y parseo."""

    def test_upload_docx(self, server_available):
        """POST /api/upload debe retornar el modelo de documento parseado."""
        file_bytes = create_test_docx_bytes()
        req = multipart_post(
            f"{BASE_URL}/api/upload",
            file_bytes,
            "integration_test.docx"
        )

        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode('utf-8'))

        assert "session_id" in data, f"No hay session_id en respuesta: {list(data.keys())}"
        assert len(data["session_id"]) > 0
        assert "elements" in data
        assert len(data["elements"]) > 0

        # Verificar estructura de elementos
        for elem in data["elements"]:
            assert "id" in elem
            assert "type" in elem
            assert "text" in elem or elem["type"] in ("image", "table")

        return data["session_id"]


class TestSessionFlow:
    """Verifica el acceso a sesiones."""

    def test_get_session(self, server_available):
        """GET /api/session/{id} debe retornar la sesión."""
        # Primero crear una sesión
        file_bytes = create_test_docx_bytes()
        req = multipart_post(f"{BASE_URL}/api/upload", file_bytes, "session_test.docx")
        with urllib.request.urlopen(req, timeout=15) as resp:
            upload_data = json.loads(resp.read().decode('utf-8'))

        session_id = upload_data["session_id"]

        # Luego obtener la sesión
        req2 = urllib.request.Request(f"{BASE_URL}/api/session/{session_id}")
        with urllib.request.urlopen(req2, timeout=5) as resp:
            session_data = json.loads(resp.read().decode('utf-8'))

        assert session_data["session_id"] == session_id
        assert len(session_data["elements"]) == len(upload_data["elements"])

    def test_nonexistent_session_returns_404(self, server_available):
        """Solicitar una sesión inexistente debe retornar 404."""
        req = urllib.request.Request(f"{BASE_URL}/api/session/nonexistent_12345")
        try:
            with urllib.request.urlopen(req, timeout=5) as resp:
                pytest.fail("Debería haber retornado 404")
        except urllib.error.HTTPError as e:
            assert e.code == 404


class TestGenerateAndDownload:
    """Verifica generación y descarga de documento APA 7."""

    def test_generate_and_download(self, server_available):
        """POST /api/generate + GET /api/download/{id}."""
        # Crear sesión
        file_bytes = create_test_docx_bytes()
        req = multipart_post(f"{BASE_URL}/api/upload", file_bytes, "gen_test.docx")
        with urllib.request.urlopen(req, timeout=15) as resp:
            upload_data = json.loads(resp.read().decode('utf-8'))

        session_id = upload_data["session_id"]

        # Generar
        gen_body = json.dumps({
            "session_id": session_id,
            "portada": {
                "title": "Estudio sobre IA en Educación",
                "author": "Juan Pérez García",
                "institution": "Universidad Nacional Autónoma de México",
                "course": "Seminario de Investigación Avanzada",
                "instructor": "Dra. María López Castillo",
                "date": "2025"
            }
        }).encode('utf-8')

        gen_req = urllib.request.Request(
            f"{BASE_URL}/api/generate",
            data=gen_body,
            headers={'Content-Type': 'application/json'}
        )

        with urllib.request.urlopen(gen_req, timeout=30) as resp:
            gen_data = json.loads(resp.read().decode('utf-8'))

        assert "download_url" in gen_data
        assert "filename" in gen_data

        # Descargar
        dl_url = f"{BASE_URL}{gen_data['download_url']}"
        dl_req = urllib.request.Request(dl_url)
        with urllib.request.urlopen(dl_req, timeout=15) as resp:
            downloaded_bytes = resp.read()

        assert len(downloaded_bytes) > 1000, (
            f"Archivo descargado muy pequeño: {len(downloaded_bytes)} bytes"
        )

        # Validar que es un .docx válido
        import zipfile
        buf = io.BytesIO(downloaded_bytes)
        assert zipfile.is_zipfile(buf), "El archivo descargado no es un ZIP válido"

        # Reabrir con python-docx
        doc = docx.Document(buf)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "inteligencia artificial" in all_text.lower()


class TestValidateEndpoint:
    """Verifica el endpoint de validación APA."""

    def test_validate_document(self, server_available):
        """POST /api/validate debe retornar issues de validación."""
        # Crear sesión
        file_bytes = create_test_docx_bytes()
        req = multipart_post(f"{BASE_URL}/api/upload", file_bytes, "validate_test.docx")
        with urllib.request.urlopen(req, timeout=15) as resp:
            upload_data = json.loads(resp.read().decode('utf-8'))

        session_id = upload_data["session_id"]

        # Validar
        validate_body = json.dumps({
            "session_id": session_id,
            "references": [
                {
                    "id": "ref_1",
                    "authors": ["García", "M."],
                    "year": "2023",
                    "title": "Inteligencia artificial en educación",
                    "source": "Revista de Educación",
                    "raw_text": "García, M. (2023). Inteligencia artificial en educación. Revista de Educación, 45(2), 123-145."
                }
            ]
        }).encode('utf-8')

        val_req = urllib.request.Request(
            f"{BASE_URL}/api/validate",
            data=validate_body,
            headers={'Content-Type': 'application/json'}
        )

        with urllib.request.urlopen(val_req, timeout=10) as resp:
            val_data = json.loads(resp.read().decode('utf-8'))

        assert "issues" in val_data
        assert isinstance(val_data["issues"], list)
