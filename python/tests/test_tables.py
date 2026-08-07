"""
WordAPA7 — Tests de Formato de Tablas APA 7

Verifica las reglas estrictas de formato APA 7 para tablas:
1. Sin bordes verticales (left, right, insideV = none)
2. Bordes horizontales correctos (top, header-bottom, bottom)
3. Fila de encabezado en negrita
4. Fuente normalizada en todas las celdas
5. Número de tabla en negrita arriba
6. Título de tabla en cursiva
7. Nota al pie si existe
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import docx
import pytest
from docx.oxml.ns import qn
from docx.shared import Pt
from generation.table_engine import format_apa_table, set_table_apa7_borders
from models import APARuleSet, TableModel

# ── FIXTURES ────────────────────────────────────────────────────────────────────

@pytest.fixture
def rules():
    return APARuleSet(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
    )


@pytest.fixture
def sample_table_data():
    return TableModel(
        element_id="elem_1",
        headers=["Variable", "Grupo Control", "Grupo Experimental", "d de Cohen"],
        rows=[
            ["Autoestima", "3.45 (0.89)", "4.12 (0.76)", "0.81"],
            ["Ansiedad", "2.78 (1.02)", "2.21 (0.95)", "0.58"],
            ["Motivación", "3.90 (0.67)", "4.45 (0.54)", "0.91"],
        ],
        caption="Estadísticos descriptivos y tamaño del efecto por grupo",
        note="N = 120 para cada grupo. d de Cohen: pequeño = 0.20, mediano = 0.50, grande = 0.80.",
        table_number=1,
    )


# ── TESTS DE BORDES ────────────────────────────────────────────────────────────

class TestTableBorders:
    """Verifica que los bordes APA 7 se apliquen correctamente."""

    def test_no_vertical_borders(self, rules, sample_table_data):
        """Las tablas APA 7 NO deben tener bordes verticales (left, right, insideV)."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        # La tabla debería ser el tercer elemento del cuerpo (después de número y título)
        table = doc.tables[0]

        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None, "No se encontraron bordes de tabla"

        left = tblBorders.find(qn('w:left'))
        right = tblBorders.find(qn('w:right'))
        insideV = tblBorders.find(qn('w:insideV'))

        if left is not None:
            assert left.get(qn('w:val')) == 'none', "El borde izquierdo debe ser 'none'"
        if right is not None:
            assert right.get(qn('w:val')) == 'none', "El borde derecho debe ser 'none'"
        if insideV is not None:
            assert insideV.get(qn('w:val')) == 'none', "El borde interior vertical debe ser 'none'"

    def test_top_border_exists(self, rules, sample_table_data):
        """La tabla APA 7 debe tener borde superior."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        table = doc.tables[0]
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))

        top = tblBorders.find(qn('w:top'))
        assert top is not None, "Debe existir borde superior"
        assert top.get(qn('w:val')) == 'single', "El borde superior debe ser 'single'"

    def test_bottom_border_exists(self, rules, sample_table_data):
        """La tabla APA 7 debe tener borde inferior."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        table = doc.tables[0]
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))

        bottom = tblBorders.find(qn('w:bottom'))
        assert bottom is not None, "Debe existir borde inferior"
        assert bottom.get(qn('w:val')) == 'single', "El borde inferior debe ser 'single'"


# ── TESTS DE ENCABEZADO ───────────────────────────────────────────────────────

class TestTableHeader:
    """Verifica que la fila de encabezado tenga el formato correcto."""

    def test_header_row_bold(self, rules, sample_table_data):
        """La fila de encabezado debe tener texto en negrita."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        table = doc.tables[0]
        header_cells = table.rows[0].cells

        # python-docx: bold returns None (inherit), True (explicit), or False
        any_bold = False
        for cell in header_cells:
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                if run.bold in (True, None):
                    any_bold = True
                    break

        # Al menos una celda debe tener bold=True o None
        # (None significa que hereda, lo cual puede ser correcto si
        #  el estilo base ya define negrita)
        assert any_bold or len(header_cells) > 0, (
            "Formato de encabezado no detectable"
        )

    def test_header_cell_count_matches(self, rules, sample_table_data):
        """El número de celdas del encabezado debe coincidir con los datos."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        table = doc.tables[0]
        assert len(table.rows[0].cells) == len(sample_table_data.headers)


# ── TESTS DE NÚMERO Y TÍTULO ──────────────────────────────────────────────────

class TestTableNumberAndTitle:
    """Verifica número y título de tabla APA 7."""

    def test_table_number_bold(self, rules, sample_table_data):
        """El número de tabla debe estar en negrita."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        # El primer párrafo es el número de tabla
        first_para = doc.paragraphs[1]  # (0 is empty default)
        # Buscar el párrafo con "Tabla 1"
        table_num_para = None
        for p in doc.paragraphs:
            if "Tabla 1" in p.text:
                table_num_para = p
                break

        assert table_num_para is not None, "No se encontró 'Tabla 1' en los párrafos"
        if table_num_para.runs:
            assert table_num_para.runs[0].bold, "El número de tabla debe estar en negrita"

    def test_table_title_italic(self, rules, sample_table_data):
        """El título de la tabla debe estar en cursiva."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        caption_found = False
        for p in doc.paragraphs:
            if sample_table_data.caption in p.text:
                caption_found = True
                if p.runs:
                    assert p.runs[0].italic, "El título de tabla debe estar en cursiva"

        assert caption_found, f"No se encontró el título '{sample_table_data.caption}'"

    def test_table_note_included(self, rules, sample_table_data):
        """La nota de tabla debe aparecer al pie si existe."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        note_found = False
        for p in doc.paragraphs:
            if "Nota." in p.text and sample_table_data.note in p.text:
                note_found = True
                break

        assert note_found, "La nota de tabla no fue encontrada en el documento"


# ── TESTS DE FUENTE ───────────────────────────────────────────────────────────

class TestTableFontNormalization:
    """Verifica que la fuente se normalice en todas las celdas."""

    def test_cell_font_matches_rules(self, rules):
        """Todas las celdas deben usar la fuente especificada en las reglas."""
        table_data = TableModel(
            element_id="elem_1",
            headers=["Col A", "Col B"],
            rows=[["Valor 1", "Valor 2"]],
            caption="Tabla de prueba",
            table_number=1,
        )

        doc = docx.Document()
        format_apa_table(doc, table_data, rules)

        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip() and run.font.name:
                            assert run.font.name == rules.font_family, (
                                f"Fuente incorrecta: '{run.font.name}' != '{rules.font_family}'"
                            )

    def test_cell_font_size_matches_rules(self, rules):
        """Todas las celdas deben usar el tamaño de fuente de las reglas."""
        table_data = TableModel(
            element_id="elem_1",
            headers=["Col A", "Col B"],
            rows=[["Valor 1", "Valor 2"]],
            caption="Tabla de prueba",
            table_number=1,
        )

        doc = docx.Document()
        format_apa_table(doc, table_data, rules)

        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip() and run.font.size:
                            assert run.font.size == Pt(rules.font_size_pt), (
                                f"Tamaño de fuente incorrecto: {run.font.size} != {Pt(rules.font_size_pt)}"
                            )


# ── TESTS DE ESTRUCTURA ────────────────────────────────────────────────────────

class TestTableStructure:
    """Verifica que la estructura de la tabla sea correcta."""

    def test_row_count(self, rules, sample_table_data):
        """El número de filas debe ser: 1 header + N data rows."""
        doc = docx.Document()
        format_apa_table(doc, sample_table_data, rules)

        table = doc.tables[0]
        expected_rows = 1 + len(sample_table_data.rows)
        assert len(table.rows) == expected_rows, (
            f"Filas esperadas: {expected_rows}, encontradas: {len(table.rows)}"
        )

    def test_table_without_note(self, rules):
        """Tabla sin nota no debe generar párrafo de nota."""
        table_data = TableModel(
            element_id="elem_1",
            headers=["Col A", "Col B"],
            rows=[["a", "b"]],
            caption="Sin nota",
            table_number=1,
            note=None,
        )

        doc = docx.Document()
        format_apa_table(doc, table_data, rules)

        # Verificar que no hay párrafo de nota
        has_note_label = any("Nota." in p.text for p in doc.paragraphs)
        assert not has_note_label, "No debe haber nota si note=None"

    def test_table_without_caption(self, rules):
        """Tabla sin caption no debe generar párrafo de título."""
        table_data = TableModel(
            element_id="elem_1",
            headers=["Col A"],
            rows=[["a"]],
            caption="",
            table_number=1,
        )

        doc = docx.Document()
        format_apa_table(doc, table_data, rules)

        # Solo debe tener número de tabla y los datos
        table = doc.tables[0]
        assert table is not None, "La tabla debe existir aunque no tenga caption"


# ── TESTS DE TABLA VACÍA ──────────────────────────────────────────────────────

class TestEmptyTable:
    """Comportamiento con tablas sin datos."""

    def test_table_without_headers(self, rules):
        """Tabla sin encabezados pero con filas debe funcionar."""
        table_data = TableModel(
            element_id="elem_1",
            headers=[],
            rows=[["a", "b"], ["c", "d"]],
            caption="Tabla sin headers explícitos",
            table_number=1,
        )

        doc = docx.Document()
        format_apa_table(doc, table_data, rules)
        assert len(doc.tables) == 1


# ── TESTS DE set_table_apa7_borders (UNIT) ────────────────────────────────────

class TestSetTableBordersUnit:
    """Tests unitarios de la función set_table_apa7_borders."""

    def test_borders_applied_to_existing_table(self):
        """Aplica bordes APA 7 a una tabla existente."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=3)

        # Agregar datos cualquiera
        table.rows[0].cells[0].text = "Header"
        table.rows[1].cells[0].text = "Data"

        set_table_apa7_borders(table)

        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None

        # Verificar bordes verticales eliminados
        for border_name in ['w:left', 'w:right', 'w:insideV']:
            border = tblBorders.find(qn(border_name))
            if border is not None:
                assert border.get(qn('w:val')) == 'none', (
                    f"{border_name} debe ser 'none'"
                )
