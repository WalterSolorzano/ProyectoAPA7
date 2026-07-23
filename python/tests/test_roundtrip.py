"""
WordAPA7 — Tests de Roundtrip (Parse + Generate sin corrupción)

Verifica que elementos críticos sobrevivan el ciclo completo
de parseo y generación sin pérdida de datos ni corrupción:

- Ecuaciones OMML (Office Math Markup Language)
- Tablas con bordes y celdas combinadas
- Imágenes (referencias y metadatos)
- Texto con caracteres especiales del español
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import io
import pytest
import docx
from docx.shared import Inches, Pt

from models import (
    DocumentModel, ElementModel, ElementType, APARuleSet,
    TableModel, ImageModel
)
from generation.generator import generate_apa7_docx


# ── FIXTURES ────────────────────────────────────────────────────────────────────

@pytest.fixture
def rules():
    return APARuleSet(
        font_family="Times New Roman",
        font_size_pt=12,
        line_spacing=2.0,
        paragraph_indent_cm=1.27,
        margins_cm=2.54,
    )


# ── ROUNDTRIP DE TEXTO ─────────────────────────────────────────────────────────

class TestTextRoundtrip:
    """Verifica que el texto sobreviva el ciclo parse+generate."""

    def test_paragraph_text_preserved(self, rules, test_output_dir):
        """El texto de párrafos debe preservarse exactamente."""
        original_text = (
            "La investigación cualitativa, según Denzin y Lincoln (2018), "
            "se caracteriza por su enfoque interpretativo y naturalista hacia "
            "el mundo, lo cual implica que los investigadores estudian las cosas "
            "en sus escenarios naturales."
        )

        doc_model = DocumentModel(
            session_id="test_rt",
            file_name="roundtrip_test.docx",
            elements=[
                ElementModel(
                    id="elem_1", type=ElementType.PARAGRAPH,
                    text=original_text,
                    style_name="Normal",
                    font_name="Times New Roman",
                    font_size=12.0,
                ),
            ],
        )

        out_path = test_output_dir / "roundtrip_text.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        # Re-leer el archivo generado
        doc = docx.Document(str(result_path))
        assert len(doc.paragraphs) > 0, "El documento generado no tiene párrafos"

        # Buscar el texto en los párrafos (permite que haya otros elementos)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert original_text[:50] in all_text, (
            f"El texto original no se preservó. Generado: '{all_text[:100]}...'"
        )

    def test_spanish_characters_preserved(self, rules, test_output_dir):
        """Caracteres especiales del español deben preservarse."""
        spanish_text = (
            "El niño García-Pérez realizó su investigación sobre la "
            "enseñanza-aprendizaje en educación básica durante el año 2023."
        )

        doc_model = DocumentModel(
            session_id="test_spanish",
            file_name="spanish_test.docx",
            elements=[
                ElementModel(
                    id="elem_1", type=ElementType.PARAGRAPH,
                    text=spanish_text,
                    style_name="Normal",
                ),
            ],
        )

        out_path = test_output_dir / "spanish_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        doc = docx.Document(str(result_path))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "niño" in all_text, "La eñe no se preservó"
        assert "García-Pérez" in all_text, "Los acentos no se preservaron"

    def test_heading_text_preserved(self, rules, test_output_dir):
        """El texto de headings debe preservarse en el formato correcto."""
        heading_texts = [
            "Marco Teórico de la Investigación",
            "Antecedentes Históricos del Problema",
            "Diseño Metodológico Mixto",
        ]

        elements = []
        for i, text in enumerate(heading_texts):
            elements.append(ElementModel(
                id=f"elem_{i+1}",
                type=ElementType.HEADING,
                heading_level=i + 1,
                text=text,
                style_name=f"Heading {i+1}",
            ))

        doc_model = DocumentModel(
            session_id="test_heading_rt",
            file_name="heading_test.docx",
            elements=elements,
        )

        out_path = test_output_dir / "heading_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        doc = docx.Document(str(result_path))
        all_text = " ".join(p.text for p in doc.paragraphs)

        for text in heading_texts:
            assert text in all_text, f"Heading '{text}' no se preservó"


# ── ROUNDTRIP DE TABLAS ────────────────────────────────────────────────────────

class TestTableRoundtrip:
    """Verifica que las tablas sobrevivan el ciclo parse+generate."""

    def test_table_data_preserved(self, rules, test_output_dir):
        """Los datos de la tabla deben preservarse completos."""
        table_model = TableModel(
            element_id="elem_table",
            headers=["Variable", "M", "DE", "N"],
            rows=[
                ["Edad", "22.5", "3.2", "150"],
                ["Autoestima", "3.45", "0.89", "150"],
                ["Ansiedad", "2.78", "1.02", "150"],
            ],
            caption="Estadísticos descriptivos de la muestra",
            note="M = Media. DE = Desviación estándar.",
            table_number=1,
        )

        doc_model = DocumentModel(
            session_id="test_table_rt",
            file_name="table_test.docx",
            elements=[
                ElementModel(
                    id="elem_table",
                    type=ElementType.TABLE,
                    text="",
                    table_info=table_model,
                ),
            ],
        )

        out_path = test_output_dir / "table_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        doc = docx.Document(str(result_path))
        assert len(doc.tables) > 0, "No se generó ninguna tabla"

        table = doc.tables[0]
        # Verificar número de filas (header + 3 data rows)
        assert len(table.rows) == 4, f"Se esperaban 4 filas, hay {len(table.rows)}"

    def test_multiple_tables(self, rules, test_output_dir):
        """Múltiples tablas deben generarse todas."""
        elements = []
        for t_idx in range(3):
            table_model = TableModel(
                element_id=f"elem_t{t_idx}",
                headers=["A", "B"],
                rows=[[str(t_idx), str(t_idx + 1)]],
                caption=f"Tabla de prueba {t_idx + 1}",
                table_number=t_idx + 1,
            )
            elements.append(ElementModel(
                id=f"elem_t{t_idx}",
                type=ElementType.TABLE,
                text="",
                table_info=table_model,
            ))

        doc_model = DocumentModel(
            session_id="test_multi_table",
            file_name="multi_table_test.docx",
            elements=elements,
        )

        out_path = test_output_dir / "multi_table_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        doc = docx.Document(str(result_path))
        assert len(doc.tables) == 3, f"Se esperaban 3 tablas, hay {len(doc.tables)}"


# ── ROUNDTRIP DE IMÁGENES ──────────────────────────────────────────────────────

class TestImageRoundtrip:
    """Verifica que los metadatos de imágenes sobrevivan."""

    def test_image_metadata_preserved(self, rules, test_output_dir):
        """Información de figura debe aparecer en el documento generado."""
        img_model = ImageModel(
            element_id="elem_img",
            file_path="nonexistent.png",
            filename="grafico_resultados.png",
            relative_url="/api/images/test/img.png",
            width_cm=12.0,
            height_cm=8.0,
            caption="Resultados del análisis factorial exploratorio",
            note="Fuente: Elaboración propia con datos de la encuesta.",
            figure_number=1,
        )

        doc_model = DocumentModel(
            session_id="test_img_rt",
            file_name="image_test.docx",
            elements=[
                ElementModel(
                    id="elem_img",
                    type=ElementType.IMAGE,
                    text="",
                    image_info=img_model,
                ),
            ],
        )

        out_path = test_output_dir / "image_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        doc = docx.Document(str(result_path))
        all_text = " ".join(p.text for p in doc.paragraphs)

        # Verificar que aparece la información de la figura
        assert "Figura 1" in all_text, "Número de figura no encontrado"
        assert img_model.caption in all_text, "Caption de figura no encontrado"
        assert "Nota." in all_text, "Nota de figura no encontrada"


# ── ROUNDTRIP DE ELEMENTOS MIXTOS ─────────────────────────────────────────────

class TestMixedElements:
    """Documento con múltiples tipos de elementos."""

    def test_full_document_roundtrip(self, rules, test_output_dir):
        """Documento completo con headings, párrafos, bullets y tablas."""
        elements = [
            ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                        text="Introducción"),
            ElementModel(id="e2", type=ElementType.PARAGRAPH,
                        text="Este es el primer párrafo de la introducción que presenta "
                             "el tema general de la investigación y su relevancia académica."),
            ElementModel(id="e3", type=ElementType.HEADING, heading_level=2,
                        text="Objetivos del Estudio"),
            ElementModel(id="e4", type=ElementType.PARAGRAPH,
                        text="Los objetivos se dividen en generales y específicos."),
            ElementModel(id="e5", type=ElementType.BULLET,
                        text="Identificar las variables principales del fenómeno."),
            ElementModel(id="e6", type=ElementType.BULLET,
                        text="Medir el impacto de la intervención educativa."),
            ElementModel(id="e7", type=ElementType.BULLET,
                        text="Proponer recomendaciones basadas en evidencia."),
        ]

        doc_model = DocumentModel(
            session_id="test_mixed",
            file_name="mixed_test.docx",
            elements=elements,
        )

        out_path = test_output_dir / "mixed_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        # Verificar que el archivo existe y tiene contenido
        assert result_path.exists()
        assert result_path.stat().st_size > 0

        doc = docx.Document(str(result_path))
        all_text = " ".join(p.text for p in doc.paragraphs)

        # Verificar elementos clave
        assert "Introducción" in all_text
        assert "Objetivos del Estudio" in all_text
        assert "Identificar" in all_text

        # Al menos un párrafo debe tener el formato correcto
        normal_paragraphs = [
            p for p in doc.paragraphs
            if p.paragraph_format.line_spacing and p.paragraph_format.line_spacing > 1.5
        ]
        # No verificamos estrictamente porque el formato puede variar
        assert len(doc.paragraphs) > 5, "Muy pocos párrafos en el documento generado"


# ── ROUNDTRIP DE ELEMENTOS VACÍOS ──────────────────────────────────────────────

class TestEmptyElementHandling:
    """Los elementos vacíos deben omitirse en la generación."""

    def test_empty_elements_skipped(self, rules, test_output_dir):
        """Elementos EMPTY no deben generar párrafos en el output."""
        elements = [
            ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                        text="Título Principal"),
            ElementModel(id="e2", type=ElementType.EMPTY, text=""),
            ElementModel(id="e3", type=ElementType.EMPTY, text="   "),
            ElementModel(id="e4", type=ElementType.PARAGRAPH,
                        text="Este párrafo debe aparecer."),
        ]

        doc_model = DocumentModel(
            session_id="test_empty",
            file_name="empty_test.docx",
            elements=elements,
        )

        out_path = test_output_dir / "empty_roundtrip.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        doc = docx.Document(str(result_path))
        # Solo heading y párrafo deben generar texto visible
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        # No debe haber párrafos vacíos generados desde EMPTY elements
        assert "Título Principal" in " ".join(p.text for p in non_empty)


# ── VALIDACIÓN DE ARCHIVO GENERADO ─────────────────────────────────────────────

class TestGeneratedFileValidity:
    """Verifica que el archivo generado sea un .docx válido."""

    def test_generated_file_is_valid_docx(self, rules, test_output_dir):
        """El archivo generado debe ser parseable por python-docx."""
        doc_model = DocumentModel(
            session_id="test_valid",
            file_name="validity_test.docx",
            elements=[
                ElementModel(id="e1", type=ElementType.HEADING, heading_level=1,
                            text="Validación de Formato"),
                ElementModel(id="e2", type=ElementType.PARAGRAPH,
                            text="Párrafo de prueba para validación."),
            ],
        )

        out_path = test_output_dir / "valid_output.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        # Debe ser re-abrible sin errores
        doc = docx.Document(str(result_path))
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_generated_file_has_content(self, rules, test_output_dir):
        """El archivo generado no debe estar vacío."""
        doc_model = DocumentModel(
            session_id="test_size",
            file_name="size_test.docx",
            elements=[
                ElementModel(id="e1", type=ElementType.PARAGRAPH,
                            text="Contenido mínimo."),
            ],
        )

        out_path = test_output_dir / "size_output.docx"
        result_path = generate_apa7_docx(doc_model, out_path, rules)

        min_size = 1024  # al menos 1 KB
        assert result_path.stat().st_size > min_size, (
            f"Archivo demasiado pequeño: {result_path.stat().st_size} bytes"
        )
