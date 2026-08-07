"""
WordAPA7 — Tests para Arquitectura XML Profunda:
- structure_scanner.py (Scan de contenedores y deduplicación mc:Choice/mc:Fallback)
- field_guard.py (Protección de campos TOC y updateFields en settings.xml)
- sanity_check.py (Gate de verificación de integridad de texto)
"""

import zipfile

import docx
from lxml import etree
from parsing.field_guard import enable_word_update_fields
from parsing.sanity_check import ExportBlockedError, verify_document_content_integrity
from parsing.structure_scanner import extract_textbox_texts, scan_document_xml_nodes


class TestStructureScanner:
    def test_scan_document_xml_nodes(self, tmp_path):
        doc_path = tmp_path / "test_shapes.docx"
        doc = docx.Document()
        doc.add_paragraph("Texto de párrafo normal en el cuerpo")
        doc.save(doc_path)

        nodes = scan_document_xml_nodes(doc_path)
        assert len(nodes) > 0
        p_node = next(n for n in nodes if "normal" in n.text)
        assert p_node.container_type == "body_paragraph"
        assert p_node.is_editable is True
        assert p_node.is_fallback is False

    def test_extract_textbox_texts(self, tmp_path):
        doc_path = tmp_path / "test_txbx.docx"
        doc = docx.Document()
        doc.add_paragraph("Prueba de extracción")
        doc.save(doc_path)

        texts = extract_textbox_texts(doc_path)
        assert isinstance(texts, list)


class TestFieldGuard:
    def test_enable_word_update_fields(self, tmp_path):
        doc_path = tmp_path / "test_toc.docx"
        doc = docx.Document()
        doc.add_paragraph("Párrafo para TOC")
        doc.save(doc_path)

        result = enable_word_update_fields(doc_path)
        assert result is True

        # Verificar que <w:updateFields w:val="true"/> se insertó en word/settings.xml
        with zipfile.ZipFile(doc_path, 'r') as zf:
            if "word/settings.xml" in zf.namelist():
                xml_data = zf.read("word/settings.xml")
                tree = etree.fromstring(xml_data)
                W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                update_fields = tree.find(f"{{{W_NS}}}updateFields")
                assert update_fields is not None
                assert update_fields.attrib.get(f"{{{W_NS}}}val") == "true"


class TestSanityCheck:
    def test_verify_document_content_integrity_passed(self, tmp_path):
        orig_path = tmp_path / "orig.docx"
        gen_path = tmp_path / "gen.docx"

        d1 = docx.Document()
        d1.add_paragraph("Este es un texto completo de prueba de sanidad de contenido.")
        d1.save(orig_path)

        d2 = docx.Document()
        d2.add_paragraph("Este es un texto completo de prueba de sanidad de contenido.")
        d2.save(gen_path)

        assert verify_document_content_integrity(orig_path, gen_path, tolerance=0.05) is True

    def test_verify_document_content_integrity_failed(self, tmp_path):
        orig_path = tmp_path / "orig_large.docx"
        gen_path = tmp_path / "gen_empty.docx"

        d1 = docx.Document()
        for i in range(10):
            d1.add_paragraph(f"Párrafo de contenido importante número {i} con información crítica de la portada.")
        d1.save(orig_path)

        d2 = docx.Document()
        d2.add_paragraph("Corto")
        d2.save(gen_path)

        try:
            verify_document_content_integrity(orig_path, gen_path, tolerance=0.02)
            assert False, "Debería haber lanzado ExportBlockedError"
        except ExportBlockedError as e:
            assert "Gate de Sanidad" in str(e)
