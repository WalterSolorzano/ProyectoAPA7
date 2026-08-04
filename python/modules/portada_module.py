"""
WordAPA7 — Módulo de Portada APA 7

Genera y formatea la página de portada en sus dos modalidades APA 7:
1. Estudiante: Título en negrita centrado en la mitad superior, autor, afiliación, curso, profesor, fecha.
2. Profesional: Incluye nota de autor (Author Note) y Running Head.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from models import PortadaData, APAFormat, APARuleSet


class _CoverBuilder:
    """Construye la portada INSERTÁNDOLA al inicio del cuerpo del documento
    (posición 0) en vez de al final — evita títulos desordenados y portada
    colgando al final del archivo generado."""

    def __init__(self, doc: docx.Document):
        self.doc = doc
        self._body = doc.element.body
        self._idx = 0

    def add_paragraph(self) -> Paragraph:
        p_elem = OxmlElement('w:p')
        self._body.insert(self._idx, p_elem)
        self._idx += 1
        return Paragraph(p_elem, self.doc)

    def add_page_break(self) -> None:
        p_elem = OxmlElement('w:p')
        r_elem = OxmlElement('w:r')
        br_elem = OxmlElement('w:br')
        br_elem.set(docx.oxml.ns.qn('w:type'), 'page')
        r_elem.append(br_elem)
        p_elem.append(r_elem)
        self._body.insert(self._idx, p_elem)
        self._idx += 1

    @property
    def paragraph_count(self) -> int:
        return self._idx


def format_apa_portada(
    doc: docx.Document,
    portada: PortadaData,
    rules: APARuleSet,
    original_textbox_texts: list[str] | None = None,
):
    """
    Inserta la portada APA 7 formateada al inicio del documento.

    Si se proporciona original_textbox_texts, estos se insertan como parrafos
    centrados preservados ANTES de la portada generada, para no perder
    contenido original de cuadros de texto, WordArt, logos, etc.
    """
    font_name = rules.font_family
    font_size = Pt(rules.font_size_pt)

    builder = _CoverBuilder(doc)

    # Contenido original de cuadros de texto (preservado antes de la portada)
    if original_textbox_texts:
        for tb_text in original_textbox_texts:
            if not tb_text or not tb_text.strip():
                continue
            p_tb = builder.add_paragraph()
            p_tb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_tb.paragraph_format.line_spacing = rules.line_spacing
            p_tb.paragraph_format.first_line_indent = Inches(0)
            p_tb.paragraph_format.space_before = Pt(0)
            p_tb.paragraph_format.space_after = Pt(6)
            r_tb = p_tb.add_run(tb_text.strip())
            r_tb.font.name = font_name
            r_tb.font.size = font_size
            r_tb.font.color.rgb = RGBColor(0, 0, 0)

        # Pequeña separacion antes de la portada generada
        p_sep = builder.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(12)
        p_sep.paragraph_format.space_after = Pt(6)

    # 3 o 4 líneas en blanco iniciales
    for _ in range(3):
        p_blank = builder.add_paragraph()
        p_blank.paragraph_format.line_spacing = rules.line_spacing
        p_blank.paragraph_format.space_before = Pt(0)
        p_blank.paragraph_format.space_after = Pt(0)

    # Título en Negrita, Centrado
    p_title = builder.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = rules.line_spacing
    p_title.paragraph_format.first_line_indent = Inches(0)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)

    r_title = p_title.add_run(portada.title or "Título del Trabajo")
    r_title.bold = True
    r_title.font.name = font_name
    r_title.font.size = font_size
    r_title.font.color.rgb = RGBColor(0, 0, 0)

    # 1 o 2 líneas en blanco de separación
    p_sep = builder.add_paragraph()
    p_sep.paragraph_format.line_spacing = rules.line_spacing

    # Datos del Autor y Afiliación
    lines = []
    if portada.author:
        lines.append(portada.author)
    if portada.institution:
        lines.append(portada.institution)

    if portada.apa_format == APAFormat.STUDENT:
        if portada.course:
            lines.append(portada.course)
        if portada.instructor:
            lines.append(portada.instructor)
        if portada.date:
            lines.append(portada.date)

    for line in lines:
        p_line = builder.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.line_spacing = rules.line_spacing
        p_line.paragraph_format.first_line_indent = Inches(0)
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(0)

        r = p_line.add_run(line)
        r.font.name = font_name
        r.font.size = font_size
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Nota de Autor para portada profesional al pie
    if portada.apa_format == APAFormat.PROFESSIONAL and portada.author_note:
        p_an_hdr = builder.add_paragraph()
        p_an_hdr.paragraph_format.space_before = Pt(36)
        p_an_hdr.paragraph_format.space_after = Pt(6)
        p_an_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_an_hdr = p_an_hdr.add_run("Nota del Autor")
        r_an_hdr.bold = True
        r_an_hdr.font.name = font_name
        r_an_hdr.font.size = font_size

        p_an = builder.add_paragraph()
        p_an.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_an.paragraph_format.line_spacing = rules.line_spacing
        p_an.paragraph_format.first_line_indent = Inches(0.5)

        r_an = p_an.add_run(portada.author_note)
        r_an.font.name = font_name
        r_an.font.size = font_size

    # Salto de página al finalizar la portada
    builder.add_page_break()

    return builder.paragraph_count
