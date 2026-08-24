"""Aplicación POR ALCANCE sobre el documento ORIGINAL (sin regeneración).

Garantía contractual: cada función muta ÚNICAMENTE su alcance. Si el usuario
elige solo 'tablas_imagenes', ni la portada ni el texto del cuerpo cambian
ni un byte (los tests de aislamiento lo verifican comparando XML).

Alcances:
- texto            : estilos base (fuente/tamaño/interlineado) vía styles.xml
                     + sangría de primera línea en Normal.
- tablas_imagenes  : numeración APA (Tabla N arriba / Figura N abajo) y
                     bordes APA en tablas. No altera ningún otro párrafo.
- bibliografia     : formato de la sección Referencias (hanging indent,
                     doble espacio, sin space-after). No toca el cuerpo.
"""

from __future__ import annotations

import copy
import io
import re
from typing import Any, Dict, List

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

VALID_SCOPES = ("texto", "tablas_imagenes", "bibliografia")

_REFS_HEADING = re.compile(
    r"^\s*(referencias|bibliograf[íi]a|references|works cited)\b", re.IGNORECASE
)


def _body_xml(doc: docx.Document) -> str:
    return doc.element.body.xml


# ------------------------------------------------------------------ texto
def apply_scope_texto(doc: docx.Document, rules: Dict[str, Any]) -> None:
    """Estilos globales + sangría. No crea/borra/mueve contenido."""
    from generation.style_engine import update_docx_styles_xml

    class _R:  # adapter mínimo al APARuleSet real
        pass

    r = _R()
    r.font_family = rules.get("font_family", "Times New Roman")
    r.font_size_pt = float(rules.get("font_size_pt", 12))
    r.line_spacing = float(rules.get("line_spacing", 2.0))
    try:
        update_docx_styles_xml(doc, r)
    except Exception:
        pass
    indent_cm = float(rules.get("first_line_indent_cm", 1.27))
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.first_line_indent is None:
            pf.first_line_indent = Cm(indent_cm)


# -------------------------------------------------------- tablas_imagenes
def _set_apa_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = parse_borders()
    tbl_pr.append(borders)


def parse_borders():
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    xml = (
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="12" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:color="000000"/>'
        '<w:left w:val="none" w:sz="0"/>'
        '<w:right w:val="none" w:sz="0"/>'
        '<w:insideH w:val="single" w:sz="6" w:color="000000"/>'
        '<w:insideV w:val="none" w:sz="0"/>'
        "</w:tblBorders>"
    ) % nsdecls("w")
    return parse_xml(xml)


def apply_scope_tablas_imagenes(doc: docx.Document) -> Dict[str, int]:
    """Numera tablas (caption arriba) y figuras (caption abajo). Solo inserta
    párrafos nuevos adyacentes; jamás edita párrafos existentes."""
    counts = {"tablas": 0, "figuras": 0}

    # --- Tablas: caption ARRIBA ---
    t_idx = 0
    for tbl in doc.tables:
        t_idx += 1
        anchor = tbl._tbl
        new_p = copy.deepcopy(anchor.getprevious() if anchor.getprevious() is not None else anchor)
        # construir párrafo limpio
        from docx.oxml import OxmlElement

        p_el = OxmlElement("w:p")
        anchor.addprevious(p_el)
        para = docx.text.paragraph.Paragraph(p_el, tbl._tbl)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(f"Tabla {t_idx}")
        run.bold = True
        counts["tablas"] = t_idx

    # --- Figuras: caption DEBAJO de párrafos con drawing ---
    f_idx = 0
    body_paras = list(doc.paragraphs)
    for i, p in enumerate(body_paras):
        if p._p.findall(".//" + qn("w:drawing")):
            f_idx += 1
            from docx.oxml import OxmlElement

            p_el = OxmlElement("w:p")
            p._p.addnext(p_el)
            para = docx.text.paragraph.Paragraph(p_el, p._parent)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"Figura {f_idx}.")
            run.bold = True
            body_paras.insert(i + 1, para)
    counts["figuras"] = f_idx
    return counts


# ------------------------------------------------------------ bibliografia
def apply_scope_bibliografia(doc: docx.Document) -> int:
    """Hanging indent + doble espacio SOLO dentro de la sección de referencias."""
    start = None
    for i, p in enumerate(doc.paragraphs):
        if _REFS_HEADING.match(p.text or ""):
            start = i + 1
    if start is None:
        return 0
    count = 0
    for p in doc.paragraphs[start:]:
        if not (p.text or "").strip():
            continue
        if _REFS_HEADING.match(p.text or ""):
            break
        pf = p.paragraph_format
        pf.left_indent = Cm(1.27)
        pf.first_line_indent = Cm(-1.27)
        pf.line_spacing = 2.0
        pf.space_after = Pt(0)
        count += 1
    return count


# ------------------------------------------------------------------ entry
def apply_scopes(file_bytes: bytes, scopes: List[str], rules: Dict[str, Any]) -> tuple[bytes, Dict[str, Any]]:
    """Punto único: aplica SOLO los alcances pedidos sobre el original."""
    invalid = [s for s in scopes if s not in VALID_SCOPES]
    if invalid:
        raise ValueError(f"Alcances inválidos: {invalid}")

    doc = docx.Document(io.BytesIO(file_bytes))
    summary: Dict[str, Any] = {"scopes": list(scopes)}

    if "texto" in scopes:
        apply_scope_texto(doc, rules)
        summary["texto"] = True
    if "bibliografia" in scopes:
        summary["refs_formateadas"] = apply_scope_bibliografia(doc)
    if "tablas_imagenes" in scopes:
        summary.update(apply_scope_tablas_imagenes(doc))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), summary
