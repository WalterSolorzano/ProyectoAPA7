"""
WordAPA7 — Disenador de Portadas (Cover Designer)

Permite al usuario elegir una plantilla de portada (imagen, docx, o integrada)
que se antepone al documento final SIN ser modificada por el sistema APA 7.
La portada se guarda como plantilla reutilizable en storage/cover_templates/.
"""

import json
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

COVER_TEMPLATES_DIR = "cover_templates"


@dataclass
class CoverTemplate:
    name: str
    description: str
    source_type: str  # "image", "docx", "builtin"
    source_path: str
    preview_path: str
    is_builtin: bool
    created_at: str = ""


def _get_storage_dir(storage_dir: Path) -> Path:
    """Retorna el directorio de almacenamiento de plantillas.
    
    Recibe STORAGE_DIR directamente (no BASE_DIR) para que en producción
    escriba en AppData en lugar de Program Files (que es de solo lectura).
    """
    cover_dir = storage_dir / COVER_TEMPLATES_DIR
    cover_dir.mkdir(parents=True, exist_ok=True)
    return cover_dir


def _template_dir(storage_dir: Path, name: str) -> Path:
    """Retorna el directorio para una plantilla especifica."""
    tdir = storage_dir / _sanitize_name(name)
    tdir.mkdir(parents=True, exist_ok=True)
    return tdir


def _sanitize_name(name: str) -> str:
    """Sanitiza el nombre para uso como nombre de directorio."""
    safe = "".join(c if c.isalnum() or c in " _-." else "_" for c in name)
    return safe.strip() or "untitled"


def _metadata_path(tdir: Path) -> Path:
    return tdir / "metadata.json"


def _source_path(tdir: Path, ext: str) -> Path:
    return tdir / f"source.{ext}"


def _preview_path(tdir: Path) -> Path:
    return tdir / "preview.png"


# ── PLANTILLAS INTEGRADAS (builtin) ────────────────────────────────────────

_BUILTIN_TEMPLATES: List[CoverTemplate] = [
    CoverTemplate(
        name="APA 7 Estudiante",
        description="Portada APA 7 clasica para trabajos de estudiante: titulo en negrita, autor, institucion, curso, profesor, fecha.",
        source_type="builtin",
        source_path="",
        preview_path="",
        is_builtin=True,
        created_at="2026-01-01T00:00:00",
    ),
    CoverTemplate(
        name="APA 7 Profesional",
        description="Portada APA 7 profesional con Nota del Autor (Author Note) y espacio para Running Head.",
        source_type="builtin",
        source_path="",
        preview_path="",
        is_builtin=True,
        created_at="2026-01-01T00:00:00",
    ),
    CoverTemplate(
        name="Portada Minimalista",
        description="Portada limpia y sobria con solo titulo centrado y linea decorativa delgada.",
        source_type="builtin",
        source_path="",
        preview_path="",
        is_builtin=True,
        created_at="2026-01-01T00:00:00",
    ),
    CoverTemplate(
        name="Portada con Logos",
        description="Portada APA 7 que preserva logotipos institucionales en la cabecera antes del titulo.",
        source_type="builtin",
        source_path="",
        preview_path="",
        is_builtin=True,
        created_at="2026-01-01T00:00:00",
    ),
]


# ── FUNCIONES PUBLICAS ─────────────────────────────────────────────────────

def list_cover_templates(base_dir: Path) -> List[CoverTemplate]:
    """
    Lista todas las plantillas de portada disponibles.
    Combina las plantillas integradas con las guardadas por el usuario.
    """
    storage_dir = _get_storage_dir(base_dir)
    templates = list(_BUILTIN_TEMPLATES)

    if not storage_dir.exists():
        return templates

    for item in sorted(storage_dir.iterdir()):
        if item.is_dir():
            meta_path = _metadata_path(item)
            if meta_path.exists():
                try:
                    with open(meta_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    # Convertir a CoverTemplate, asegurando campos por defecto
                    templates.append(CoverTemplate(
                        name=data.get("name", item.name),
                        description=data.get("description", ""),
                        source_type=data.get("source_type", "docx"),
                        source_path=str(_source_path(item, "docx")),
                        preview_path=str(_preview_path(item)) if _preview_path(item).exists() else "",
                        is_builtin=False,
                        created_at=data.get("created_at", ""),
                    ))
                except (json.JSONDecodeError, KeyError) as e:
                    print(f"[WARN] Error leyendo metadatos de plantilla {item.name}: {e}")
                    continue

    return templates


def create_cover_from_image(
    image_path: Path,
    name: str,
    description: str,
    base_dir: Path,
) -> CoverTemplate:
    """
    Crea una plantilla de portada a partir de una imagen.
    La imagen se copia al almacenamiento de plantillas y se genera un preview.
    """
    storage_dir = _get_storage_dir(base_dir)
    tdir = _template_dir(storage_dir, name)

    # Copiar imagen al almacenamiento
    ext = image_path.suffix.lower().lstrip(".") or "png"
    dest = _source_path(tdir, ext)
    shutil.copy2(image_path, dest)

    # Crear preview (copia directa por ahora)
    preview = _preview_path(tdir)
    shutil.copy2(image_path, preview)

    now = datetime.now().isoformat()
    metadata = {
        "name": name,
        "description": description or f"Portada desde imagen: {image_path.name}",
        "source_type": "image",
        "source_ext": ext,
        "created_at": now,
    }
    with open(_metadata_path(tdir), "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)

    return CoverTemplate(
        name=name,
        description=metadata["description"],
        source_type="image",
        source_path=str(dest),
        preview_path=str(preview),
        is_builtin=False,
        created_at=now,
    )


def create_cover_from_docx(
    docx_path: Path,
    name: str,
    description: str,
    base_dir: Path,
) -> CoverTemplate:
    """
    Crea una plantilla de portada a partir de un documento Word.
    Extrae solo la primera pagina del DOCX y la guarda como plantilla.
    """
    storage_dir = _get_storage_dir(base_dir)
    tdir = _template_dir(storage_dir, name)

    # Copiar el DOCX completo al almacenamiento
    dest = _source_path(tdir, "docx")
    shutil.copy2(docx_path, dest)

    # Intentar generar preview desde la primera pagina del DOCX
    preview = _preview_path(tdir)
    _generate_docx_preview(docx_path, preview)

    now = datetime.now().isoformat()
    metadata = {
        "name": name,
        "description": description or f"Portada desde Word: {docx_path.name}",
        "source_type": "docx",
        "created_at": now,
    }
    with open(_metadata_path(tdir), "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)

    return CoverTemplate(
        name=name,
        description=metadata["description"],
        source_type="docx",
        source_path=str(dest),
        preview_path=str(preview) if preview.exists() else "",
        is_builtin=False,
        created_at=now,
    )


def delete_cover_template(name: str, base_dir: Path) -> bool:
    """
    Elimina una plantilla de portada creada por el usuario.
    Las plantillas integradas no se pueden eliminar.
    """
    storage_dir = _get_storage_dir(base_dir)
    tdir = storage_dir / _sanitize_name(name)

    if not tdir.exists():
        return False

    shutil.rmtree(tdir)
    return True


def apply_cover_to_document(
    doc: docx.Document,
    cover_template: CoverTemplate,
    base_dir: Path,
    portada_title: str = "",
    portada_author: str = "",
    portada_institution: str = "",
    portada_course: str = "",
    portada_instructor: str = "",
    portada_date: str = "",
) -> int:
    """
    Antepone la portada al documento.
    Retorna cuantos parrafos se agregaron.

    - IMAGEN: crea pagina completa con imagen centrada
    - DOCX: copia contenido del DOCX origen al inicio del doc destino
    - BUILTIN: genera portada APA 7 con formato predefinido
    """
    paragraphs_before = len(doc.paragraphs)
    body = doc.element.body
    children_before = len(list(body))

    if cover_template.source_type == "image":
        _apply_image_cover(doc, cover_template, base_dir)
    elif cover_template.source_type == "docx":
        _apply_docx_cover(doc, cover_template, base_dir)
    elif cover_template.source_type == "builtin":
        _apply_builtin_cover(
            doc, cover_template, base_dir,
            portada_title, portada_author, portada_institution,
            portada_course, portada_instructor, portada_date,
        )

    # Salto de pagina obligatorio despues de la portada
    doc.add_page_break()

    # Los builders anexan al final del cuerpo; mover los elementos recién
    # creados a la posición 0 para ANTEPONER la portada al contenido existente.
    new_elements = list(body)[children_before:]
    for offset, element in enumerate(new_elements):
        body.insert(offset, element)

    # Activar different_first_page_header_footer para que header APA solo aparezca desde pagina 2
    for section in doc.sections:
        section.different_first_page_header_footer = True

    paragraphs_added = len(doc.paragraphs) - paragraphs_before
    return paragraphs_added


# ── FUNCIONES INTERNAS DE APLICACION ──────────────────────────────────────

def _apply_image_cover(doc: docx.Document, template: CoverTemplate, base_dir: Path) -> None:
    """
    Inserta la imagen como portada a pagina completa.
    La imagen se escala para ocupar toda la pagina con margen APA.
    """
    source = Path(template.source_path)
    if not source.exists():
        # Fallback: texto plano
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Portada: {template.name}]")
        r.font.size = Pt(12)
        return

    # Agregar la imagen centrada, escalada a ~15cm de ancho (margenes APA)
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run()
        run.add_picture(str(source), width=Inches(5.5))
    except Exception as e:
        print(f"[WARN] Error insertando portada imagen: {e}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Portada: {template.name}]")
        r.font.size = Pt(12)


def _apply_docx_cover(doc: docx.Document, template: CoverTemplate, base_dir: Path) -> None:
    """
    Extrae el contenido del DOCX plantilla y lo inserta al inicio del documento destino.
    Solo toma la primera seccion (portada) del DOCX origen.
    """
    source = Path(template.source_path)
    if not source.exists():
        _apply_builtin_cover(doc, None, base_dir, "", "", "", "", "", "")
        return

    try:
        cover_doc = docx.Document(source)

        # Preservar el primer encabezado/logo si existe
        for section in cover_doc.sections:
            header = section.header
            if header and not header.is_linked_to_previous:
                # Copiar header si existe
                for header_para in header.paragraphs:
                    if header_para.text.strip():
                        new_p = doc.add_paragraph()
                        new_p.alignment = header_para.alignment
                        for run in header_para.runs:
                            new_run = new_p.add_run(run.text)
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            new_run.bold = run.bold
                            new_run.italic = run.italic

        # Copiar parrafos del DOCX plantilla al documento destino
        for para in cover_doc.paragraphs:
            new_p = doc.add_paragraph()
            new_p.alignment = para.alignment
            new_p.paragraph_format.space_before = para.paragraph_format.space_before
            new_p.paragraph_format.space_after = para.paragraph_format.space_after
            new_p.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            new_p.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent

            for run in para.runs:
                new_run = new_p.add_run(run.text)
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

        # NOTA: Las imágenes incrustadas en el DOCX plantilla no se copian en esta
        # versión. Para portadas con imágenes, usa el tipo "image" en lugar de "docx".

    except Exception as e:
        print(f"[WARN] Error copiando portada DOCX: {e}")
        _apply_builtin_cover(doc, None, base_dir, "", "", "", "", "", "")


def _apply_builtin_cover(
    doc: docx.Document,
    template: Optional[CoverTemplate],
    base_dir: Path,
    title: str = "",
    author: str = "",
    institution: str = "",
    course: str = "",
    instructor: str = "",
    date_str: str = "",
) -> None:
    """
    Genera una portada con formato APA 7 predefinido.
    Funciona como fallback cuando no hay plantilla.
    """
    font_name = "Times New Roman"
    font_size = Pt(12)

    if template and template.name == "APA 7 Profesional":
        # Profesional: incluye Author Note
        _generate_professional_cover(doc, title, author, institution, author_note="")
    elif template and template.name == "Portada Minimalista":
        # Minimalista: solo titulo con linea decorativa
        for _ in range(6):
            doc.add_paragraph().paragraph_format.line_spacing = 2.0

        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_line = p_line.add_run("─" * 50)
        r_line.font.name = font_name
        r_line.font.size = Pt(10)
        r_line.font.color.rgb = RGBColor(100, 100, 100)

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.line_spacing = 2.0
        p_title.paragraph_format.first_line_indent = Inches(0)
        r_title = p_title.add_run(title or "Titulo del Trabajo")
        r_title.bold = True
        r_title.font.name = font_name
        r_title.font.size = Pt(16)

        p_line2 = doc.add_paragraph()
        p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_line2 = p_line2.add_run("─" * 50)
        r_line2.font.name = font_name
        r_line2.font.size = Pt(10)
        r_line2.font.color.rgb = RGBColor(100, 100, 100)

        for _ in range(3):
            doc.add_paragraph().paragraph_format.line_spacing = 2.0

        if author:
            for author_line in [ln.strip() for ln in author.split('\n') if ln.strip()]:
                p_author = doc.add_paragraph()
                p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_author.paragraph_format.line_spacing = 2.0
                r_author = p_author.add_run(author_line)
                r_author.font.name = font_name
                r_author.font.size = font_size

        if institution:
            p_inst = doc.add_paragraph()
            p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_inst.paragraph_format.line_spacing = 2.0
            r_inst = p_inst.add_run(institution)
            r_inst.font.name = font_name
            r_inst.font.size = font_size

    elif template and template.name == "Portada con Logos":
        # Portada con espacio para logos
        p_logo_placeholder = doc.add_paragraph()
        p_logo_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_logo = p_logo_placeholder.add_run("[LOGOS INSTITUCIONALES]")
        r_logo.font.name = font_name
        r_logo.font.size = Pt(14)
        r_logo.font.color.rgb = RGBColor(150, 150, 150)
        r_logo.italic = True

        for _ in range(4):
            doc.add_paragraph().paragraph_format.line_spacing = 2.0

        # Titulo centrado en negrita
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.line_spacing = 2.0
        p_title.paragraph_format.first_line_indent = Inches(0)
        r_title = p_title.add_run(title or "Titulo del Trabajo")
        r_title.bold = True
        r_title.font.name = font_name
        r_title.font.size = Pt(14)

        for _ in range(2):
            doc.add_paragraph().paragraph_format.line_spacing = 2.0

        if author:
            for author_line in [ln.strip() for ln in author.split('\n') if ln.strip()]:
                p_a = doc.add_paragraph()
                p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_a.paragraph_format.line_spacing = 2.0
                r_a = p_a.add_run(author_line)
                r_a.font.name = font_name
                r_a.font.size = font_size

        if institution:
            p_i = doc.add_paragraph()
            p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_i.paragraph_format.line_spacing = 2.0
            r_i = p_i.add_run(institution)
            r_i.font.name = font_name
            r_i.font.size = font_size
    else:
        # Estudiante (default APA 7)
        _generate_student_cover(doc, title, author, institution, course, instructor, date_str)


def _generate_student_cover(
    doc: docx.Document,
    title: str,
    author: str,
    institution: str,
    course: str,
    instructor: str,
    date_str: str,
) -> None:
    """Portada APA 7 estilo estudiante."""
    font_name = "Times New Roman"
    font_size = Pt(12)

    # 3-4 lineas en blanco
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Titulo en negrita centrado
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 2.0
    p_title.paragraph_format.first_line_indent = Inches(0)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    r_title = p_title.add_run(title or "Titulo del Trabajo")
    r_title.bold = True
    r_title.font.name = font_name
    r_title.font.size = font_size
    r_title.font.color.rgb = RGBColor(0, 0, 0)

    # Linea en blanco
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.line_spacing = 2.0

    # Datos del autor
    lines = []
    if author:
        for _ln in author.split('\n'):
            _ln = _ln.strip()
            if _ln:
                lines.append(_ln)
    if institution:
        lines.append(institution)
    if course:
        lines.append(course)
    if instructor:
        lines.append(instructor)
    if date_str:
        lines.append(date_str)

    for line in lines:
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.line_spacing = 2.0
        p_line.paragraph_format.first_line_indent = Inches(0)
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(0)
        r = p_line.add_run(line)
        r.font.name = font_name
        r.font.size = font_size
        r.font.color.rgb = RGBColor(0, 0, 0)


def _generate_professional_cover(
    doc: docx.Document,
    title: str,
    author: str,
    institution: str,
    author_note: str,
) -> None:
    """Portada APA 7 estilo profesional con Author Note."""
    font_name = "Times New Roman"
    font_size = Pt(12)

    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 2.0
    p_title.paragraph_format.first_line_indent = Inches(0)
    r_title = p_title.add_run(title or "Titulo del Trabajo")
    r_title.bold = True
    r_title.font.name = font_name
    r_title.font.size = font_size

    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.line_spacing = 2.0

    if author:
        for author_line in [ln.strip() for ln in author.split('\n') if ln.strip()]:
            p_a = doc.add_paragraph()
            p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_a.paragraph_format.line_spacing = 2.0
            p_a.paragraph_format.first_line_indent = Inches(0)
            r_a = p_a.add_run(author_line)
            r_a.font.name = font_name
            r_a.font.size = font_size

    if institution:
        p_i = doc.add_paragraph()
        p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_i.paragraph_format.line_spacing = 2.0
        p_i.paragraph_format.first_line_indent = Inches(0)
        r_i = p_i.add_run(institution)
        r_i.font.name = font_name
        r_i.font.size = font_size

    if author_note:
        p_an_hdr = doc.add_paragraph()
        p_an_hdr.paragraph_format.space_before = Pt(36)
        p_an_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_an_hdr = p_an_hdr.add_run("Nota del Autor")
        r_an_hdr.bold = True
        r_an_hdr.font.name = font_name
        r_an_hdr.font.size = font_size

        p_an = doc.add_paragraph()
        p_an.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_an.paragraph_format.line_spacing = 2.0
        p_an.paragraph_format.first_line_indent = Inches(0.5)
        r_an = p_an.add_run(author_note)
        r_an.font.name = font_name
        r_an.font.size = font_size


def _generate_docx_preview(docx_path: Path, preview_path: Path) -> None:
    """
    Genera un preview del DOCX. Como no podemos renderizar imagenes directamente,
    creamos un marcador de preview. En produccion, se podria convertir a PNG
    usando LibreOffice o un servicio externo.
    """
    if not docx_path.exists():
        return
    try:
        # Crear un archivo de texto con el nombre del template como preview placeholder
        with open(str(preview_path) + ".txt", "w", encoding="utf-8") as f:
            f.write(f"Preview: {docx_path.stem}\n")
    except Exception:
        pass
