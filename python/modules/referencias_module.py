"""
WordAPA7 — Modulo de Referencias Bibliograficas APA 7

Genera y formatea la seccion de Referencias:
1. Titulo "Referencias" centrado en negrita en pagina nueva.
2. Sangria francesa (hanging indent) de 1.27 cm (0.5 in).
3. Interlineado doble (2.0) sin espacios adicionales.
4. Orden alfabetico por primer autor.
5. Cursiva en titulos de libros, revistas y volumen.
6. DOI content negotiation via Crossref (gratis, sin API key).
"""

import re
from typing import List, Optional
from urllib.parse import quote

import httpx

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from models import ReferenciaModel, APARuleSet
from generation.style_engine import set_run_font


_REF_PREFIX_RE = re.compile(
    r'^(?:[•○▪–\-\*]|\(?\d+[\.\)]|\(?[a-zA-Z][\.\)])\s*'
)


def _strip_ref_prefix(text: str) -> str:
    return _REF_PREFIX_RE.sub('', text.strip())


# ── DOI Content Negotiation ────────────────────────────────────────────────────

DOI_RESOLVER_URL: str = "https://doi.org/"
CROSSREF_ACCEPT_HEADER: str = "text/x-bibliography; style=apa; locale=es-419"


async def resolve_doi(doi_input: str) -> Optional[str]:
    """
    Usa DOI content negotiation — GRATIS, sin API key.

    GET https://doi.org/{doi}
    Header: Accept: text/x-bibliography; style=apa; locale=es-419

    Retorna el texto APA 7 formateado directamente desde Crossref.
    Si falla, retorna None para que se use el fallback LLM.
    """
    # Limpiar DOI: remover prefijo "doi:" o "https://doi.org/" si existe
    doi_clean: str = doi_input.strip()
    doi_clean = re.sub(r'^https?://doi\.org/', '', doi_clean)
    doi_clean = re.sub(r'^doi:\s*', '', doi_clean, flags=re.IGNORECASE)
    doi_clean = doi_clean.strip()

    if not doi_clean:
        return None

    url: str = f"{DOI_RESOLVER_URL}{doi_clean}"

    try:
        async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
            resp = await client.get(
                url,
                headers={"Accept": CROSSREF_ACCEPT_HEADER},
            )
            if resp.status_code == 200 and resp.text.strip():
                return resp.text.strip()
            elif resp.status_code == 404:
                print(f"[INFO] DOI no encontrado en Crossref: {doi_clean}")
                return None
            else:
                print(f"[WARN] Crossref devolvio status {resp.status_code} para DOI: {doi_clean}")
                return None
    except httpx.TimeoutException:
        print(f"[WARN] Timeout consultando DOI: {doi_clean}")
        return None
    except Exception as err:
        print(f"[WARN] Error consultando DOI {doi_clean}: {err}")
        return None


import re

def extract_doi_or_isbn(text: str):
    doi_match = re.search(r'(10\.\d{4,9}/[-._;()/:A-Z0-9]+)', text, re.IGNORECASE)
    if doi_match:
        return {"type": "doi", "value": doi_match.group(1)}
    
    isbn_match = re.search(r'(?:ISBN(?:-1[03])?:? )?(?=[-0-9 ]{13,17})(?:97[89][- ]?)?[0-9]{1,5}[- ]?[0-9]+[- ]?[0-9]+[- ]?[0-9X]', text, re.IGNORECASE)
    if isbn_match:
        isbn_clean = re.sub(r'[^0-9X]', '', isbn_match.group(0).upper())
        if len(isbn_clean) in (10, 13):
            return {"type": "isbn", "value": isbn_clean}
    return None

async def fetch_crossref_metadata(doi: str) -> Optional[dict]:
    url = f"https://api.crossref.org/works/{doi}"
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.get(url)
            if resp.status_code == 200:
                data = resp.json()["message"]
                authors = []
                for author in data.get("author", []):
                    if "family" in author and "given" in author:
                        authors.append(f"{author['family']}, {author['given'][0]}.")
                    elif "family" in author:
                        authors.append(author["family"])
                
                author_str = ", & ".join(authors) if authors else "Autor desconocido"
                year = data.get("published-print", {}).get("date-parts", [[None]])[0][0]
                if not year:
                    year = data.get("published-online", {}).get("date-parts", [[None]])[0][0]
                year_str = str(year) if year else "s.f."
                
                title = data.get("title", [""])[0]
                container_title = data.get("container-title", [""])[0]
                volume = data.get("volume", "")
                issue = data.get("issue", "")
                page = data.get("page", "")
                
                formatted = f"{author_str} ({year_str}). {title}."
                if container_title:
                    formatted += f" {container_title}"
                    if volume:
                        formatted += f", {volume}"
                        if issue:
                            formatted += f"({issue})"
                    if page:
                        formatted += f", {page}."
                    else:
                        formatted += "."
                
                return {
                    "author": author_str,
                    "year": year_str,
                    "title": title,
                    "journal": container_title,
                    "volume": volume,
                    "issue": issue,
                    "page": page,
                    "doi": doi,
                    "formatted": formatted
                }
    except Exception as e:
        print(f"[WARN] Error in Crossref API: {e}")
    return None


async def search_crossref_by_author_year(authors: list[str], year: str) -> Optional[dict]:
    """
    Busca en Crossref API por apellido(s) de autor + año.
    API gratuita, no requiere key. Retorna el mejor match formateado en APA 7.

    Args:
        authors: lista de apellidos, ej: ["Garcia", "Lopez"]
        year: año de publicación, ej: "2023"

    Returns: dict con author, year, title, source, doi, formatted_apa, o None
    """
    query = " ".join(authors) + " " + year
    url = f"https://api.crossref.org/works?query={quote(query)}&rows=5"
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.get(url, headers={"User-Agent": "WordAPA7/1.0 (mailto:wordapa7@antigravity.dev)"})
            if resp.status_code != 200:
                return None
            data = resp.json()
            items = data.get("message", {}).get("items", [])
            if not items:
                return None

            results = []
            for item in items:
                # Formatear autores APA 7
                apa_authors = []
                item_authors = item.get("author", [])
                for author in item_authors:
                    family = author.get("family", "")
                    given = author.get("given", "")
                    if family and given:
                        apa_authors.append(f"{family}, {given[0]}.")
                    elif family:
                        apa_authors.append(family)
                if not apa_authors:
                    apa_authors = ["Autor desconocido"]
                author_str = ", & ".join(apa_authors) if len(apa_authors) <= 2 else (
                    ", ".join(apa_authors[:-1]) + ", & " + apa_authors[-1]
                )

                # Año
                pub_year = item.get("published-print", {}).get("date-parts", [[None]])[0][0]
                if not pub_year:
                    pub_year = item.get("published-online", {}).get("date-parts", [[None]])[0][0]
                if not pub_year:
                    pub_year = item.get("created", {}).get("date-parts", [[None]])[0][0]
                year_str = str(pub_year) if pub_year else year

                # Título
                titles = item.get("title", [])
                title_str = titles[0] if titles else "Sin título"

                # Source
                container = item.get("container-title", [])
                source_str = container[0] if container else ""
                publisher = item.get("publisher", "")
                if not source_str and publisher:
                    source_str = publisher

                # DOI
                doi = item.get("DOI", "")

                # Volume, issue, pages
                volume = item.get("volume", "")
                issue = item.get("issue", "")
                page = item.get("page", "")

                # Formato APA 7
                formatted = f"{author_str} ({year_str}). {title_str}."
                if source_str:
                    formatted += f" {source_str}"
                    if volume:
                        formatted += f", {volume}"
                        if issue:
                            formatted += f"({issue})"
                    if page:
                        formatted += f", {page}"
                formatted += "."
                if doi:
                    formatted += f" https://doi.org/{doi}"

                # Score de relevancia: el primer autor coincide = alta relevancia
                item_surnames = [a.get("family", "").lower() for a in item_authors]
                query_surnames = [a.lower() for a in authors]
                relevance = "high" if (item_surnames and query_surnames and item_surnames[0] == query_surnames[0]) else "medium"

                results.append({
                    "authors": apa_authors,
                    "year": year_str,
                    "title": title_str,
                    "source": source_str or publisher,
                    "doi": doi,
                    "formatted_apa": formatted,
                    "relevance": relevance,
                })

            if not results:
                return None

            return {
                "candidates": results,
                "found": True,
                "total_results": data.get("message", {}).get("total-results", len(results)),
            }
    except Exception as e:
        print(f"[WARN] Error searching Crossref: {e}")
    return None


async def fetch_openlibrary_metadata(isbn: str) -> Optional[dict]:
    url = f"https://openlibrary.org/api/books?bibkeys=ISBN:{isbn}&format=json&jscmd=data"
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            resp = await client.get(url)
            if resp.status_code == 200:
                data = resp.json()
                key = f"ISBN:{isbn}"
                if key in data:
                    book = data[key]
                    authors = [a["name"] for a in book.get("authors", [])]
                    author_str = ", & ".join(authors) if authors else "Autor desconocido"
                    year = book.get("publish_date", "s.f.")
                    title = book.get("title", "")
                    publisher = book.get("publishers", [{"name": ""}])[0].get("name", "")
                    
                    formatted = f"{author_str} ({year}). {title}. {publisher}."
                    
                    return {
                        "author": author_str,
                        "year": year,
                        "title": title,
                        "publisher": publisher,
                        "isbn": isbn,
                        "formatted": formatted
                    }
    except Exception as e:
        print(f"[WARN] Error in OpenLibrary API: {e}")
    return None

async def resolve_dois_batch(raw_references: List[str]) -> List[dict]:
    """
    Recibe una lista de referencias en texto libre. Extrae DOIs/ISBNs y 
    resuelve los metadatos de forma concurrente usando Crossref u OpenLibrary.
    Retorna lista de diccionarios con los resultados encontrados, manteniendo el índice.
    """
    import asyncio
    
    async def process_ref(idx, text):
        info = extract_doi_or_isbn(text)
        if not info:
            return {"index": idx, "original": text, "resolved": False}
        
        if info["type"] == "doi":
            metadata = await fetch_crossref_metadata(info["value"])
            if metadata:
                return {"index": idx, "original": text, "resolved": True, "type": "doi", "metadata": metadata}
        elif info["type"] == "isbn":
            metadata = await fetch_openlibrary_metadata(info["value"])
            if metadata:
                return {"index": idx, "original": text, "resolved": True, "type": "isbn", "metadata": metadata}
                
        return {"index": idx, "original": text, "resolved": False}

    tasks = [process_ref(idx, text) for idx, text in enumerate(raw_references)]
    results = await asyncio.gather(*tasks)
    return list(results)


# ── Seccion de Referencias ─────────────────────────────────────────────────────

def sort_referencias_alphabetically(references: List[ReferenciaModel]) -> List[ReferenciaModel]:
    """
    Ordena alfabeticamente las referencias por el apellido del primer autor
    o por titulo si no hay autor.
    """
    def get_sort_key(ref: ReferenciaModel) -> str:
        if ref.authors:
            return ref.authors[0].lower()
        elif ref.title:
            return ref.title.lower()
        return ref.raw_text.lower()

    return sorted(references, key=get_sort_key)


def format_apa_referencias_section(
    doc: docx.Document,
    references: List[ReferenciaModel],
    rules: APARuleSet,
) -> None:
    """
    Inserta la seccion de Referencias al final del documento con sangria francesa estricta.
    """
    if not references:
        return

    # Salto de pagina antes de Referencias
    doc.add_page_break()

    # Titulo Nivel 1 Centrado
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.line_spacing = rules.line_spacing
    p_hdr.paragraph_format.first_line_indent = Inches(0)
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(0)

    r_hdr = p_hdr.add_run("Referencias")
    r_hdr.bold = True
    set_run_font(r_hdr, rules.font_family, rules.font_size_pt)
    r_hdr.font.color.rgb = RGBColor(0, 0, 0)

    # Ordenar referencias
    sorted_refs = sort_referencias_alphabetically(references)

    for ref in sorted_refs:
        p_ref = doc.add_paragraph()
        try:
            p_ref.style = None
        except Exception:
            pass
        p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_ref.paragraph_format.line_spacing = rules.line_spacing

        pPr = p_ref._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)

        # Sangria Francesa (Hanging Indent) 1.27 cm
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.first_line_indent = Inches(-0.5)
        p_ref.paragraph_format.space_before = Pt(0)
        p_ref.paragraph_format.space_after = Pt(0)

        # Si tenemos texto crudo formateado o campos individuales
        text: str = ref.formatted_apa if ref.formatted_apa else ref.raw_text

        # Validacion: si text esta vacio, intentar raw_text; si aun vacio, saltar
        if not text or not text.strip():
            text = ref.raw_text
        if not text or not text.strip():
            continue

        text = _strip_ref_prefix(text)

        run = p_ref.add_run(text)
        run.bold = False
        set_run_font(run, rules.font_family, rules.font_size_pt)
        run.font.color.rgb = RGBColor(0, 0, 0)
