"""
WordAPA7 — Módulo de Recreación de Portada UNI (Nicaragua)

Replica la portada institucional de la Universidad Nacional de Ingeniería
tal como aparece en los trabajos reales de los estudiantes:

- Logo UNI (imagen que ya incluye el nombre de la universidad), centrado arriba
- Área de Conocimiento (fuente Butler, 20pt) — se sustituye por Times si no está instalada
- Título del trabajo (Montserrat Black, 20pt)
- Asignatura (Butler, 20pt)
- "Elaborado por" (Montserrat Bold, 11pt)
- Autores en 4 columnas (3 de estudiantes + 1 de docente) con separadores
  verticales negros, Montserrat 10pt
- Fecha (izquierda) y lugar (Managua, Nicaragua)
- Salto de página al final

La portada se INSERTA al inicio del documento (posición 0).
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


def _resolve_logo_path() -> Path:
    meipass = getattr(sys, '_MEIPASS', None)
    candidates = []
    if meipass:
        candidates.append(Path(meipass) / "assets" / "logo_uni.png")  # PyInstaller frozen
    candidates.append(Path(__file__).parent.parent / "assets" / "logo_uni.png")  # desarrollo
    for cand in candidates:
        if cand.exists():
            return cand
    return candidates[0]


LOGO_PATH: Path = _resolve_logo_path()

# Tipografía institucional (idéntica a la de los trabajos UNI reales)
FONT_DEPARTMENT = "Times New Roman"  # Área de Conocimiento / asignatura (20pt, serif)
FONT_TITLE = "Montserrat Black"      # Título del trabajo
FONT_BODY = "Montserrat"             # Autores, fecha, lugar
BLACK = RGBColor(0x00, 0x00, 0x00)

# Conversión de centímetros a "twentieths of a point" (dxa): 1 cm = 567 dxa
_CM_TO_DXA = 567


def _add_run_styled(paragraph, text: str, bold=False, size_pt=12,
                    font: str = FONT_BODY, color: RGBColor = BLACK):
    """Agrega un run con estilo específico a un párrafo."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    # Fijar también la fuente en el atributo eastAsia para respetar el nombre exacto
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    return run


def _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0):
    """Configura el espaciado de un párrafo."""
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


class _UniCoverBuilder:
    """Construye la portada UNI INSERTÁNDOLA al inicio del documento (posición 0)."""

    def __init__(self, doc: Document):
        self.doc = doc
        self._body = doc.element.body
        self._idx = 0

    def add_paragraph(self) -> Paragraph:
        p_elem = OxmlElement('w:p')
        self._body.insert(self._idx, p_elem)
        self._idx += 1
        return Paragraph(p_elem, self.doc)

    def add_page_break(self) -> None:
        p_elem = OxmlElement('w:p')
        r_elem = OxmlElement('w:r')
        br_elem = OxmlElement('w:br')
        br_elem.set(qn('w:type'), 'page')
        r_elem.append(br_elem)
        p_elem.append(r_elem)
        self._body.insert(self._idx, p_elem)
        self._idx += 1

    def insert_table(self, table) -> None:
        """Inserta una tabla en la posición actual del builder."""
        self._body.insert(self._idx, table._element)
        self._idx += 1

    @property
    def paragraph_count(self) -> int:
        return self._idx


def _clear_cell_borders(cell, keep_right: bool = False):
    """Elimina todos los bordes de la celda; opcionalmente conserva el derecho.

    El separador vertical derecho solo se añade cuando *keep_right* es True,
    lo que debe reservarse para celdas que contienen contenido real. Las
    celdas vacías (sin autor) reciben ``keep_right=False`` para que no quede
    un separador colgando sobre el vacío.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    if keep_right:
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')
        right_border.set(qn('w:color'), '000000')
        tcBorders.append(right_border)
    else:
        border = OxmlElement('w:right')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_width(cell, width_cm: float):
    """Fija el ancho de una celda de tabla en centímetros."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * _CM_TO_DXA)))
    tcW.set(qn('w:type'), 'dxa')


def _set_cell_vertical_alignment(cell, align: str = "center"):
    """Fija la alineación vertical de una celda de tabla.

    Valores válidos en OOXML: ``top``, ``center``, ``bottom``.
    Por defecto centra verticalmente para que todas las celdas de una misma
    fila queden alineadas aunque tengan distinta cantidad de texto.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), align)


def _set_cell_padding(cell, top_cm: float = 0.2, bottom_cm: float = 0.2,
                      left_cm: float = 0.2, right_cm: float = 0.2):
    """Agrega relleno (padding) interno a una celda de tabla.

    Evita que el texto quede pegado a las líneas separadoras verticales,
    dando un respiro visual uniforme de 0.2 cm por defecto.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for side, val_cm in [('top', top_cm), ('bottom', bottom_cm),
                         ('left', left_cm), ('right', right_cm)]:
        mar = tcMar.find(qn(f'w:{side}'))
        if mar is None:
            mar = OxmlElement(f'w:{side}')
            tcMar.append(mar)
        mar.set(qn('w:w'), str(int(val_cm * _CM_TO_DXA)))
        mar.set(qn('w:type'), 'dxa')


def _set_row_min_height(row, height_cm: float = 1.5):
    """Fija la altura mínima de una fila de tabla (regla ``atLeast``).

    Garantiza que las filas de autores no colapsen cuando hay poco texto
    (p. ej. un solo estudiante). ``atLeast`` permite que la fila crezca si el
    contenido la rebasa.
    """
    trPr = row._tr.get_or_add_trPr()
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        trHeight = OxmlElement('w:trHeight')
        trPr.append(trHeight)
    trHeight.set(qn('w:val'), str(int(height_cm * _CM_TO_DXA)))
    trHeight.set(qn('w:hRule'), 'atLeast')


def generate_uni_cover(
    doc: Document,
    titulo: str,
    asignatura: str,
    autores: list[dict],   # [{nombre, carnet}, ...]
    tutor: str,
    grupo: str,
    departamento: str = "Área de Conocimiento de Ingeniería y Afines",
    fecha: str = "",
    lugar: str = "Managua, Nicaragua",
) -> int:
    """
    Inserta al inicio del documento una portada institucional UNI fiel al
    formato de los trabajos reales:

        [Logo UNI centrado]
        Área de Conocimiento de Ingeniería y Afines        (Butler 20pt)
        Título del trabajo                                  (Montserrat Black 20pt)
        Asignatura                                          (Butler 20pt)

        Elaborado por                                       (Montserrat Bold 11pt)
        ──────────┬──────────┬──────────┬─────────────────
        │ Autor 1  │ Autor 3  │ Autor 5  │ Tutor           │
        │ Carnet   │ Carnet   │ Carnet   │ Grupo: XXX      │
        │ Autor 2  │ Autor 4  │          │                 │
        │ Carnet   │ Carnet   │          │                 │
        ──────────┴──────────┴──────────┴─────────────────

        25 de junio del año 2025
        Managua, Nicaragua

    Returns: Número de párrafos insertados (para cover_paragraph_count).
    """
    if not fecha:
        from datetime import date
        d = date.today()
        months = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                  "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        fecha = f"{d.day} de {months[d.month-1]} del año {d.year}"

    builder = _UniCoverBuilder(doc)

    # ── 1. Logo UNI centrado (la imagen incluye el nombre de la universidad) ──
    logo_p = builder.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(logo_p, before=0, after=6)
    if LOGO_PATH.exists():
        run_logo = logo_p.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=Cm(5.2))

    # ── 2. Departamento (Butler 20pt, centrado) ──────────────────────────────
    dept_p = builder.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_styled(dept_p, departamento, bold=False, size_pt=20, font=FONT_DEPARTMENT)
    _set_paragraph_spacing(dept_p, before=6, after=30)

    # ── 3. Título (Montserrat Black 20pt, centrado) ──────────────────────────
    titulo_p = builder.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_styled(titulo_p, titulo, bold=False, size_pt=20, font=FONT_TITLE)
    _set_paragraph_spacing(titulo_p, before=0, after=30)

    # ── 4. Asignatura (Butler 20pt, centrado) ────────────────────────────────
    if asignatura:
        asig_p = builder.add_paragraph()
        asig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run_styled(asig_p, asignatura, bold=False, size_pt=20, font=FONT_DEPARTMENT)
        _set_paragraph_spacing(asig_p, before=0, after=60)

    # ── 5. "Elaborado por" (Montserrat Bold 11pt, izquierda) ─────────────────
    elab_p = builder.add_paragraph()
    elab_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run_styled(elab_p, "Elaborado por", bold=True, size_pt=11, font=FONT_BODY)
    _set_paragraph_spacing(elab_p, before=0, after=10)

    # ── 6. Autores en columnas con separadores verticales negros ─────────────
    estudiantes = [a for a in autores if not a.get("es_tutor", False)]

    # Distribuir en columnas de estudiantes (la última es el docente).
    # Columnas dinámicas: pocos estudiantes = pocas columnas, para no dejar
    # columnas vacías con separadores (portada rota con 1-2 integrantes).
    n_student_cols = min(3, max(1, len(estudiantes)))
    cols_students: list[list[dict]] = [[] for _ in range(n_student_cols)]
    for i, est in enumerate(estudiantes):
        cols_students[i % n_student_cols].append(est)

    tutor_entry = None
    if tutor:
        tutor_entry = {"nombre": tutor, "carnet": grupo if grupo else "", "es_tutor": True}

    n_cols = n_student_cols + (1 if tutor_entry else 0)
    max_rows = max((len(c) for c in cols_students), default=0)
    if tutor_entry:
        max_rows = max(max_rows, 1)
    max_rows = max(max_rows, 1)

    if n_cols > 0:
        # ── 6.1 Anchos de columna balanceados según nº de estudiantes ──────────
        # Pocos integrantes → columnas más anchas y equilibrio con el docente.
        #   3 estudiantes → 3.5 cm cada uno  + 5.0 cm docente  (total 15.5 cm)
        #   2 estudiantes → 5.0 cm cada uno  + 5.0 cm docente  (total 15.0 cm)
        #   1 estudiante  → 7.0 cm           + 5.0 cm docente  (total 12.0 cm)
        student_col_width = {3: 3.5, 2: 5.0, 1: 7.0}.get(n_student_cols, 3.5)
        tutor_col_width = 5.0
        col_widths_cm = ([student_col_width] * n_student_cols) + \
                        ([tutor_col_width] if tutor_entry else [])

        authors_table = doc.add_table(rows=max_rows, cols=n_cols)
        authors_table.autofit = False

        for row_i in range(max_rows):
            # Altura mínima de fila para que las celdas no colapsen.
            _set_row_min_height(authors_table.rows[row_i], height_cm=1.5)

        for col_i in range(n_cols):
            is_tutor_col = (col_i == n_student_cols and tutor_entry)
            col_data = cols_students[col_i] if not is_tutor_col else [tutor_entry]

            for row_i in range(max_rows):
                cell = authors_table.cell(row_i, col_i)
                _set_cell_width(cell, col_widths_cm[col_i])

                # Centrar verticalmente todas las celdas de la misma fila.
                _set_cell_vertical_alignment(cell, align="center")

                # Padding interno para que el texto no toque las líneas.
                _set_cell_padding(cell, top_cm=0.2, bottom_cm=0.2,
                                   left_cm=0.2, right_cm=0.2)

                # ── 6.4 Separador vertical solo en celdas con contenido ─────────
                # Una celda vacía (sin autor en esa fila) NO debe mostrar el
                # separador derecho; de lo contrario queda una línea colgando
                # sobre el vacío y la portada se ve rota.
                has_content = row_i < len(col_data)
                keep_right = (col_i < n_cols - 1) and has_content
                _clear_cell_borders(cell, keep_right=keep_right)

                if has_content:
                    autor = col_data[row_i]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)

                    nombre = autor.get("nombre", "")
                    carnet = autor.get("carnet", "")

                    if is_tutor_col:
                        _add_run_styled(p, nombre, bold=True, size_pt=10, font=FONT_BODY)
                        if carnet:
                            p2 = cell.add_paragraph()
                            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            _set_paragraph_spacing(p2, before=0, after=0, line_spacing=1.1)
                            _add_run_styled(p2, f"Grupo: {carnet}", bold=True, size_pt=10, font=FONT_BODY)
                    else:
                        _add_run_styled(p, nombre, bold=False, size_pt=10, font=FONT_BODY)
                        if carnet:
                            p2 = cell.add_paragraph()
                            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            _set_paragraph_spacing(p2, before=0, after=0, line_spacing=1.1)
                            _add_run_styled(p2, carnet, bold=False, size_pt=10, font=FONT_BODY)

        builder.insert_table(authors_table)

    # ── 6.5 Espaciador dinámico: empuja fecha/lugar hacia abajo de la hoja ──
    # Pocos autores = más espacio, para que la portada NO quede pegada arriba
    # con todo el vacío abajo (replica el flex:1 del preview).
    PAGE_USABLE_CM = 24.6    # A4 29.7cm - márgenes 2.54cm ×2
    TOP_BLOCK_CM = 8.0       # logo + dept + título + asignatura + "Elaborado por"
    AUTH_ROW_CM = 1.5        # altura mínima por fila de autores (con padding)
    DATE_BLOCK_CM = 1.4      # fecha + lugar
    TARGET_FILL = 0.86       # la fecha queda cerca del 86% del alto de la hoja
    used_cm = TOP_BLOCK_CM + max(1, max_rows) * AUTH_ROW_CM + DATE_BLOCK_CM
    spacer_cm = max(0.0, PAGE_USABLE_CM * TARGET_FILL - used_cm)
    empty_lines = int(spacer_cm / 0.5)
    for _ in range(empty_lines):
        sp = builder.add_paragraph()
        _set_paragraph_spacing(sp, before=0, after=0, line_spacing=1.0)

    # ── 7. Fecha (izquierda) ─────────────────────────────────────────────────
    fecha_p = builder.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run_styled(fecha_p, fecha, bold=False, size_pt=11, font=FONT_BODY)
    _set_paragraph_spacing(fecha_p, before=24, after=2)

    # ── 8. Lugar (Managua, Nicaragua) ────────────────────────────────────────
    lugar_p = builder.add_paragraph()
    lugar_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run_styled(lugar_p, lugar, bold=False, size_pt=11, font=FONT_BODY)
    _set_paragraph_spacing(lugar_p, before=0, after=0)

    # ── 9. Salto de página ───────────────────────────────────────────────────
    builder.add_page_break()

    return builder.paragraph_count
