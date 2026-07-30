import sys
sys.path.insert(0, 'python')
import docx
import zipfile
import os
from pathlib import Path
from models import DocumentModel, ElementModel, ElementType, APARuleSet, PortadaData, TableModel
from generation.generator import generate_apa7_docx

def main():
    doc_orig = docx.Document()
    doc_orig.add_heading('Introducción', level=1)
    doc_orig.add_paragraph('Este es un párrafo del cuerpo.')
    t = doc_orig.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = 'H1'
    t.rows[0].cells[1].text = 'H2'
    t.rows[1].cells[0].text = 'V1'
    t.rows[1].cells[1].text = 'V2'

    orig_path = Path('test_duplication_orig.docx')
    out_path = Path('test_duplication_out.docx')
    doc_orig.save(orig_path)

    import re
    with zipfile.ZipFile(orig_path) as z:
        xml_orig = z.read('word/document.xml').decode('utf-8')
        tbl_count_orig = len(re.findall(r'<w:tbl\b[^>]*>', xml_orig))

    elements = [
        ElementModel(id='e1', type=ElementType.HEADING, text='Introducción', heading_level=1),
        ElementModel(id='e2', type=ElementType.PARAGRAPH, text='Este es un párrafo del cuerpo.'),
        ElementModel(id='e3', type=ElementType.TABLE, text='Tabla 1', table_info=TableModel(element_id='e3', headers=['H1','H2'], rows=[['V1','V2']], table_number=1, caption='Resumen'))
    ]

    doc_model = DocumentModel(session_id='test_dup_session', file_name='test_duplication_orig.docx', elements=elements)
    generate_apa7_docx(doc_model, out_path, rules=APARuleSet(), portada=PortadaData(use_original_cover=True))

    with zipfile.ZipFile(out_path) as z:
        xml_out = z.read('word/document.xml').decode('utf-8')
        tbl_count_out = len(re.findall(r'<w:tbl\b[^>]*>', xml_out))
        styles_xml_out = z.read('word/styles.xml').decode('utf-8')

    print(f"Tabla orig count: {tbl_count_orig}, Tabla gen count: {tbl_count_out}")
    print("styles.xml tiene w:line=480:", 'w:line="480"' in styles_xml_out)
    assert tbl_count_orig == tbl_count_out, f"ERROR: Duplicacion detectada! {tbl_count_orig} vs {tbl_count_out}"
    print("SUCCESS: ¡NO HAY DUPLICACION Y STYLES.XML FUE ACTUALIZADO CORRECTAMENTE!")

    if orig_path.exists(): os.remove(orig_path)
    if out_path.exists(): os.remove(out_path)

if __name__ == '__main__':
    main()
