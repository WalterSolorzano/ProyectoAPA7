"""
WordAPA7 — Layered DOCX Generator (From-Scratch Pipeline)

Creates a clean APA 7 DOCX from DocumentModel without modifying the original.
Eliminates bugs from: paragraph matching, style inheritance, numbering conflicts,
OOXML artifact preservation issues.
"""

from pathlib import Path
from typing import List

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from models import (
    APARuleSet,
    DocumentModel,
    ElementModel,
    ElementType,
    PortadaData,
    ReferenciaModel,
)
from modules.portada_module import format_apa_portada
from modules.referencias_module import format_apa_referencias_section

from generation.bullet_engine import format_bullet_item, format_numbered_item
from generation.document_structure import setup_apa_header
from generation.image_handler import format_apa_figure
from generation.style_engine import (
    apply_page_setup,
    format_block_quote,
    format_heading_paragraph,
    format_normal_paragraph,
    normalize_global_body_spacing,
    update_docx_styles_xml,
)
from generation.table_engine import format_apa_table


def generate_apa7_from_scratch(
    doc_model: DocumentModel,
    output_filepath: Path | str,
    rules: APARuleSet | None = None,
    portada: PortadaData | None = None,
    references: List[ReferenciaModel] | None = None,
) -> Path:
    """
    Generate a clean APA 7 DOCX from DocumentModel (no original file needed).

    Three-phase pipeline:
        1. Model Preparation  — validate and normalize elements
        2. Layer 1: Skeleton — page setup, sections, headers
        3. Layer 2: Content  — cover page + body elements (headings, paragraphs,
            lists, tables, images, equations, quotes)
        4. Layer 3: Polish   — references, global spacing normalization
    """
    if rules is None:
        rules = APARuleSet()

    out_path = Path(output_filepath)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # ═══════════════════════════════════════════════════════════════════
    # PHASE 1: Model Preparation — filter out cover, empty, break elements
    # ═══════════════════════════════════════════════════════════════════

    body_elements: list[ElementModel] = []
    for item in doc_model.elements:
        elem = ElementModel.model_validate(item) if isinstance(item, dict) else item
        if elem.type == ElementType.EMPTY:
            continue
        if elem.is_cover_section or elem.type == ElementType.PORTADA_BLOCK:
            continue
        body_elements.append(elem)

    # ═══════════════════════════════════════════════════════════════════
    # PHASE 2: Skeleton — fresh document + page setup + header
    # ═══════════════════════════════════════════════════════════════════

    doc = docx.Document()
    apply_page_setup(doc, rules)
    update_docx_styles_xml(doc, rules)

    running_head = portada.running_head if portada else ""
    setup_apa_header(doc, doc_model.apa_format, running_head, rules)

    # ═══════════════════════════════════════════════════════════════════
    # PHASE 3: Content — cover page + body elements
    # ═══════════════════════════════════════════════════════════════════

    # 3a. Cover page (APA 7 synthetic)
    if portada:
        format_apa_portada(doc, portada, rules)

    # 3b. Cover designer template
    cover_template_applied = False
    if doc_model.portada:
        cover_template_id = None
        if isinstance(doc_model.portada, dict):
            cover_template_id = doc_model.portada.get('cover_template_id')
        elif hasattr(doc_model.portada, 'get'):
            cover_template_id = doc_model.portada.get('cover_template_id')

        if cover_template_id:
            try:
                from modules.cover_designer import apply_cover_to_document, list_cover_templates
                base_dir = Path(__file__).resolve().parent.parent
                templates = list_cover_templates(base_dir)
                selected = None
                for t in templates:
                    if t.name == cover_template_id:
                        selected = t
                        break
                if selected:
                    apply_cover_to_document(
                        doc, selected, base_dir,
                        portada_title=portada.title if portada else "",
                        portada_author=portada.author if portada else "",
                        portada_institution=portada.institution if portada else "",
                        portada_course=portada.course if portada else "",
                        portada_instructor=portada.instructor if portada else "",
                        portada_date=portada.date if portada else "",
                    )
                    cover_template_applied = True
            except Exception as e:
                print(f"[LAYERED-GEN] Cover designer skip: {e}")

    # 3c. Body elements
    heading_counters: dict[int, int] = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    numbered_counters: dict[int, int] = {1: 0, 2: 0, 3: 0}
    last_numbered_level = 0

    def _to_roman(num: int) -> str:
        if num <= 0:
            return str(num)
        val = [50, 40, 10, 9, 5, 4, 1]
        syms = ["L", "XL", "X", "IX", "V", "IV", "I"]
        roman = ""
        for i, v in enumerate(val):
            while num >= v:
                roman += syms[i]
                num -= v
        return roman

    def _build_prefix(counters, level: int, style: str) -> str:
        if style == 'none' or level > 2:
            return ""
        parts = []
        for l in range(1, level + 1):
            c = counters.get(l, 0)
            if c > 0:
                if style == 'roman' and l == 1:
                    parts.append(_to_roman(c))
                else:
                    parts.append(str(c))
        return ".".join(parts) + ". " if parts else ""

    for elem in body_elements:
        elem_type = elem.type
        elem_level = elem.heading_level or 1
        list_lvl = elem.list_level or 1

        if elem_type == ElementType.PAGE_BREAK:
            doc.add_page_break()
            continue

        if elem_type == ElementType.HEADING:
            lvl = elem.heading_level or 1
            # Reset numbered list counters on heading
            if lvl <= 2:
                numbered_counters = {1: 0, 2: 0, 3: 0}
                last_numbered_level = 0

            heading_counters[lvl] = heading_counters.get(lvl, 0) + 1
            for lower in range(lvl + 1, 6):
                heading_counters[lower] = 0

            level_style = getattr(rules, f'heading_numbering_style_lvl{lvl}', 'decimal')
            prefix = _build_prefix(heading_counters, lvl, level_style)

            # Page break before H1
            if lvl == 1:
                doc.add_page_break()

            p = doc.add_paragraph()
            format_heading_paragraph(p, lvl, prefix + (elem.text or ""), rules)

        elif elem_type == ElementType.PARAGRAPH:
            numbered_counters = {1: 0, 2: 0, 3: 0}
            last_numbered_level = 0
            p = doc.add_paragraph()
            format_normal_paragraph(p, elem.text or "", rules)

        elif elem_type == ElementType.BLOCK_QUOTE:
            numbered_counters = {1: 0, 2: 0, 3: 0}
            p = doc.add_paragraph()
            format_block_quote(p, elem.text or "", rules)

        elif elem_type == ElementType.BULLET:
            numbered_counters = {1: 0, 2: 0, 3: 0}
            p = doc.add_paragraph()
            bullet_style = (
                rules.bullet_style_level1 if list_lvl == 1
                else rules.bullet_style_level2 if list_lvl == 2
                else rules.bullet_style_level3
            )
            format_bullet_item(p, elem.text or "", level=list_lvl, rules=rules,
                             bullet_style=bullet_style, is_bold=elem.is_bold,
                             is_italic=elem.is_italic)

        elif elem_type == ElementType.NUMBERED_LIST:
            num_lvl = list_lvl if list_lvl in (1, 2, 3) else 1
            if num_lvl <= last_numbered_level:
                for l in range(num_lvl + 1, 4):
                    numbered_counters[l] = 0
            numbered_counters[num_lvl] = numbered_counters.get(num_lvl, 0) + 1
            last_numbered_level = num_lvl
            num_style = (
                rules.number_style_level1 if num_lvl == 1
                else rules.number_style_level2 if num_lvl == 2
                else rules.number_style_level3
            )
            p = doc.add_paragraph()
            format_numbered_item(p, elem.text or "", level=num_lvl,
                               index=numbered_counters[num_lvl], rules=rules,
                               num_style=num_style, is_bold=elem.is_bold,
                               is_italic=elem.is_italic)

        elif elem_type == ElementType.TABLE:
            if elem.table_info:
                format_apa_table(doc, elem.table_info, rules)

        elif elem_type == ElementType.IMAGE:
            if elem.image_info:
                format_apa_figure(doc, elem.image_info, rules)

        elif elem_type == ElementType.EQUATION:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = rules.line_spacing
            eq_cfg = getattr(elem, 'equation', None)
            if eq_cfg is not None:
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                }
                p.paragraph_format.alignment = align_map.get(eq_cfg.alignment or "center", WD_ALIGN_PARAGRAPH.CENTER)
            if elem.text:
                r = p.add_run(elem.text)
                r.font.name = rules.font_family
                r.font.size = Pt(rules.font_size_pt)

        elif elem_type == ElementType.TOC:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("Indice / Tabla de Contenidos")
            r.bold = True
            r.font.name = rules.font_family
            r.font.size = Pt(rules.font_size_pt)

    # ═══════════════════════════════════════════════════════════════════
    # PHASE 4: Polish — references + global spacing
    # ═══════════════════════════════════════════════════════════════════

    if references:
        format_apa_referencias_section(doc, references, rules)

    normalize_global_body_spacing(doc, rules, 0)

    doc.save(out_path)
    return out_path
