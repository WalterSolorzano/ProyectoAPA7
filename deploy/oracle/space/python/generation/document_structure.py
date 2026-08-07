"""
WordAPA7 — Estructura de Documento y Encabezados (Running Head)

Maneja los elementos de nivel de documento:
- Numeración de páginas en la esquina superior derecha (mediante fldChar PAGE en XML de Word).
- Running Head (Título corto en mayúsculas a la izquierda en modo Profesional).
- Salto de página previo a secciones mayores (Referencias, Portada).
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor
from models import APAFormat, APARuleSet


def add_page_number_field(run):
    """
    Inserta un campo XML de Word `PAGE` dentro del run para numeración dinámica.
    """
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _clean_existing_page_numbers(doc: docx.Document):
    """
    Elimina cualquier campo PAGE / NUMPAGES existente en los headers y footers
    de todas las secciones del documento.  Esto evita que el numero de pagina
    APA se superponga a uno ya presente en el documento original (doble numero).
    Es seguro llamarlo aunque el documento no tenga headers (no-op).
    """
    for section in doc.sections:
        for attr in ('header', 'first_page_header', 'even_page_header',
                     'footer', 'first_page_footer', 'even_page_footer'):
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            try:
                element = hf._element
            except (AttributeError, Exception):
                continue

            for p in element.findall(f'.//{qn("w:p")}'):
                should_clean = False
                for instr in p.findall(f'.//{qn("w:instrText")}'):
                    text = (instr.text or '').upper()
                    if 'PAGE' in text or 'NUMPAGES' in text:
                        should_clean = True
                        break
                if should_clean:
                    # Elimina todos los runs del parrafo que contienen el field
                    for r in list(p.findall(qn("w:r"))):
                        p.remove(r)


def setup_apa_header(doc: docx.Document, format_type: APAFormat, running_head_text: str, rules: APARuleSet):
    """
    Configura el encabezado de las páginas según el tipo APA (Estudiante vs Profesional).
    - Estudiante: Número de página a la derecha.
    - Profesional: Running head en mayúsculas a la izquierda + Número de página a la derecha.
    """
    # Clean existing page number fields to prevent double numbering
    _clean_existing_page_numbers(doc)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False

    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if format_type == APAFormat.PROFESSIONAL and running_head_text:
        # Running Head a la izquierda + Tabulación a la derecha + Número de página
        r_rh = p.add_run(f"{running_head_text.upper()}\t")
        r_rh.font.name = rules.font_family
        r_rh.font.size = Pt(rules.font_size_pt)
        r_rh.font.color.rgb = RGBColor(0, 0, 0)

    # Run para el número de página
    r_pg = p.add_run()
    r_pg.font.name = rules.font_family
    r_pg.font.size = Pt(rules.font_size_pt)
    r_pg.font.color.rgb = RGBColor(0, 0, 0)
    add_page_number_field(r_pg)
