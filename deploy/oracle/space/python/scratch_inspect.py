from pathlib import Path

import docx

doc_path = Path(r"C:\Users\--X\.gemini\antigravity\scratch\wordapa7\10mo Trabajo Contabilidad_APA7.docx")
print("Cargando:", doc_path)

doc = docx.Document(str(doc_path))

print("=== PARRAFOS ===")
for idx, p in enumerate(doc.paragraphs[:35]):
    print(f"P{idx}: '{p.text.strip()}' | estilo: {p.style.name} | runs: {len(p.runs)}")

print("\n=== DIBUJOS / SHAPES / TEXTBOXES ===")
for idx, p in enumerate(doc.paragraphs[:35]):
    for r in p.runs:
        drawings = r._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        txbx = r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent')
        if drawings or txbx:
            print(f"P{idx} tiene drawings: {len(drawings)}, txbxContent: {len(txbx)}, texto p: '{p.text.strip()}'")
            if txbx:
                for t in txbx:
                    print("   TEXTBOX CONTENIDO:", [tp.text for tp in t.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')])
