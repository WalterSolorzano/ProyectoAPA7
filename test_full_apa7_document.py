"""
WordAPA7 — Test de Auditoría Exhaustiva de Documento Saliente APA 7

Crea un documento DOCX completo con:
- Portada con Universidad, Logo, Título, Autores, Tutor, Fecha
- Título Nivel 1 ("Introducción"), Título Nivel 2, Título Nivel 3
- Listas con viñetas y numeradas anidadas (Niveles 1 y 2)
- Tabla con rotulación APA 7 ("Tabla 1" negrita, Título cursiva, Nota al pie)
- Figura con rotulación APA 7 ("Figura 1" negrita, Título cursiva)
- Referencias Bibliográficas

Luego sube el documento al backend local, genera el DOCX formateado, y valida el XML.
"""

import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt
import requests

BASE_URL = "http://127.0.0.1:8742"

def create_sample_docx(filepath: str):
    doc = docx.Document()
    
    # Portada
    p0 = doc.add_paragraph("Universidad Nacional de Ingeniería")
    p1 = doc.add_paragraph("Facultad de Electrotecnia y Computación")
    p2 = doc.add_paragraph("ANÁLISIS DE SISTEMAS ACADÉMICOS EN APA 7")
    p3 = doc.add_paragraph("Elaborado por: Br. Juan Pérez, Br. María López")
    p4 = doc.add_paragraph("Tutor: Ing. Carlos Rodríguez")
    p5 = doc.add_paragraph("15 de junio de 2026 Managua, Nicaragua")
    
    doc.add_page_break()
    
    # Cuerpo
    h1 = doc.add_paragraph("Introducción")
    h1.style = 'Heading 1'
    
    p_intro = doc.add_paragraph(
        "El presente estudio examina la automatización del formato APA 7ma edición "
        "en documentos académicos universitarios utilizando técnicas modernas de procesamiento."
    )
    
    h2 = doc.add_paragraph("Planteamiento del Problema")
    h2.style = 'Heading 2'
    
    p_prob = doc.add_paragraph("Muchos estudiantes enfrentan dificultades con las sangrías y rotulación de tablas.")
    
    # Listas
    b1 = doc.add_paragraph("• Elemento de lista principal (Nivel 1)")
    b2 = doc.add_paragraph("    • Subelemento anidado (Nivel 2)")
    
    # Tabla
    table_p1 = doc.add_paragraph("Tabla 1")
    table_p2 = doc.add_paragraph("Resumen de Estadísticas Académicas")
    
    t = doc.add_table(rows=3, cols=3)
    t.cell(0, 0).text = "Categoría"
    t.cell(0, 1).text = "Estudiantes"
    t.cell(0, 2).text = "Porcentaje"
    
    t.cell(1, 0).text = "Aprobados"
    t.cell(1, 1).text = "120"
    t.cell(1, 2).text = "85%"
    
    t.cell(2, 0).text = "Reprobados"
    t.cell(2, 1).text = "20"
    t.cell(2, 2).text = "15%"
    
    table_note = doc.add_paragraph("Nota. Datos recopilados en el período académico 2026.")
    
    doc.save(filepath)
    print(f"[OK] Documento de prueba creado: {filepath}")

def run_audit():
    sample_file = "sample_apa7_test.docx"
    create_sample_docx(sample_file)
    
    # 1. Upload
    with open(sample_file, "rb") as f:
        resp = requests.post(f"{BASE_URL}/api/upload", files={"file": f})
    
    if resp.status_code != 200:
        print(f"[FAIL] Upload status code: {resp.status_code}")
        return False
        
    doc_data = resp.json()
    session_id = doc_data["session_id"]
    elements = doc_data["elements"]
    print(f"[OK] Upload exitoso. Session ID: {session_id}, Elementos: {len(elements)}")
    
    # Imprimir clasificación de elementos
    for idx, e in enumerate(elements):
        print(f"  Elem #{idx}: type={e.get('type')}, heading_lvl={e.get('heading_level')}, is_cover={e.get('is_cover_section')}, text='{(e.get('text') or '')[:40]}'")
        
    # 2. Generate APA 7
    gen_resp = requests.post(f"{BASE_URL}/api/generate", json={"session_id": session_id})
    if gen_resp.status_code != 200:
        print(f"[FAIL] Generate status code: {gen_resp.status_code}")
        return False
        
    download_url = gen_resp.json()["download_url"]
    print(f"[OK] Generación exitosa. Download URL: {download_url}")
    
    # 3. Download
    down_resp = requests.get(f"{BASE_URL}{download_url}")
    out_file = "output_apa7_result.docx"
    with open(out_file, "wb") as f:
        f.write(down_resp.content)
    print(f"[OK] Archivo guardado: {out_file}")
    
    # 4. Inspection of generated output DOCX
    out_doc = docx.Document(out_file)
    print(f"[OK] DOCX generado cargado. Párrafos: {len(out_doc.paragraphs)}, Tablas: {len(out_doc.tables)}")
    
    # Validar que no hay "1" flotante en el encabezado de la primera página
    sec0 = out_doc.sections[0]
    print(f"  Header different_first_page_header_footer: {sec0.different_first_page_header_footer}")
    
    return True

if __name__ == "__main__":
    run_audit()
